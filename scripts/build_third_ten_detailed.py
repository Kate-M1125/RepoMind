from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches,Pt,RGBColor

OUT=Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第三批10题.docx')
BLUE=RGBColor(46,116,181);DARK=RGBColor(31,77,120);GRAY=RGBColor(95,103,112)
QA=[
('21. 讲一下 RepoMind 项目里 RAG 的完整流程',[
'我的回答是：RepoMind 的 RAG 不是“把 Issue 向量化后查一次 ChromaDB”，而是从仓库预处理、多路召回、融合精排到补充结构化证据的一条完整链路。目标不是直接让检索结果替代分析，而是先从大型仓库中缩小候选范围，再交给 ReAct Agent 继续验证。',
'索引阶段，系统先 clone 仓库，遍历可分析文件并做代码切片。Python 代码优先按 AST 中的函数和类保留语义边界，非 Python 文件使用通用切分；每个 Chunk 保存文件路径、语言、行号、imports 等元数据。随后使用 Sentence Transformer 生成 embedding，写入以仓库为单位管理的 ChromaDB Collection。',
'查询阶段位于 workflow/nodes/retrieve.py。系统并行收集三类上下文，其中 RAG 分支会构造 HyDE 假设文档、Issue 原始标题、问题类型加正文、堆栈信息等多路 Query。每一路先召回 top-5，之后使用 RRF 根据排名融合得到大约 top-10 候选，再用 cross-encoder/ms-marco-MiniLM-L-6-v2 对 Query 和代码块做成对打分，最终保留 top-3。',
'精排后的结果还不是全部上下文。系统会根据 Chunk metadata 中的 imports 检索相关依赖源码，并结合轻量调用图补充上下游关系。最后 RAG 结果与 entry 分支的精确 grep 结果、regression 场景的 Git 历史在 merge_context 中合并。后续分析 Agent 仍可以用 read_file、find_definition 和 find_callers 继续取证。'
]),
('22. 为什么不能只把 Issue 和代码直接交给 LLM？为什么还需要 RAG？',[
'我的回答是：主要有三个原因。第一，仓库代码通常远大于模型上下文窗口，不可能把所有文件完整放进去；第二，即使上下文装得下，大量无关代码也会稀释真正证据，增加 Token、延迟和注意力噪声；第三，Issue 描述与代码实现之间经常存在词汇差异，用户描述的是现象，真正代码里可能是另一个内部函数或配置项。',
'RAG 的作用是先建立候选集合，让模型优先看到最可能相关的文件和函数。RepoMind 使用多查询而不是只用 Issue 正文，是因为标题、类型、堆栈和 HyDE 各自覆盖不同信号；之后再用 CrossEncoder 精排，避免仅凭 embedding 相似度把“主题相近但与根因无关”的代码放到最前面。',
'但我也不会把 RAG 说成万能的。向量检索擅长发现语义相关代码，却可能漏掉精确符号、配置项和调用关系，所以项目另外保留 entry+grep、定义查找、调用者追踪和 Git 历史。我的理解是：RAG 负责快速缩小搜索空间，工具负责验证与扩展证据，LLM 负责综合推理。',
'如果只做 RAG 后直接生成，模型容易把候选代码当成完整事实；如果只让 Agent 从空白开始全库搜索，又会浪费大量工具轮次。RepoMind 把两者结合，是在召回效率和分析可靠性之间做平衡。'
]),
('23. 代码是如何切片的？为什么不能固定按字符数切分？',[
'我的回答是：代码切片首先要保证语义完整，其次才是控制 Chunk 大小。固定按字符数或固定行数切分虽然简单，但可能把函数签名、函数体和返回逻辑切到不同块里，也可能让类方法失去所属类上下文，最终召回到的片段无法独立理解。',
'RepoMind 对 Python 代码优先使用 AST 识别函数和类，以语法节点作为语义边界；Chunk 会保留文件路径、起止行、语言和 imports 等元数据。对于过长的函数或无法用 AST 解析的文件，再使用通用递归切分进行兜底。其他语言目前没有完整 AST/LSP 语义解析，更多依赖文本切分，这是项目当前的能力边界。',
'Chunk 太小会导致上下文碎片化：模型只看到一条语句，却不知道参数来源和调用关系；Chunk 太大又会降低向量表示的聚焦程度，召回后也浪费 Token。因此合理策略是以函数、类或配置段为基本单元，对超长单元二次拆分，并保留少量 overlap 处理跨边界信息。',
'评估切片不能只看平均长度。我会检查关键函数是否能作为完整块被召回、文件路径和行号是否正确、同一代码是否产生大量重复 Chunk，以及最终 Recall@K 和根因定位是否提升。切片策略最终要为检索和后续工具导航服务。'
]),
('24. HyDE 是什么？为什么代码检索也使用 HyDE？',[
'我的回答是：HyDE 全称 Hypothetical Document Embeddings。它不是直接拿用户 Query 去检索，而是先让模型根据问题生成一段“假设中的相关文档或代码说明”，再用这段假设文本做 embedding 检索。它的核心价值是弥补用户问题和目标文档之间的表达鸿沟。',
'Issue 通常描述的是行为现象，例如“某个参数传入子路由后失效”，代码中真正出现的可能是 include_router、default 参数或内部合并逻辑。如果只对 Issue 原文做向量检索，现象词和实现词不一定重合。HyDE 可以把现象改写成更接近代码结构的假设，例如可能涉及哪个模块、函数和数据传递过程。',
'RepoMind 没有只依赖 HyDE，而是把它作为多查询中的一路。Issue 标题保留用户最集中的描述，类型加正文保留完整业务语义，堆栈提供精确符号，HyDE 提供可能的实现语言。多路结果通过 RRF 融合，可以降低某一次模型假设错误造成的影响。',
'HyDE 的风险是模型可能编造不存在的函数名。如果把 HyDE 当答案使用会很危险，所以项目只把它作为召回 Query，不把假设内容直接当证据。最终结论必须回到真实仓库文件、工具输出和 Git 记录上验证。'
]),
('25. 为什么要做多查询召回？RRF 是怎么融合的？',[
'我的回答是：一个 Issue 往往同时包含现象、对象、错误信息和版本信息，单一 Query 很难覆盖所有检索意图。多查询召回就是从不同角度表达同一个问题，让每一路检索发现不同候选，再进行统一融合。',
'RepoMind 使用 HyDE、Issue 标题、类型加正文以及堆栈信息等查询。标题短而聚焦，正文信息完整但噪声较多，堆栈对符号定位很强，HyDE 尝试把问题翻译成实现语言。每一路先独立取 top-5，这样不会因为某个长 Query 的 embedding 主导全部候选。',
'融合阶段使用 Reciprocal Rank Fusion。它不直接相加不同查询的相似度，而是根据某个文档在各结果列表中的排名累计分数，常见形式是 1/(k+rank)。同一个 Chunk 如果在多路查询中都排名靠前，融合得分会更高；只在某一路偶然出现的结果则相对靠后。',
'选择 RRF 的原因是不同 Query 的相似度分布不完全可比，直接加分需要复杂归一化。RRF 简单、稳定，也适合没有大规模标注数据时使用。它的局限是只使用排名，忽略原始分数差距，所以 RepoMind 在融合后还使用 CrossEncoder 做更精细的相关性判断。'
]),
('26. 为什么召回后还要用 CrossEncoder Rerank？一般保留几个代码块？',[
'我的回答是：向量召回使用 Bi-Encoder，把 Query 和代码块分别编码后计算向量相似度，速度快，适合从大量 Chunk 中找候选；但它在编码时看不到 Query 与代码之间逐词交互，容易把主题相似但不能解释当前问题的代码排在前面。',
'CrossEncoder 会把 Query 和候选代码作为一对输入，让模型同时阅读二者并输出相关性分数。它的精度通常更高，但每个候选都要单独前向计算，成本不适合全库使用，所以最合理的方式是先快速召回，再对少量候选精排。',
'RepoMind 的实现是多查询每路 top-5，经 RRF 融合后取大约 top-10，再使用 cross-encoder/ms-marco-MiniLM-L-6-v2 精排，最后保留 top-3 进入核心上下文。top-3 不是通用真理，而是这个项目在上下文质量和 Token 成本之间采用的配置。',
'如果 top-k 太小，关键证据可能被截断；太大则会带入相似但无关的文件。调参时应该看关键文件 Recall@K、精排后的命中顺序、最终根因准确性和 Token 消耗，而不是只追求某个固定数字。对于调用链分散的问题，还需要让工具继续读取关联函数，不能指望三个 Chunk 覆盖完整根因。'
]),
('27. BM25、向量检索和混合检索有什么区别？RepoMind 为什么还需要 grep？',[
'我的回答是：BM25 是稀疏关键词检索，依赖词频、逆文档频率和长度归一化，适合函数名、异常文本、配置键等精确词匹配；向量检索把 Query 和文档映射到稠密语义空间，适合表达不同但语义相近的内容。两者解决的问题不同。',
'代码场景中，用户可能准确给出 generate_unique_id_function，这时关键词或 grep 非常有效；也可能只描述“子路由没有继承 ID 生成策略”，这时向量检索更容易发现相关实现。混合检索通常同时运行 BM25 和向量召回，再通过 RRF 或学习排序融合。',
'RepoMind 当前核心管线实现了多查询向量检索、RRF 和 CrossEncoder，并通过 grep_code、find_definition 等工具补足精确检索能力。项目没有完整实现独立 BM25 索引，所以面试时不能把 BM25 说成已经落地。可以说目前使用向量语义召回加 grep 符号检索，后续可把 BM25 纳入统一候选融合。',
'grep 仍然必要，因为 RAG 返回的是概率相关结果，不能保证所有精确标识符都在 top-k 中。尤其是入口函数、错误码和配置项，文本搜索更直接。RepoMind 的 entry 分支就是先识别可能入口，再用 grep 定位真实定义，从而把语义召回和确定性搜索结合起来。'
]),
('28. 为什么有些 RAG 系统会引入父子索引？RepoMind 是否使用了？',[
'我的回答是：父子索引主要解决“检索粒度”和“阅读粒度”之间的冲突。小 Chunk 的 embedding 更聚焦，容易准确召回；但模型真正阅读时往往需要更大的父段落、完整函数或整个章节。父子索引会对子 Chunk 建索引，命中后返回对应父 Chunk，从而兼顾召回精度和上下文完整性。',
'在代码场景中，可以把一个类或文件作为父块，把方法或局部逻辑作为子块。检索命中某个方法后，返回完整方法加所属类摘要；文档场景则可以用段落作为子块、章节作为父块。父子关系必须通过稳定 ID 和版本管理维护，否则更新后容易返回过期父内容。',
'RepoMind 当前主要采用语义代码切片、元数据和后续工具扩展，没有实现标准的父子索引。它解决上下文完整性的方式是：召回 Chunk 后保留路径和行号，Agent 可以调用 read_file 阅读更大范围，也可以通过 find_definition 和 find_callers 扩展上下游。',
'这两种方式各有取舍。父子索引一次检索就能提供更完整上下文，延迟稳定；工具扩展更动态，模型只在需要时读取，但会增加 Agent 轮次。面试中我会明确父子索引是可选改进，不会把它包装成 RepoMind 已有功能。'
]),
('29. RepoMind 的上下文工程是怎么设计的？',[
'我的回答是：上下文工程不是简单地把所有历史、检索结果和工具输出拼到 Prompt 中，而是根据当前任务阶段选择必要信息、明确来源和结构，并控制 Token 预算。RepoMind 的上下文由任务信息、项目结构、历史案例、代码证据和工具轨迹几部分组成。',
'任务信息包括 Issue 标题、正文、labels、评论、关联 PR、图片描述和堆栈；project_map 提供仓库级目录结构；memory_context 提供历史相似案例；code_context 则由 entry、RAG、Git 三路证据合并。进入不同分析节点时还会注入 issue_type、is_regression、fix_type 或 feature_exists，使 Prompt 关注不同目标。',
'项目没有把整个仓库和所有 Memory 全部放入上下文。RAG 精排后只保留少量高价值 Chunk；工具结果按轮次加入 ReAct messages；工具集合也按任务裁剪，例如 style Review 不需要 Git 回归工具。结构上尽量把外部代码标记为证据，把系统规则、用户问题和工具结果分开，减少模型把仓库文本误当指令。',
'当前改进空间包括更严格的 Token 预算、长工具结果摘要、证据去重和上下文优先级。当上下文过长时，我会优先保留用户目标、关键约束、已确认结论和证据引用，压缩中间推理与重复代码，而不是简单从最早消息开始截断。'
]),
('30. RAG 召回率低或答案相关性低时，你会怎么排查和评测？',[
'我的回答是：我会先把问题拆成“解析与切片、查询构造、初始召回、融合精排、上下文组装、最终生成”六个阶段，不能看到答案错了就直接调 Prompt。每个阶段需要有独立可观测数据。',
'如果关键文件根本没进入初始 top-k，先检查仓库是否完整索引、文件类型是否被过滤、Chunk 是否把关键逻辑切碎、embedding 是否适合代码，以及 Query 是否包含有效实现线索。可以比较 Issue 原文、标题、HyDE、堆栈各路的 Recall@K，判断是哪一路失效。',
'如果关键文件已召回但精排后消失，就检查 RRF 候选数、CrossEncoder 模型是否适合代码、Query 与候选是否被截断以及 top-k 是否过小。如果正确代码已经进入最终上下文但答案仍然错，则问题更可能在 Prompt、证据组织、ReAct 路径或模型推理，而不是继续扩大召回。',
'评测上，检索层看关键文件/函数 Recall@K、MRR、NDCG；生成层看根因正确性、影响文件、严重度、修复可执行性和引用忠实度；系统层看延迟、Token 和工具次数。RepoMind 的 eval 目录已有 Dataset、Runner、LLM Judge 和 Report，改动前后要在固定回归集和独立留出集上比较，避免只修复当前 Badcase。'
]),
]

def sf(r,size=11,bold=None,color=None):
 n='PingFang SC';r.font.name=n;r.font.size=Pt(size);rf=r._element.get_or_add_rPr().get_or_add_rFonts()
 for a in('ascii','hAnsi','eastAsia','cs'):rf.set(qn('w:'+a),n)
 if bold is not None:r.bold=bold
 if color is not None:r.font.color.rgb=color
doc=Document();s=doc.sections[0];s.page_width,s.page_height=Inches(8.5),Inches(11);s.top_margin=s.bottom_margin=Inches(.82);s.left_margin=s.right_margin=Inches(.88)
st=doc.styles['Normal'];st.font.name='PingFang SC';st.font.size=Pt(11);st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_after=Pt(6);st.paragraph_format.line_spacing=1.25
for nm,sz,bf,af,c in [('Heading 1',16,16,8,BLUE),('Heading 2',13,12,6,DARK)]:
 st=doc.styles[nm];st.font.name='PingFang SC';st.font.size=Pt(sz);st.font.bold=True;st.font.color.rgb=c;st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_before=Pt(bf);st.paragraph_format.space_after=Pt(af);st.paragraph_format.keep_with_next=True
h=s.header.paragraphs[0];sf(h.add_run('RepoMind · Agent 开发面试问答'),9,color=GRAY);f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;sf(f.add_run('第三批 · 第 21—30 题'),9,color=GRAY)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(70);p.paragraph_format.space_after=Pt(10);sf(p.add_run('RepoMind Agent 开发面试问答'),24,True,DARK)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(35);sf(p.add_run('第三批：RAG 与上下文工程（10 题）'),14,color=BLUE)
doc.add_paragraph('本批为第 21—30 题，逐题结合 RepoMind 的索引、召回、精排和上下文构建代码整理。');doc.add_page_break()
for q,ps in QA:
 doc.add_heading(q,level=1)
 for t in ps:p=doc.add_paragraph(t);p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第三批10题';doc.core_properties.author='RepoMind';doc.save(OUT);print(OUT)
