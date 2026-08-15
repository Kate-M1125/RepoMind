from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Users/msy/Desktop/github/RepoMind/RepoMind_字节Agent开发二面_第一篇面经参考回答.docx")
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(95, 103, 112)


def font(run, name="PingFang SC", size=11, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    font(run, size=9, color=MUTED)


def add_answer(doc, question, paragraphs, bullets=None, code=None, note=None):
    doc.add_heading(question, level=1)
    label = doc.add_paragraph()
    label.paragraph_format.space_after = Pt(4)
    font(label.add_run("参考回答"), size=10, bold=True, color=DARK)
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.keep_together = True
    if bullets:
        for text in bullets:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(text)
    if code:
        p = doc.add_paragraph(style="Code Block")
        p.add_run(code)
    if note:
        p = doc.add_paragraph(style="Interview Note")
        font(p.add_run("补充说明："), size=10, bold=True, color=DARK)
        font(p.add_run(note), size=10, color=MUTED)


doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.85)
sec.left_margin = sec.right_margin = Inches(0.9)
sec.header_distance = sec.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "PingFang SC"
for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
    normal._element.rPr.rFonts.set(qn(f"w:{attr}"), "PingFang SC")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, before, after, color in (
    ("Heading 1", 16, 18, 8, BLUE),
    ("Heading 2", 13, 14, 7, BLUE),
    ("Heading 3", 12, 10, 5, DARK),
):
    st = styles[name]
    st.font.name = "PingFang SC"
    for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
        st._element.rPr.rFonts.set(qn(f"w:{attr}"), "PingFang SC")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    st = styles[list_name]
    st.font.name = "PingFang SC"
    for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
        st._element.rPr.rFonts.set(qn(f"w:{attr}"), "PingFang SC")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.375)
    st.paragraph_format.first_line_indent = Inches(-0.188)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.25

code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = "Menlo"
code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
code_style.font.size = Pt(9)
code_style.paragraph_format.left_indent = Inches(0.22)
code_style.paragraph_format.right_indent = Inches(0.22)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(8)
code_style.paragraph_format.line_spacing = 1.1
shd = OxmlElement("w:shd")
shd.set(qn("w:fill"), "F2F4F7")
code_style._element.get_or_add_pPr().append(shd)

note_style = styles.add_style("Interview Note", WD_STYLE_TYPE.PARAGRAPH)
note_style.font.name = "PingFang SC"
for attr in ("eastAsia", "ascii", "hAnsi", "cs"):
    note_style._element.rPr.rFonts.set(qn(f"w:{attr}"), "PingFang SC")
note_style.font.size = Pt(10)
note_style.paragraph_format.left_indent = Inches(0.18)
note_style.paragraph_format.space_before = Pt(4)
note_style.paragraph_format.space_after = Pt(8)
note_style.paragraph_format.line_spacing = 1.2
note_shd = OxmlElement("w:shd")
note_shd.set(qn("w:fill"), "F4F6F9")
note_style._element.get_or_add_pPr().append(note_shd)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run("RepoMind · Agent 开发面试问答"), size=9, color=MUTED)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("第 "), size=9, color=MUTED)
add_field(footer, "PAGE")
font(footer.add_run(" 页"), size=9, color=MUTED)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(70)
p.paragraph_format.space_after = Pt(10)
font(p.add_run("RepoMind Agent 开发面试问答"), size=25, bold=True, color=DARK)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(22)
font(p.add_run("字节二面 · 第一篇面经参考回答"), size=14, color=BLUE)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(55)
font(p.add_run("基于当前 RepoMind 仓库实现整理 · 2026 年 8 月"), size=10, color=MUTED)

doc.add_heading("使用说明", level=1)
doc.add_paragraph("本手册按照“面试问题—参考回答—必要补充”的方式整理。回答以 RepoMind 当前代码与项目文档为依据，重点说明设计动机、实现方式和技术取舍。涉及个人教育、实验室或工作经历的内容需由使用者按真实情况补充。")
doc.add_page_break()

add_answer(doc, "1. 你这两个项目都是学校里做的吗？硕士期间实验室做的什么工作？", [
    "RepoMind 是我围绕 Agent 工程化和代码分析场景完成的项目，主要目标是验证大模型能不能结合仓库检索、工具调用和工作流编排，完成 GitHub Issue 根因分析和 PR Code Review。",
    "介绍个人经历时，我会从项目来源、本人负责的模块和最终结果三个方面展开，避免只讲项目背景，没有突出个人贡献。",
], note="这一题涉及另一个项目、学校和实验室经历，必须按个人真实情况补充，不能根据 RepoMind 项目虚构。")

add_answer(doc, "2. 你做项目的出发点是什么？", [
    "这个项目的出发点是，分析 GitHub Issue 时，开发者通常要在问题描述、源码、调用链和 Git 历史之间反复切换；做 PR Review 时，也要同时关注安全、逻辑和代码风格，人工处理比较耗时。",
    "如果直接把 Issue 和一部分代码交给大模型，效果并不稳定。模型不知道应该看哪些文件，也缺乏继续搜索源码、追踪调用关系和验证结论的能力。",
    "所以我做了 RepoMind，希望把 RAG、代码工具和 Agent 工作流结合起来。RAG 负责缩小检索范围，ReAct Agent 负责继续探索代码，LangGraph 负责分类、路由、并行分析、反思重试和人工审批。项目重点不是简单调用一次 LLM，而是让分析过程可控、可恢复，并让最终结论尽量有代码证据支撑。",
])

add_answer(doc, "3. 那介绍一下你的项目吧？", [
    "RepoMind 是一个面向 GitHub 仓库的智能代码分析系统。用户输入 Issue 或 Pull Request URL，系统会自动识别任务类型，并进入不同的分析工作流。",
    "Issue 链路会获取 Issue、评论和关联信息，clone 仓库并建立代码索引，再通过 RAG 和代码工具定位相关源码、分析根因，最终输出问题类型、根因、影响文件、严重程度和修复建议。",
    "PR 链路会读取 PR diff，然后从 security、logic 和 style 三个角度并行审查，最终生成 APPROVE 或 REQUEST CHANGES。",
    "技术上，我使用 LangGraph 编排状态机，用 ChromaDB 和 Sentence Transformer 做向量检索，用 CrossEncoder 精排，用 ReAct 实现多轮工具调用，并使用 SQLite Checkpoint 保存工作流状态。分析结束后，独立 Reflection 节点会检查证据是否充分，置信度低于 0.7 时触发重新分析。",
])

add_answer(doc, "4. 你选 Agent 项目框架时有什么考虑？除了 LangChain，还了解或调研过别的框架吗？", [
    "我选择框架时主要考虑四点：复杂状态能不能显式管理、是否支持条件路由和循环、是否支持并行任务，以及长任务能不能中断恢复。",
    "RepoMind 不是简单的问答链。Issue 要先分类，再根据 bug、feature、question 或 security 进入不同分支；上下文收集需要并行执行；分析阶段存在 ReAct 循环；低置信度时还需要回退重试。因此相比普通 LangChain Chain，LangGraph 的图结构和状态机更适合这个项目。",
    "我也了解过 AutoGen、CrewAI，以及直接自研状态机。AutoGen 和 CrewAI 更强调多 Agent 对话和角色协作；但 RepoMind 的核心不是让多个角色自由聊天，而是让代码分析按照可控流程执行。完全自研虽然自由度更高，但要自己处理状态合并、条件路由、中断恢复和 Checkpoint，工程成本较高。",
    "因此我采用 LangGraph：外层用确定性的 Workflow 保证流程可控，内层在分析节点中使用 ReAct，让模型自主决定调用什么工具。",
])

add_answer(doc, "5. 谈谈你对 Agent 的理解。Agent 包括哪些部分？运行机制怎样？", [
    "我认为 Agent 不是接了大模型的聊天机器人，而是一个能够围绕目标持续感知、决策、执行和反馈的系统。",
], bullets=[
    "模型：负责理解任务和作出决策。",
    "状态与 Memory：保存当前目标、上下文、执行进度和可复用经验。",
    "工具：负责与代码仓库、GitHub、Git 历史等外部环境交互。",
    "Planning：决定任务如何拆解以及下一步做什么。",
    "Workflow / Runtime：控制路由、循环、重试、权限和停止条件。",
    "Eval / Reflection：判断结果是否达到要求。",
], code="获取任务 → 构建上下文 → 模型判断下一步动作\n→ 调用代码或 Git 工具 → 读取结果 → 更新状态\n→ 判断继续探索还是输出结论 → Reflection 检查结论")

add_answer(doc, "6. Planning 是 Agent 框架做的，还是大模型做的？", [
    "两者都可以参与，但负责的层次不同。框架更适合确定性的宏观规划，例如先获取 Issue，再建立索引、分析并生成报告；这些步骤有明确依赖，需要框架保证顺序、状态和失败恢复。",
    "大模型更适合动态的微观规划。例如拿到相关代码后，下一步应该读哪个文件、查哪个函数、是否继续追踪调用者，这种路径难以提前写死。",
    "RepoMind 使用混合方式：外层由 LangGraph 控制固定流程；进入 analyze_bug 等节点后，由 ReAct Agent 自主选择 read_file、grep_code、find_definition、find_callers 和 Git 历史工具。我的总结是：框架负责流程边界和可靠性，大模型负责边界内的不确定决策。",
])

add_answer(doc, "7. Planning 有哪些方法可以实现？", [
    "常见的 Planning 方法可以分为四类。",
], bullets=[
    "预定义 Workflow：提前写好 DAG、状态机和条件分支，稳定、易观测，但未知路径下不够灵活。",
    "ReAct：按照思考—行动—观察的方式边执行边规划，适合下一步依赖工具结果的代码探索任务。",
    "Plan-and-Execute：先生成完整或阶段性计划，再逐步执行；全局性较好，但需要支持计划失效后的重新规划。",
    "Reflection / Critic-Actor：Actor 生成方案，Critic 检查证据和结果，质量不足时携带反馈重新执行。",
], note="RepoMind 组合使用了预定义 Workflow、阶段性分类规划、ReAct 工具循环和 Reflection 重试。")

add_answer(doc, "8. 写 Agent 时有用过长期记忆和短期记忆吗？", [
    "两种都使用了，但是用途不同。短期记忆主要存在当前 LangGraph State 和 ReAct 消息上下文中，包括 Issue 内容、检索代码、工具调用、工具结果、当前结论和重试次数，服务于当前任务。",
    "长期记忆使用 ChromaDB 保存已经完成的 Issue 分析和 PR Review 结果。新任务开始时，系统根据仓库、问题类型和语义相似度召回历史案例；任务结束后，再把结构化结果写回 Memory。",
    "我不会把所有历史对话直接写入长期记忆。长期记忆应该保存经过整理、能够复用的结论，否则容易引入噪声和错误记忆污染。",
])

add_answer(doc, "9. 短期记忆和长期记忆有什么区别？", [
    "最核心的区别是生命周期和使用目的。短期记忆服务于当前任务，保存最近消息、工具调用、临时上下文和执行状态；它更新频繁，任务结束后可以清理。",
    "长期记忆可以跨任务、跨会话保留，保存历史案例、用户偏好、项目规则和经过验证的经验。它不会每次全部加载，而是根据任务按需召回。",
    "在 RepoMind 中，LangGraph State、Checkpoint 和 ReAct messages 属于短期或任务级记忆；ChromaDB 中的历史分析结果属于长期记忆。Checkpoint 解决当前任务中断后从哪里继续，长期记忆解决未来遇到类似任务时能否复用以前经验。",
])

add_answer(doc, "10. 短期记忆里会存储 Function Call 吗？不放会有什么问题？", [
    "需要保存，而且不仅要保存模型发出的 Function Call，还要保存对应的工具结果。ReAct 的下一步判断依赖完整消息链。",
    "如果不保存，模型可能重复调用同一个工具，无法沿着上一次结果继续分析，最终结论也缺少证据链。在模型 API 中，assistant tool call 和 tool result 通常还需要保持配对，否则可能产生消息协议错误。",
    "RepoMind 的 LLM Client 会把 assistant 工具调用和执行结果都追加到 messages 中，然后发起下一轮推理；框架同时限制最大轮数，避免模型不断重复调用。",
])

add_answer(doc, "11. 工具会使用 MCP 吗，还是自己写？大模型工具机制有什么好处？", [
    "RepoMind 的工具执行层是自己实现的，但接口设计遵循 MCP 风格的工具规范。每个工具都有统一名称、描述、JSON Schema 参数和执行器，并通过 ToolRegistry 注册。",
    "项目提供 read_file、grep_code、find_definition、find_callers、git_blame 和 search_commits 等工具。底层通过 GitHub API、本地仓库和 Git 命令执行，上层模型只需要按照 Schema 生成调用参数。",
    "这种设计把模型和底层实现解耦，也让能力边界显式化：模型只能使用系统暴露的工具，参数可以经过 Schema 校验，高风险操作还能接入权限检查和人工审批。相比直接执行模型生成的任意代码，它更可控，也更容易记录和审计。",
])

add_answer(doc, "12. 关于 RAG 有什么看法？召回时调过哪些参数？", [
    "我认为 RAG 的作用不是把更多代码塞给模型，而是从大型仓库中找到高概率相关的证据，降低模型的搜索空间。",
    "RepoMind 使用多阶段检索：对代码做语义切片；查询阶段同时使用 HyDE 假设文档、Issue 标题、问题类型加正文和堆栈信息；每一路召回 top-5，经 RRF 融合得到 top-10，再用 CrossEncoder 精排，最终保留 top-3。",
    "我重点关注和调节 Chunk 大小、Overlap、每路 top-k、融合候选数、精排保留数以及不同 Issue 类型的检索权重。评价时观察相关文件能否进入 top-k，同时关注最终根因准确率、上下文 Token 数和检索延迟。",
    "RAG 负责发现候选代码，不能代替完整分析。因此检索之后仍然需要 ReAct 工具循环，让模型继续查定义、调用者和 Git 历史。",
])

add_answer(doc, "13. Python 的一些内容？", [
    "这道题范围较宽，我会根据面试官的具体追问回答。结合 RepoMind，比较容易延伸到异步编程、类型系统、并发模型和资源管理。",
    "RepoMind 使用 TypedDict 定义 LangGraph State，因为它既保留字典的动态更新方式，也明确每个状态字段的类型；输入输出再使用 Pydantic 校验，防止 LLM 返回结构合法但业务字段不符合要求。",
    "如果问并发，我会区分线程、进程和异步 IO。GIL 会限制单进程内 CPU 密集型 Python 线程并行，但模型调用和 GitHub API 请求主要是 IO 密集型，更适合异步 IO 或框架并行调度。Embedding 和 CrossEncoder 等 CPU/GPU 密集计算，可以交给底层计算库、独立进程或专门推理服务。",
])

add_answer(doc, "14. 手撕：合并两个有序数组", [
    "如果允许新建数组，可以使用双指针，从两个数组头部开始比较，每次取较小值放入结果数组，时间复杂度为 O(m+n)，空间复杂度为 O(m+n)。",
    "如果要求把第二个数组合并进第一个数组，并且第一个数组尾部已经预留空间，就从后向前写，避免覆盖尚未处理的元素。",
], code="def merge(nums1, m, nums2, n):\n    i, j, k = m - 1, n - 1, m + n - 1\n\n    while i >= 0 and j >= 0:\n        if nums1[i] > nums2[j]:\n            nums1[k] = nums1[i]\n            i -= 1\n        else:\n            nums1[k] = nums2[j]\n            j -= 1\n        k -= 1\n\n    while j >= 0:\n        nums1[k] = nums2[j]\n        j -= 1\n        k -= 1", note="不需要单独处理 nums1 剩余元素，因为它们本来就在正确位置。该实现时间复杂度 O(m+n)，额外空间 O(1)。")

doc.core_properties.title = "RepoMind Agent 开发面试问答"
doc.core_properties.subject = "字节 Agent 开发二面第一篇面经参考回答"
doc.core_properties.author = "RepoMind"
doc.save(OUT)
print(OUT)
