from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches,Pt,RGBColor
OUT=Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第六批20题.docx')
QA=[
('71. Issue 意图识别为什么不能只做一个普通的单标签分类？',[
'我的回答是：单标签分类只能回答“它属于哪一类”，却不能完整表达后续分析需要的控制信息。例如同一个 bug 还可能是版本回归，修复方式可能是代码修改、依赖升级或配置调整；security 还涉及是否需要人工批准。',
'RepoMind 先用 route_issue 输出 bug、feature、question、security，随后在特定分支继续规划：bug 经过 detect_regression 和 classify_fix，feature 经过 detect_existing，security 进入 human_gate。也就是说，第一层分类用于选择大路径，第二层判断用于决定证据和 Prompt。',
'如果把所有组合做成一个巨大枚举，类别会快速膨胀，训练和维护都困难。更合理的是分层识别：先稳定地判断主意图，再由专门节点提取该意图需要的属性。'
]),
('72. 澄清节点应该由模型决定，还是由规则决定？',[
'我的回答是：是否需要澄清可以由模型发现信息缺口，但是否阻塞流程最好由规则和风险策略决定。模型擅长判断语义是否模糊，代码擅长判断必填字段、权限和高风险条件。',
'RepoMind 当前输入主要是标准 GitHub URL，可以通过正则提取 owner、repo 和编号，通常不需要用户澄清；Issue 类型由拉取完整正文后判断。security 分支是否继续则不是模型随意决定，而是固定进入 human_gate。',
'如果未来支持自然语言任务，我会让模型输出 missing_fields 和 clarification_question，再由规则判断这些字段是否影响执行。可安全使用默认值的信息继续运行，涉及仓库、目标分支和写操作范围时必须询问用户。'
]),
('73. RepoMind 的 Human-in-the-Loop 为什么只放在 Security 分支？',[
'我的回答是：HITL 应放在风险与不确定性最高、错误后果不可接受的位置，而不是每个节点都让人确认。过多审批会把 Agent 退化成人工操作表单，失去自动化价值。',
'RepoMind 的 route_issue 识别为 security 后固定进入 human_gate，批准后才继续 Memory、上下文和分析。这体现的是风险分级：普通 bug、feature 和 question 主要执行读取与分析；安全问题可能包含敏感漏洞信息，需要控制处理和传播。',
'当前实现仍较基础。若未来加入自动改代码、创建 PR 或外部写工具，HITL 应按操作风险而不是只按 Issue 类型触发，并展示具体目标、参数、影响范围和可回滚性。'
]),
('74. RepoMind 的 PR Code Review 工作流是怎么设计的？',[
'我的回答是：PRGraph 先执行 fetch_pr、build_project_map、describe_pr_images 和 memory_retrieve_pr，获取 diff、仓库结构、图片上下文和历史 Review 案例。随后 prepare_reviews 作为并行派发点。',
'dispatch_reviews 使用 LangGraph Send 同时创建 security、style、logic 三个 run_review。三路共享 PR 基础数据，但使用不同 Prompt、工具集合和输出字段；LangGraph 在三路全部完成后才进入 reflect_pr。',
'reflect_pr 检查三类问题是否有证据并输出 confidence。低于 0.7 且 iteration<2 时回 prepare_reviews 重审，不重跑抓取和 Memory；通过后 generate_pr_report 生成 Verdict，最后 memory_save_pr 保存案例。'
]),
('75. PR Review 为什么拆成 security、style、logic 三路？怎么避免重复和冲突？',[
'我的回答是：三类 Review 的关注目标和证据不同。security 关注权限、注入、敏感信息和危险调用；logic 关注边界条件、异常路径和行为回归；style 关注命名、可读性和维护成本。拆分后 Prompt 更聚焦，也可以并行降低总延迟。',
'RepoMind 让三路分别写 security_issues、style_issues、logic_issues，避免并行覆盖。每条问题使用结构化字段保存文件、行号、严重度、类别、描述和建议，最后统一合并。',
'重复问题可以按 file+line+normalized description 做去重；分类冲突时优先保留风险更高、证据更具体的记录，并由 reflect_pr 检查。三路不通过自然语言互相协商，减少信息变形。'
]),
('76. PR 的 APPROVE / REQUEST CHANGES 逻辑为什么这样设计？',[
'我的回答是：RepoMind 当前规则是存在 critical 或 high 问题时 REQUEST CHANGES；只有 low 和 medium 时仍然 APPROVE，但在报告中给出改进建议。这样把阻塞合并与非阻塞建议区分开，避免风格或一般维护性问题阻止所有 PR。',
'Verdict 由确定性代码根据结构化 severity 计算，而不是让模型自由写一句 APPROVE。这样规则可测试、可解释，也方便团队根据治理策略调整。',
'需要说明，medium 是否阻塞没有通用答案。真实团队可能要求某些 logic medium 必须修复，因此更成熟的实现应按 category、仓库规则和分支保护配置阈值，而不是全球固定。'
]),
('77. RepoMind 当前有哪些明确不足？面试时应该怎么讲？',[
'我的回答是：首先，后端 jobs 是进程内字典，虽然 Graph 有 SQLite Checkpoint，但 Web 任务发现和多实例调度不够生产化。其次，非 Python 代码的语义切片和调用图能力有限，主要依赖通用切分与搜索。',
'第三，Memory 缺少版本、过期、冲突和人工纠错治理；安全方面只有 security human_gate，还没有完整 RBAC、租户隔离和写工具两阶段提交。第四，Reflection 和 Judge 依赖 LLM，评测集规模与线上指标仍需要扩充。',
'面试时我会把不足与改进路径对应起来：外置 Job+MQ、语言服务/Tree-sitter、多版本 Memory、风险分级工具网关、规则+人工混合评测。承认边界比虚构“全部生产化”更可信。'
]),
('78. 跨仓库 Orchestrator 是怎么设计的？什么时候需要多仓库分析？',[
'我的回答是：当一个 Issue 可能涉及主仓库、SDK、依赖库或多个服务时，只分析单仓库会漏掉真实根因。RepoMind 在 agents/orchestrator 中提供独立 Graph，先识别相关仓库，再通过 Send 并行执行 repo_analysis，最后 merge_reports。',
'每个 Repo Agent 负责自己的仓库上下文和分析结果，Orchestrator 不让它们自由聊天，而是收集结构化 report、confidence 和 repo_name。合并时优先考虑成功结果和高置信度证据。',
'跨仓库分析的难点是版本对齐、权限和成本。必须明确依赖 commit/tag，限制最大仓库数量，并避免同一依赖被多个 Agent 重复索引。只有证据确实跨边界时才启用。'
]),
('79. fetch_dependency 为什么需要单独工具？依赖源码怎么参与根因分析？',[
'我的回答是：很多问题表面发生在当前仓库，根因实际来自第三方包的行为或版本不兼容。如果只看业务仓库，模型可能错误建议修改调用方，而正确修复是升级依赖或适配 API。',
'RepoMind 的 dep_indexer 从 RAG Chunk metadata 的 imports 中识别候选依赖，fetch_dependency 允许 ReAct 按需检索依赖源码。bug 在 classify_fix 阶段还会区分 dependency_upgrade，使 Prompt 关注 pyproject.toml 等依赖配置。',
'依赖证据必须绑定版本。不能用最新主分支源码解释旧版本行为；生产化应从 lockfile、包元数据或 Git tag 确认实际版本，并控制递归依赖深度。'
]),
('80. 项目地图和调用图分别解决什么问题？',[
'我的回答是：项目地图提供宏观结构，让模型知道代码、测试、配置和文档大致位于哪里；调用图提供函数之间的关系，帮助从入口追踪到实现或从被调函数反向寻找调用者。',
'RepoMind 在 build_project_map 阶段先生成仓库目录视图，避免模型完全盲搜。RAG 分支后续补充轻量调用图，ReAct 还可以使用 find_definition 和 find_callers 做动态导航。',
'项目地图不能替代检索，调用图也不一定完整。Python 动态调用、装饰器、反射和跨语言边界会降低静态分析准确率，所以调用关系应作为候选线索，再通过 read_file 验证。'
]),
('81. Issue 中的图片和堆栈信息怎么处理？为什么要单独节点？',[
'我的回答是：图片和堆栈是两种结构完全不同的证据。截图可能包含 UI、错误弹窗和网络信息，需要视觉模型描述；堆栈则可以确定性提取文件、行号、异常类型和调用顺序。',
'RepoMind 在 build_project_map 后分别执行 describe_images 和 parse_stack_trace，把结果写入 image_descriptions 与 stack_trace_context。之后 route、检索和分析节点使用这些字段，而不是把原始二进制或杂乱正文重复解析。',
'单独节点便于失败降级：没有图片就返回空描述，堆栈解析失败也不阻断普通检索。堆栈中的文件和函数可以成为高精度 Query，但仍要防止用户粘贴的路径与当前仓库版本不一致。'
]),
('82. Self-Attention 的原理是什么？为什么分成 Q、K、V？',[
'我的回答是：输入向量分别乘三个可学习矩阵得到 Q、K、V。Q 表示当前 token 想查询什么，K 表示每个 token 可以用什么特征被匹配，QK 转置点积经过缩放和 softmax 得到注意力权重，再用权重对 V 做加权求和。',
'分成三种投影使“如何匹配”和“传递什么内容”可以在不同子空间学习。如果直接对原向量做相似度并聚合，表达能力受限。除以根号 d_k 是为了防止维度增大后点积过大导致 softmax 饱和。',
'多头注意力把维度拆成多个头，每个头可学习不同关系，最后拼接投影。它的标准序列计算复杂度约为 O(n²d)，长上下文成本主要来自 n² 注意力和 KV Cache。'
]),
('83. 同一个 Token 在不同位置的向量一样吗？',[
'我的回答是：词嵌入查表阶段，同一个 token id 通常得到相同基础 embedding；加入位置编码后，不同位置的初始表示已经不同。经过 Self-Attention 后，它还会根据左右上下文形成不同的 contextual representation。',
'例如“bank”在金融和河岸语境中 token 相同，但注意到的邻近词不同，后续层向量会明显不同。即使词和上下文相同，绝对或旋转位置编码也会使不同位置表示不同。',
'因此不能把 token embedding 与模型某一层的最终 hidden state 混为一谈。前者偏词表表示，后者已经融合位置和上下文。'
]),
('84. KV Cache 是什么？大小怎么估算？',[
'我的回答是：自回归生成时，历史 token 在每一层的 K、V 不会变化。KV Cache 保存这些张量，新 token 只计算自己的 Q/K/V 并与缓存交互，避免每一步重新计算全部历史。',
'粗略大小为 batch×layers×sequence_length×2(K和V)×kv_heads×head_dim×每元素字节数。使用 MHA 时 kv_heads 通常等于 attention heads；GQA/MQA 会减少 KV 头数，从而显著降低缓存。',
'例如估算时必须说明 batch、层数、上下文长度、精度和头配置。KV Cache 随序列长度线性增长，但每步注意力仍需读取历史缓存，长上下文会增加显存带宽和延迟。'
]),
('85. Continuous Batching 和 vLLM 解决什么问题？RepoMind 使用了吗？',[
'我的回答是：静态批处理必须等待一批请求全部结束，长短请求混合会造成空闲。Continuous Batching 在 token 生成步动态加入新请求、移除完成请求，提高 GPU 利用率。vLLM 还通过 PagedAttention 以分页方式管理 KV Cache，减少碎片和预留浪费。',
'它们主要优化自建模型推理吞吐，不直接提升模型答案质量。使用时要平衡首 token 延迟、吞吐、最大上下文和调度公平性。',
'RepoMind 当前调用 OpenAI-compatible 外部 API，没有部署 vLLM，也没有线上吞吐数据。面试时可以说明原理和适用场景，但不能说项目已完成 Continuous Batching 优化。'
]),
('86. Python 有真正的多线程吗？为什么有 GIL？',[
'我的回答是：Python 语言本身支持线程，但在常见 CPython 中，GIL 使同一进程同一时刻通常只有一个线程执行 Python 字节码。它主要简化引用计数内存管理和 C 扩展兼容，并不等于线程完全无用。',
'IO 操作会释放 GIL，所以网络请求、文件 IO 和等待模型 API 的任务仍能从线程或 asyncio 获益；CPU 密集的纯 Python 计算更适合多进程、原生库或 GPU。NumPy/PyTorch 等底层计算也可能释放 GIL。',
'RepoMind 的 GitHub、LLM 和依赖请求偏 IO，适合并发；Embedding、CrossEncoder 和 AST 大规模处理偏计算，需要限制并发或独立服务。LangGraph Send 表达的是任务并行，不应直接等同于 CPython 线程并行。'
]),
('87. Lock 和 RLock 有什么区别？Agent 服务里怎么使用锁？',[
'我的回答是：Lock 是普通互斥锁，同一线程重复 acquire 会阻塞自己；RLock 记录持有线程和递归次数，同一线程可以重复获取，但必须对应次数释放。RLock 适合嵌套调用可能重复进入同一临界区的场景。',
'进程内锁不能解决多实例并发。Agent 服务若要串行同一 conversation，应使用数据库乐观锁、Redis 租约锁或按 key 分区队列，并设置超时、续租和 owner token，释放时验证所有者。',
'RepoMind 当前使用 thread_id 隔离 Checkpoint，但 jobs 字典不是分布式并发控制。不能因为 Python 字典单次操作受 GIL 保护，就认为业务状态更新已经原子。'
]),
('88. TypedDict 和 Pydantic 在 RepoMind 中分别起什么作用？',[
'我的回答是：TypedDict 主要用于静态描述 LangGraph State 的字段，让节点仍以普通 dict 方式读写，同时获得类型提示；Pydantic 用于运行时校验外部或模型输出，能检查必填字段、类型、枚举和自定义约束。',
'IssueState 使用 TypedDict 定义 issue_type、root_cause、confidence、iteration 等状态。模型返回 JSON 时不能因为能 parse 就直接写入 State，需要经过结构校验；confidence 还要限制到 0 到 1。',
'二者不是替代关系。TypedDict 本身不会在运行时阻止错误值，Pydantic 则有校验成本。内部图状态用 TypedDict 保持轻量，边界数据用 Pydantic 防止污染，是更合适的职责划分。'
]),
('89. 代码单测的分支覆盖率怎么统计？插桩原理是什么？',[
'我的回答是：语句覆盖只看某行是否执行，分支覆盖还要记录 if/else、条件跳转和异常路径的每个出口是否走过。coverage.py 会通过解释器 tracing/字节码信息记录已执行行和跳转边，再与静态可达分支比较。',
'RepoMind 的 tests 包含节点、代码导航、PR Review、鲁棒性和 regression 测试。仅看覆盖率不够，LLM 节点还应 Mock 模型响应，验证不同 issue_type 路由、低置信度回边、工具失败和最大 iteration。',
'覆盖率高不等于断言有效。关键是针对状态转换和边界条件设计测试，并把历史 Badcase 固化为回归 Case。'
]),
('90. 单元测试中的 Mock 怎么做？AST 和 LSP 都无法解析的代码怎么办？',[
'我的回答是：Mock 应替换系统边界而不是内部实现细节，例如 GitHub API、LLM、向量库、文件系统和工具 executor。固定输入输出后验证节点更新了哪些 State 字段、走了哪条边，而不是只断言函数被调用。',
'对于 AST 解析失败的代码，RepoMind 使用通用文本切分和 grep/read_file 兜底。LSP 不可用时仍可做关键词搜索、正则定义定位、目录结构和 Git 历史；同时把解析失败记录为元数据，避免模型误以为调用图完整。',
'复杂生成代码、宏、反射和多语言项目不可能完全依赖单一静态分析。合理策略是多级降级，并在报告中表达证据限制；测试要覆盖这些降级路径。'
]),
]
B=RGBColor(46,116,181);D=RGBColor(31,77,120);G=RGBColor(95,103,112)
def sf(r,z=11,b=None,c=None):
 n='PingFang SC';r.font.name=n;r.font.size=Pt(z);rf=r._element.get_or_add_rPr().get_or_add_rFonts()
 for a in('ascii','hAnsi','eastAsia','cs'):rf.set(qn('w:'+a),n)
 if b is not None:r.bold=b
 if c is not None:r.font.color.rgb=c
doc=Document();s=doc.sections[0];s.page_width,s.page_height=Inches(8.5),Inches(11);s.top_margin=s.bottom_margin=Inches(.82);s.left_margin=s.right_margin=Inches(.88)
st=doc.styles['Normal'];st.font.name='PingFang SC';st.font.size=Pt(11);st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_after=Pt(6);st.paragraph_format.line_spacing=1.25
for nm,z,bf,af,c in [('Heading 1',16,16,8,B),('Heading 2',13,12,6,D)]:
 st=doc.styles[nm];st.font.name='PingFang SC';st.font.size=Pt(z);st.font.bold=True;st.font.color.rgb=c;st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_before=Pt(bf);st.paragraph_format.space_after=Pt(af);st.paragraph_format.keep_with_next=True
sf(s.header.paragraphs[0].add_run('RepoMind · Agent 开发面试问答'),9,c=G);f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;sf(f.add_run('第六批 · 第 71—90 题'),9,c=G)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(70);p.paragraph_format.space_after=Pt(10);sf(p.add_run('RepoMind Agent 开发面试问答'),24,True,D)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(35);sf(p.add_run('第六批：业务图、模型基础与测试工程（20 题）'),14,c=B)
doc.add_paragraph('本批为第 71—90 题，覆盖 Issue/PR 设计、跨仓库分析、模型推理基础、Python 并发和测试。');doc.add_page_break()
for q,ps in QA:
 doc.add_heading(q,level=1)
 for t in ps:p=doc.add_paragraph(t);p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第六批20题';doc.core_properties.author='RepoMind';doc.save(OUT);print(OUT)
