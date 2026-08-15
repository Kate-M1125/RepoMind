from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SRC = Path('/Users/msy/Documents/Codex/2026-08-11/biu/outputs/字节agent开发二面面经_重新排版.docx')
OUT = Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_统一详细回答版.docx')
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 99, 108)


def normalize_question(s):
    s = s.replace('\u00a0', ' ').strip()
    s = re.sub(r'^[•\-\s]+', '', s)
    s = re.sub(r'^\d+(?:\.\d+)?[\.、）)]\s*', '', s)
    s = re.sub(r'\s+', '', s)
    s = s.replace('?', '？').replace('；', '。').rstrip('。')
    return s


def looks_like_question(text, style, para_index):
    s = text.strip()
    if len(s) < 4 or len(s) > 260:
        return False
    if '？' in s or '?' in s:
        return True
    if style.startswith('Heading'):
        return not any(x in s for x in ('最后说两句', '反问环节', '隐藏信息'))
    if para_index == 87:
        return bool(re.match(r'^\d+[\.、]', s))
    # Include compact numbered/bulleted interview prompts, but reject answer-like prose.
    if re.match(r'^(?:[•]|\d+(?:\.\d+)?[\.、）)])', s) and len(s) <= 72:
        reject = ('我的回答', '答：', '第一层', '第二层', '第三层', '核心是', '可以分', '包括：', '复制代码')
        return not any(x in s for x in reject)
    return False


def extract_questions():
    src = Document(SRC)
    found = []
    seen = set()
    for pi, para in enumerate(src.paragraphs[87:], 87):
        for raw in para.text.splitlines():
            s = raw.strip()
            if not looks_like_question(s, para.style.name, pi):
                continue
            key = normalize_question(s)
            if not key or key in seen:
                continue
            seen.add(key)
            clean = re.sub(r'^[•\-\s]+', '', s)
            clean = re.sub(r'^\d+(?:\.\d+)?[\.、）)]\s*', '', clean)
            found.append(clean.strip())
    return found


def category(q):
    x = q.lower()
    if any(k in x for k in ('算法', '手撕', '二分', '最长', '股票', '链表', '二叉树', '岛屿', '三数之和', '编辑距离', '爬楼梯', '大数', '数组', '字符串')):
        return '算法与数据结构'
    if any(k in x for k in ('rag', '检索', '向量', 'embedding', 'chunk', 'rerank', 'bm25', '知识库', '召回', 'hyde', '索引')):
        return 'RAG 与上下文工程'
    if any(k in x for k in ('memory', '记忆', '上下文窗口', '上下文压缩', 'checkpoint', '状态', 'state')):
        return 'Memory、状态与长任务'
    if any(k in x for k in ('tool', 'mcp', 'function call', 'function call', 'skill', '权限', '沙箱', '注入', '安全', '调用协议', 'a2a')):
        return '工具、MCP、Skill 与安全'
    if any(k in x for k in ('评测', '评估', 'eval', 'badcase', 'bad case', 'judge', '指标', '准确率', '幻觉', 'trace', '测试集', '质量')):
        return '评测、Badcase 与可观测性'
    if any(k in x for k in ('agent', 'workflow', 'langgraph', 'langchain', 'planning', 'react', 'multi-agent', 'multi agent', '多agent', '单agent', '架构', '项目', '模型选型')):
        return 'Agent 项目架构与设计'
    if any(k in x for k in ('mysql', 'redis', 'kafka', '消息队列', 'tcp', 'udp', 'http', 'https', 'jvm', 'java', 'python', 'gil', '锁', '并发', '缓存', '数据库', 'linux', 'git', 'qps', '分布式', '微服务', 'pod', '服务')):
        return '工程基础与系统设计'
    if any(k in x for k in ('attention', 'transformer', 'qkv', 'kv cache', 'sft', 'dpo', 'grpo', 'ppo', 'lora', '微调', '强化学习', 'moe', '模型')):
        return '大模型基础与训练'
    if any(k in x for k in ('自我介绍', '实习', '团队', '离职', '到岗', '职级', '职业', '个人', '为什么投', '规划')):
        return '个人经历与行为问题'
    return '场景题与其他'


def answer(q):
    x = q.lower()
    unsupported = 'RepoMind 当前实现没有直接覆盖这一点，面试时我会明确区分“已经实现”和“可落地的改进方案”，不把设计设想说成线上事实。'
    if any(k in x for k in ('自我介绍', '实习状况', '团队规模', '离开当前公司', '到岗', '职级预期', '实验室', '论文')):
        return '这道题必须结合个人真实经历回答。建议用“背景—本人职责—关键难点—量化结果—与岗位匹配点”的结构，并把 RepoMind 作为 Agent 工程化案例介绍。' + unsupported
    if '介绍' in x and ('项目' in x or 'agent' in x):
        return ('RepoMind 是一个面向 GitHub 仓库的代码分析 Agent。输入 Issue 或 PR URL 后，系统拉取 GitHub 数据、clone 并索引仓库；Issue 链路完成分类、根因定位和修复建议，PR 链路从 security、logic、style 三路并行审查。外层使用 LangGraph 管理状态、路由、并行、循环和 checkpoint，内层使用 ReAct 配合代码导航与 Git 工具动态取证，最后用独立 Reflection 节点检查证据和置信度。')
    if '出发点' in x or ('为什么' in x and '项目' in x):
        return ('出发点是解决代码问题分析中“上下文分散、搜索耗时、单次 LLM 结论缺少证据”的问题。RepoMind 用 RAG 缩小代码范围，用工具继续追踪定义、调用链和 Git 历史，再用可恢复工作流约束模型行为；价值不在一次生成，而在可控、可追踪的分析链路。')
    if 'langgraph' in x and 'langchain' in x:
        return ('LangChain 更偏通用 LLM 组件和线性 Chain；LangGraph 用显式 StateGraph 表达条件分支、并行 Send、循环、中断与 checkpoint。RepoMind 同时需要 Issue 类型路由、三路上下文收集、PR 三路审查和低置信度重试，因此选择 LangGraph，而不是只用线性 Chain。')
    if any(k in x for k in ('多agent', 'multi-agent', 'multi agent')):
        return ('RepoMind 主体是 Workflow 编排的专职分析节点，而不是让多个 Agent 自由对话。Issue 与 PR 使用不同 Graph；PR 的 security、logic、style 三路并行使用相同执行节点加 review_type 路由，避免角色聊天带来的通信成本和状态漂移。只有任务具有清晰边界、可并行、输出协议稳定时才值得拆成多 Agent。')
    if 'workflow' in x and 'agent' in x:
        return ('Workflow 由代码决定路径，适合确定性步骤；Agent 由模型根据观察动态选择动作，适合未知搜索路径。RepoMind 采用“外层 Workflow、内层 Agent Loop”：LangGraph 控制抓取、索引、路由、重试和报告，analyze 节点内部由 ReAct 自主选择工具。')
    if any(k in x for k in ('planning', '规划', '任务拆解')):
        return ('RepoMind 使用分层 Planning。宏观流程由 LangGraph 固定；detect_regression、classify_fix 等节点先做阶段性判断；进入分析节点后由 ReAct 根据工具结果逐步规划。框架负责边界、预算和恢复，模型负责边界内的不确定决策。')
    if '死循环' in x or '最大步数' in x or 'loop' in x:
        return ('循环控制必须放在代码层。RepoMind 的 ReAct 设置最大轮数，Reflection 低于 0.7 才重试且最多 2 次，并用 iteration 字段阻止无限回边。生产化还应记录工具调用指纹、状态变化量、Token/时间预算，对重复调用或无状态增量提前终止。')
    if '反思' in x or 'reflection' in x:
        return ('RepoMind 把 Reflection 设计成独立 critic，不共享 analyze 节点的即时推理过程，只读取结构化结论和证据，检查证据是否充分、是否存在其他原因并输出置信度。低于 0.7 时携带反馈重试，最多两次，避免无效反思无限循环。')
    if any(k in x for k in ('memory', '记忆')):
        if 'rag' in x:
            return ('RAG 面向外部知识与仓库源码，解决“去哪里找证据”；Memory 面向 Agent 历史状态和经验，解决“过去确认过什么”。RepoMind 的代码索引属于 RAG，ChromaDB 中保存的历史 Issue/PR 分析结果属于长期 Memory，两者使用相似检索技术但写入来源、生命周期和治理规则不同。')
        return ('RepoMind 的短期/任务记忆包括 LangGraph State、ReAct messages 和 SQLite checkpoint，保存当前输入、工具调用、证据、进度与重试次数；长期记忆用 ChromaDB 保存完成后的结构化 Issue/PR 案例，并在新任务开始时按语义召回。长期写入应去重、带来源和置信度，不能把临时猜测直接固化。')
    if 'checkpoint' in x or '恢复' in x or '中断' in x or '长任务' in x:
        return ('RepoMind 使用 LangGraph SQLite Checkpointer，以 thread_id 隔离任务并持久化图状态；安全中断后可以用相同 thread_id 继续。Checkpoint 保存的是恢复所需状态，不等同长期 Memory。当前 Job API 仍是单进程内存任务表，生产化应把 Job 元数据外置，并为有副作用工具增加幂等键、租约和补偿记录。')
    if any(k in x for k in ('rag', '检索', '召回', 'rerank', 'bm25', '向量')):
        if '评' in x or '指标' in x:
            return ('检索层应分别评 Recall@K、MRR、NDCG 和关键文件命中率；生成层评结论正确性、引用忠实度和修复可执行性。RepoMind 已有 eval 数据集、Runner 和 LLM Judge，改动前后应在固定回归集和独立留出集上比较，同时记录延迟、Token 和调用次数。')
        if 'bm25' in x or '混合' in x or '稀疏' in x:
            return ('关键词检索擅长精确标识符，向量检索擅长语义近似，二者分数不能直接相加。可以使用 RRF 按排名融合，再用 CrossEncoder 精排。RepoMind 当前核心实现是多查询向量召回、RRF 与 CrossEncoder，并用 grep/定义查找补足精确符号检索；BM25 是合理的后续增强，但不能说成已经落地。')
        if '切' in x or 'chunk' in x:
            return ('RepoMind 对 Python 优先按 AST 中的函数和类做语义切片，超长块再递归拆分；非 Python 文件使用通用文本切分并保留路径、语言和行号元数据。Chunk 要在语义完整性、召回粒度和 Token 成本之间平衡，使用 overlap 只解决跨边界信息，不应过大。')
        return ('RepoMind 的 RAG 管线是：仓库加载与语义切片→embedding 写入 ChromaDB→HyDE、标题、类型+正文、堆栈等多查询各取 top-5→RRF 融合 top-10→CrossEncoder 精排 top-3→补充依赖源码和调用图。RAG 只提供候选证据，后续 ReAct 仍会通过 grep、读文件、查定义和调用者继续验证。')
    if any(k in x for k in ('mcp', 'function call', 'function call', 'tool', '工具')):
        if any(k in x for k in ('安全', '权限', '危险', '注入')):
            return ('工具安全不能只靠 Prompt。RepoMind 用 JSON Schema/Pydantic 校验结构，工具白名单限制能力，security Issue 进入 human_gate；外部仓库文本只能视为不可信证据。生产化还应加入身份与作用域校验、路径/参数 allowlist、只读与写操作分级、幂等键、审计日志和高风险两阶段确认。')
        if '失败' in x or '超时' in x or '重试' in x:
            return ('先区分瞬时错误、参数错误、权限错误和结果不确定。网络 5xx/限流可指数退避；参数错误回模型修复且限制次数；权限错误不重试；写操作超时先用幂等键或状态查询确认是否已生效。RepoMind 的 LLM 调用使用 tenacity 重试，ReAct 有轮数上限，生产化写工具还需补偿和幂等控制。')
        return ('RepoMind 的工具层自行实现，但采用 MCP 风格抽象：Tool 包含 name、description、JSON Schema parameters 和 executor，通过 ToolRegistry 统一注册。模型以 assistant tool_calls 发起调用，执行结果以 tool role 和对应 tool_call_id 返回；项目提供读文件、grep、查定义、查调用者及 Git log/blame 等 9 类工具。MCP 进一步标准化的是客户端与外部工具服务之间的发现、通信和生命周期。')
    if 'skill' in x:
        return ('Skill 是按需加载的任务方法包，通常包含触发说明、流程规则、参考资料、脚本和模板；MCP 是连接外部工具/资源的协议，Function Calling 是模型表达调用意图的消息机制。RepoMind 当前主要实现 ToolRegistry 和工作流节点，没有完整通用 Skill 运行时；可把稳定的 Issue/PR 分析流程沉淀为 Skill，但应明确这是演进方向。')
    if any(k in x for k in ('评测', '评估', 'eval', 'badcase', 'judge', '指标', '准确率')):
        return ('RepoMind 将评测拆成检索、分析和链路三层：检索看关键文件命中与排序；分析看问题分类、根因、影响文件、严重度和建议是否正确；链路看工具选择、重试恢复、延迟、Token 与成本。项目包含数据集、Runner、LLM Judge 和回归测试。Badcase 需结合 trace 定位到 fetch、retrieval、tool、analysis、reflection 或 report，修复后同时跑固定回归集与留出集，避免只修一类却破坏另一类。')
    if '幻觉' in x:
        return ('RepoMind 通过“检索候选证据+工具继续验证+结构化输出校验+独立 Reflection”降低幻觉。Prompt 要求结论关联具体文件和代码证据；证据不足时降低置信度而不是猜测。高准确性场景还应增加引用一致性校验、规则检查、拒答阈值和人工审批。')
    if 'trace' in x or '可观测' in x:
        return ('Trace 至少记录 thread/job id、节点、模型与 Prompt 版本、输入输出摘要、Token、耗时、工具名与参数摘要、错误类型、重试、状态转移和最终评测分。敏感正文、密钥和个人信息应脱敏或只存哈希；通过分层 trace 可以区分检索没召回、规划选错工具、工具失败、模型误判或报告格式问题。')
    if 'python' in x or 'gil' in x or 'rlock' in x:
        return ('CPython 的 GIL 保证同一进程同一时刻通常只有一个线程执行 Python 字节码，主要简化引用计数和 C 扩展生态；IO 密集任务仍能从多线程/异步中获益，CPU 密集任务更适合多进程或原生计算库。Lock 不可重入，同一线程重复获取会死锁；RLock 记录所有者和递归次数，需对应次数释放。RepoMind 的模型/GitHub 请求偏 IO，适合并发调度。')
    if any(k in x for k in ('tcp', 'udp')):
        return ('TCP 面向连接、可靠、有序、字节流并带拥塞控制；UDP 无连接、尽力而为、保留报文边界、开销低。TCP 通过序列号、确认应答、校验、超时重传、滑动窗口和接收端重排保证可靠有序。')
    if 'http' in x or 'https' in x:
        return ('HTTP 请求通常经历 DNS、TCP 建连、TLS 握手（HTTPS）、发送请求、服务端处理、返回响应和连接复用。HTTPS 通过证书链验证服务端身份，用非对称密钥协商会话密钥，再用对称加密保护传输，并通过完整性校验防篡改。')
    if any(k in x for k in ('mysql', '数据库索引', 'like')):
        return ('MySQL InnoDB 常用 B+Tree 索引，树高低、页式存储和叶子有序使等值、范围及排序查询高效。联合索引需遵守最左前缀；对索引列做函数/隐式类型转换、范围后继续匹配、前导通配符 LIKE、低选择性或优化器判断全表扫描更便宜时可能不走索引。应以 EXPLAIN 和真实数据分布验证。')
    if 'redis' in x or '缓存' in x:
        return ('缓存设计先明确 key、TTL、容量、淘汰和一致性。RepoMind 可缓存 GitHub 元数据、仓库索引版本和稳定检索结果，但当前主要依赖本地/Chroma 持久化。更新可采用 cache-aside：先更新数据源，再删除缓存；热点 key 加随机 TTL 和请求合并，防穿透用参数校验/布隆过滤器，防击穿用互斥或逻辑过期。')
    if any(k in x for k in ('消息队列', 'kafka', 'mq')):
        return ('消息队列用于异步解耦、削峰、重试和事件驱动，不适合把数据库当轮询消息总线。顺序通常靠同一业务 key 进入同一分区并由单消费者串行处理；至少一次语义下消费者必须幂等；生产端确认、持久化副本、消费位点和死信队列共同降低丢失风险。RepoMind 当前 Job API 未引入 MQ，生产化长任务可用它解耦提交与执行。')
    if 'jvm' in x or '垃圾回收' in x or 'java' in x:
        return ('这属于通用 Java/JVM 题，不是 RepoMind 的 Python 实现。回答应区分线程私有的程序计数器、虚拟机栈、本地方法栈和共享的堆、元空间；GC Roots 可来自栈引用、静态字段、常量和 JNI，引擎通过可达性分析判活，再按收集器采用标记清除、复制、整理和分代策略。')
    if any(k in x for k in ('attention', 'transformer', 'qkv', 'kv cache')):
        return ('Self-Attention 将输入分别线性投影成 Q、K、V：Q 表示查询需求，K 表示可匹配特征，V 表示实际聚合内容；softmax(QKᵀ/√d) 得到权重后对 V 加权。多头机制在不同子空间学习关系。相同词的初始词向量可以相同，但叠加位置编码并经过上下文层后表示不同。KV Cache 保存历史层的 K/V，避免自回归生成时重复计算，空间随层数、头维度、序列长度和 batch 增长。')
    if any(k in x for k in ('sft', 'dpo', 'grpo', 'ppo', 'lora', '微调', '强化学习')):
        return ('这是模型训练题，RepoMind 当前通过 API 使用基座模型，没有声称做过 SFT/RL。SFT 用高质量指令—回答学习行为；DPO 用偏好对直接优化相对概率；PPO/GRPO 使用奖励进行策略优化，GRPO 用组内相对优势减少对独立价值模型的依赖；LoRA 冻结原权重，只训练低秩增量以降低显存和训练成本。面试时应明确理论理解与实际经验边界。')
    if any(k in x for k in ('二分', '升序数组')):
        return ('使用二分查找，维护闭区间 [left, right]。普通命中立即返回；若要求第一次出现，命中后记录答案并令 right=mid-1，继续向左收缩。时间复杂度 O(log n)，空间 O(1)。')
    if '股票' in x:
        return ('只能交易一次时，遍历维护历史最低价和当前最大利润；可多次交易时，把每个相邻上涨差值累加，等价于拆分所有上升区间。两种都可 O(n) 时间、O(1) 空间完成。')
    if '最长公共子串' in x:
        return ('使用二维 DP：字符相等时 dp[i][j]=dp[i-1][j-1]+1，否则置 0，并维护最大长度及结束位置。时间 O(mn)，空间可滚动压缩到 O(n)。注意它与最长公共子序列不同，子串要求连续。')
    if '无重复' in x and '子串' in x:
        return ('使用滑动窗口和哈希表记录字符最近位置。右指针扩张，遇到窗口内重复字符时把左边界移动到 last[ch]+1，持续更新最大窗口长度；时间 O(n)，空间 O(字符集)。')
    if '合并区间' in x:
        return ('先按区间起点排序，再遍历：当前区间起点不大于结果末区间终点时合并终点，否则追加新区间。排序主导时间复杂度 O(n log n)，结果数组之外额外空间取决于排序实现。')
    if '大数加法' in x:
        return ('从两个字符串末尾向前模拟竖式加法，维护进位 carry，每轮写入 sum%10，最后反转结果；时间 O(max(m,n))。减法需要先比较绝对值、处理借位并去除前导零。')
    if '岛屿' in x:
        return ('遍历网格，遇到未访问陆地就从该点 DFS/BFS，把连通陆地标记并累计面积；最大岛屿面积取所有连通分量面积最大值。每个格子最多访问一次，时间 O(mn)，空间最坏 O(mn)。')
    if '爬楼梯' in x:
        return ('状态转移为 f(n)=f(n-1)+f(n-2)，因为最后一步只能来自 n-1 或 n-2。用两个变量滚动计算即可，时间 O(n)，空间 O(1)。')
    if '大数' in x or '算法' in x or '手撕' in x:
        return ('先明确输入输出、边界条件和复杂度目标，再选择数据结构与不变量；编码后用空输入、单元素、重复值、极值和规模数据验证。若题目来自 RepoMind 之外，我会把算法解法与项目实现明确区分。')
    if '高并发' in x or 'qps' in x or '限流' in x:
        return ('先做容量估算和读写路径拆分，再通过无状态水平扩展、负载均衡、缓存、异步队列、批处理和连接池提高吞吐；使用令牌桶限流、超时、熔断、隔离和分级降级保护依赖。Agent 场景还要限制每任务 Token、工具轮数和总成本，对高频确定性请求绕过大模型走缓存/规则。')
    return ('我会先澄清业务目标、输入输出、正确性要求、延迟/成本和风险边界，再把确定性步骤交给代码或 Workflow，把需要语义判断的步骤交给模型，并为每一步设计状态、校验、超时、重试、观测和降级。' + unsupported)


def project_detail(q):
    """Return code-grounded RepoMind detail for project-related questions."""
    x = q.lower()
    cat = category(q)
    if cat == 'Agent 项目架构与设计':
        if any(k in x for k in ('pr', 'code review', 'review')):
            return ('在 RepoMind 的 PRGraph 中，fetch_pr、build_project_map、describe_pr_images 和 memory_retrieve_pr 完成准备；prepare_reviews 通过 LangGraph Send 派发 security、style、logic 三个 run_review。三路分别写入不同 State 字段，全部完成后才进入 reflect_pr。若置信度低于 0.7 且 iteration<2，就回 prepare_reviews 重新并行审查；最终只有 critical/high 会触发 REQUEST CHANGES，全部 low/medium 则 APPROVE。')
        if any(k in x for k in ('issue', '路由', '意图', '分类')):
            return ('IssueGraph 的固定链路是 fetch_issue→build_project_map→describe_images→parse_stack_trace→route_issue。route_issue 把问题分为 bug、feature、question、security：bug 先 detect_regression，feature 后续 detect_existing，question 走专门回答节点，security 先进入 human_gate。这样不是让一个“大而全”的 Prompt 处理所有问题，而是让不同问题使用不同证据和分析策略。')
        if any(k in x for k in ('并行', 'send', '协作')):
            return ('项目有两处真实并行：Issue 的 prepare_context 用 Send 派发 entry、rag、git 三路 gather_context_part；PR 的 prepare_reviews 用 Send 派发 security、style、logic 三路 run_review。它们采用“同一节点名+类型字段”模式，各分支写不同 State 字段，避免并行覆盖，汇聚节点只触发一次。')
        if any(k in x for k in ('状态', 'state', 'messages')):
            return ('IssueState 不只保存 messages，而是显式保存 issue 元数据、issue_type、stack_info、project_map、retrieved_context、analysis、confidence、reflection_note、iteration、fix_type 和 HITL 信息。显式字段让路由、并行合并、恢复和评测可以由代码读取；如果只存 messages，控制逻辑就只能反复让模型从自然语言历史中猜状态。')
        return ('RepoMind 选择“确定性图编排+节点内 ReAct”的混合架构。workflow/graph.py 和 workflow/pr_graph.py 负责业务边界、条件路由、并行、循环和 Checkpoint；workflow/nodes/analyze.py 与 analyze_pr.py 内部把工具 Schema 和 executor 交给 LLM Client，模型在最多 6～12 轮内动态取证。这个划分的核心取舍是：外层保证可控性，内层保留代码探索的灵活性。')
    if cat == 'RAG 与上下文工程':
        if any(k in x for k in ('切分', 'chunk', '切片')):
            return ('context/rag_index.py 建索引时优先保留文件路径、语言和代码结构元数据；项目文档明确 Python 使用 AST 语义切片，超长块再拆分。检索结果不仅要能被模型阅读，还要保留行号和 imports，后续 dep_indexer 会直接从 Chroma metadata 读取依赖信息，所以切片元数据也是后续工具链的一部分。')
        if any(k in x for k in ('query', '改写', 'hyde', '多路')):
            return ('workflow/nodes/retrieve.py 的 RAG 分支实际构造多路查询：HyDE 假设文档、Issue 标题、类型+正文以及堆栈信息。每路召回候选后使用 RRF 按排名融合，再调用 cross-encoder/ms-marco-MiniLM-L-6-v2 精排。多查询解决不同表述的召回覆盖，RRF 避免直接相加不同查询分数，CrossEncoder 再解决最终相关性。')
        return ('RepoMind 的完整上下文不是单一向量检索结果。entry 分支先由模型识别入口函数并 grep 定位；rag 分支完成 HyDE、多查询、RRF、CrossEncoder、依赖源码和调用图补充；git 分支只在 regression 场景追踪 log/blame/diff。merge_context 最后合并三类证据。这样设计是因为语义召回、精确符号定位和版本时间线解决的是三类不同问题。')
    if cat == 'Memory、状态与长任务':
        if 'checkpoint' in x or '恢复' in x or '中断' in x:
            return ('workflow/graph.py 和 pr_graph.py 分别用 SqliteSaver 保存 checkpoints.sqlite 与 pr_checkpoints.sqlite；backend/app.py 为每次 Job 生成 job_id，并把它作为 LangGraph thread_id。恢复依赖相同 thread_id 和已持久化 State。需要诚实说明：当前 backend 的 jobs 表仍在进程内存中，所以“图状态可恢复”不等于整个 Web 任务系统已经达到分布式生产级。')
        return ('短期状态存在 TypedDict State 和 ReAct 消息链中；长期案例由 workflow/nodes/memory.py 写入 context/memory_store.py 管理的 Chroma 集合。memory_retrieve 在正式分析前召回类似历史案例，memory_save 在报告生成后保存标题、类型、仓库、结论和置信度。它更接近案例型 episodic memory，而不是完整用户画像或自动自进化系统。')
    if cat == '工具、MCP、Skill 与安全':
        if 'skill' in x:
            return ('RepoMind 当前真实落地的是 ToolRegistry，不是通用 Skills Runtime。tools/base.py 的 Tool 定义 name、description、JSON Schema parameters 和 executor，build_tools 同时返回模型可见 Schema 与本地执行映射。面试中应说：稳定的 Issue/PR 分析方法可以继续封装成 Skill，但不能把尚未实现的触发、按需加载和版本治理说成项目已有能力。')
        if any(k in x for k in ('安全', '权限', '注入', '人工')):
            return ('项目当前最明确的高风险控制是 security Issue 的 human_gate：route_issue 识别为 security 后先中断并等待人工确认，再进入 Memory 和分析链路。代码导航工具以读取和搜索为主，没有开放任意 shell 写操作。边界是：项目尚未实现完整 RBAC/ABAC、租户隔离和写工具两阶段提交，这些应作为生产化改进回答。')
        return ('analyze_bug 注册 read_file、grep_code、list_dir、find_definition、find_callers、git_log_file、git_blame、search_commits、fetch_dependency 等工具。core/llm/client.py 保存 assistant tool_calls 和 tool role 结果，直到模型给最终答案或达到轮数上限。工具不是把源码一次性塞给模型，而是允许模型沿“入口→定义→调用者→Git 历史→依赖源码”逐步建立证据链。')
    if cat == '评测、Badcase 与可观测性':
        if any(k in x for k in ('reflection', '反思')):
            return ('workflow/nodes/reflect.py 把 critic 与分析节点分离，要求检查证据是否充分、是否存在其他可能原因，并输出 confidence 与 note。confidence 被夹在 0～1，iteration 每次加 1；graph.py 在 confidence<0.7 且 iteration<2 时按 issue_type 回对应分析节点。它是有停止条件的质量回路，不是无限“再想一遍”。')
        return ('eval/ 下有 dataset、runner、judge 和 report：Runner 用独立 thread_id 执行真实 Graph，收集 AI 结果；Judge 比较根因、文件、严重度和修复可执行性；report 汇总分数与置信度。tests/regression 还保存 smoke cases 与 baselines。项目可以证明“有离线评测和回归框架”，但不能把归档结果直接包装成线上用户指标。')
    return ''


def set_font(run, size=10.5, bold=None, color=None, name='PingFang SC'):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        rfonts.set(qn('w:' + a), name)
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color is not None: run.font.color.rgb = color


PERSONAL_EXCLUDES = (
    '实习', '最快到岗', '到岗时间', '实习时长', '离开当前公司', '离职原因',
    '职级预期', '团队规模', '个人规划', '学校里做', '实验室做', '论文拷打',
    '为什么投递', '为什么要投', '未来的发展规划', '对这边的意向',
)
questions = [
    q for q in extract_questions()
    if category(q) != '个人经历与行为问题'
    and not any(term in q for term in PERSONAL_EXCLUDES)
]
groups = {}
order = ['Agent 项目架构与设计','RAG 与上下文工程','Memory、状态与长任务','工具、MCP、Skill 与安全','评测、Badcase 与可观测性','工程基础与系统设计','大模型基础与训练','算法与数据结构','场景题与其他']
for q in questions:
    groups.setdefault(category(q), []).append(q)

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.72)
sec.left_margin = sec.right_margin = Inches(0.78)
sec.header_distance = sec.footer_distance = Inches(0.35)
normal = doc.styles['Normal']
normal.font.name = 'PingFang SC'; normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')
normal.paragraph_format.space_after = Pt(4); normal.paragraph_format.line_spacing = 1.18
for nm,sz,bef,aft,col in [('Heading 1',17,15,7,BLUE),('Heading 2',12.5,10,4,DARK),('Heading 3',11,7,3,DARK)]:
    st=doc.styles[nm]; st.font.name='PingFang SC'; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=col
    st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC'); st.paragraph_format.space_before=Pt(bef); st.paragraph_format.space_after=Pt(aft); st.paragraph_format.keep_with_next=True
ans_style=doc.styles.add_style('Reference Answer',WD_STYLE_TYPE.PARAGRAPH)
ans_style.font.name='PingFang SC'; ans_style.font.size=Pt(10.5); ans_style._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC')
ans_style.paragraph_format.left_indent=Inches(0.16); ans_style.paragraph_format.space_after=Pt(6); ans_style.paragraph_format.line_spacing=1.18
shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),'F4F6F9'); ans_style._element.get_or_add_pPr().append(shd)

h=sec.header.paragraphs[0]; set_font(h.add_run('RepoMind · Agent 开发完整面试手册'),9,color=GRAY)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_font(f.add_run('RepoMind 项目参考回答'),9,color=GRAY)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(90); p.paragraph_format.space_after=Pt(12)
set_font(p.add_run('RepoMind Agent 开发面试问答'),24,True,DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run('字节 Agent 开发二面面经 · 严格去重主题版'),14,False,BLUE)
p.paragraph_format.space_after=Pt(45)
doc.add_heading('整理规则',1)
doc.add_paragraph(f'本手册从原文件第二个面经之后提取技术题，共保留 {len(questions)} 道严格去重后的独立问题。只有去掉序号、空格以及统一全半角问号后文字完全一致的题目才合并；限定条件、追问角度或措辞存在差异时均分别保留。')
doc.add_paragraph('已删除实习经历、最快到岗、实习时长、离职原因、职级预期等个人求职问题；保留 Java/JVM、算法及其他技术题。项目题以 RepoMind 当前代码为准，未实现能力会明确写成改进方案。')
doc.add_page_break()

for cat in order:
    qs=groups.get(cat,[])
    if not qs: continue
    doc.add_heading(f'{cat}（{len(qs)} 题）',1)
    for i,q in enumerate(qs,1):
        doc.add_heading(f'{i}. {q}',2)
        detail = project_detail(q)
        unified = detail if detail else answer(q)
        p=doc.add_paragraph(style='Reference Answer')
        set_font(p.add_run('我的回答是：'),10.5,True,DARK)
        set_font(p.add_run(unified),10.5)

doc.core_properties.title='RepoMind Agent 开发面试问答'
doc.core_properties.subject='每题一个统一详细回答；保留技术题，删除个人求职问题'
doc.core_properties.author='RepoMind'
doc.save(OUT)
print({'output':str(OUT),'questions':len(questions),'categories':{k:len(v) for k,v in groups.items()}})
