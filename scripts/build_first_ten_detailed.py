from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第一批10题.docx')
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY = RGBColor(95, 103, 112)

QA = [
('1. 介绍一下你这个 Agent 项目的背景、架构和技术选型', [
'我的回答是：RepoMind 是一个面向 GitHub 仓库的代码分析 Agent，主要解决两个问题。第一个是 Issue 根因分析：用户输入 Issue URL，系统自动判断问题类型，定位相关代码和根因，并给出影响文件、严重程度和修复建议。第二个是 PR Code Review：系统读取 PR diff，从 security、logic、style 三个视角并行审查，最后给出 APPROVE 或 REQUEST CHANGES。',
'我做这个项目的出发点是，真实的代码问题分析并不是把 Issue 文本交给大模型就结束了。开发者通常要查看仓库目录、搜索符号、追踪调用链、检查 Git 历史，有时还要看依赖库源码。单次 LLM 调用既缺少完整上下文，也不能主动验证自己的判断，所以我把检索、工具调用和工作流编排组合成了一条可控的分析链路。',
'架构上我采用“外层 Workflow、内层 Agent Loop”。外层使用 LangGraph 管理状态、条件路由、并行任务、Reflection 重试、人工审批和 Checkpoint；内层分析节点使用 ReAct，让模型根据当前证据自主调用 read_file、grep_code、find_definition、find_callers、git_blame 等工具。RAG 使用 ChromaDB、Sentence Transformer、多查询召回、RRF 融合和 CrossEncoder 精排；长期案例 Memory 也保存在 ChromaDB；工作流状态通过 SQLite Checkpointer 持久化。',
'这个项目的核心不是堆技术名词，而是做职责划分：确定性的流程交给 LangGraph，需要语义判断和动态探索的部分交给大模型，精确读写和查询交给工具，质量检查交给独立 Reflection。这样既保留 Agent 的灵活性，又避免整个系统完全由模型自由发挥。'
]),
('2. 你之前的 Agent 项目用的什么架构？LangGraph 还是自研？是 Master+Sub Agent 还是 Workflow？为什么这么选型？', [
'我的回答是：RepoMind 的主体是基于 LangGraph 的有状态 Workflow，不是完全自研编排框架，也不是典型的 Master 加多个 Sub Agent 自由协商模式。项目里确实存在并行分析任务，但它们由图中的确定性节点统一调度，不依赖多个 Agent 通过自然语言反复对话。',
'选择 LangGraph 的原因是 RepoMind 的控制流比较复杂。Issue 要先获取数据、构建项目地图、解析图片和堆栈，再分类成 bug、feature、question 或 security；不同类型走不同分支。上下文收集需要 entry、RAG、Git 三路并行；分析后置信度不足还要回到对应节点重试；security 类型继续执行前还要人工确认。这些条件路由、并行、循环和中断恢复，用普通线性 Chain 表达会比较别扭。',
'我没有完全自研，是因为状态合并、并行汇聚、Checkpoint 和中断恢复都有较高工程成本，LangGraph 已经提供了合适的基础能力。但业务节点、State 字段、路由条件、工具注册、Prompt、检索管线和停止条件都是项目自己设计的，所以不是简单套一个现成 Agent。',
'我也没有使用 Master+Sub Agent 作为主架构，因为 Issue 分析的步骤之间存在明确依赖，使用 Workflow 更容易保证可观测性和可恢复性。PR 的 security、style、logic 虽然并行执行，但采用的是同一个 run_review 节点加 review_type 路由，三路写入不同 State 字段。我的取舍是：只有任务边界天然独立时才并行拆分，不为了“多 Agent”这个形式增加协商开销。'
]),
('3. 这是单 Agent 还是多 Agent？各个 Agent 的核心任务和分工是怎么划分的？', [
'我的回答是：如果按运行形态来讲，RepoMind 更接近“多个专职分析角色组成的 Workflow”，而不是多个自治 Agent 互相聊天。Issue 和 PR 分别有独立的 Graph；在 Graph 内，大部分节点是确定性处理节点，真正使用 ReAct 工具循环的分析节点才具备较强的 Agent 特征。',
'Issue 链路会根据问题类型分工。bug 和 security 最终进入 analyze_bug，重点追踪根因、调用链、回归提交和修复类型；feature 进入 analyze_feature，先判断功能是否已经存在，再评估改动位置和兼容性；question 进入 answer_question，更偏向从现有实现和文档中给出解释。它们不是同时竞争一个答案，而是由 route_issue 先确定任务类型，再选择专门分析策略。',
'PR 链路是最接近多 Agent 的部分。prepare_reviews 使用 LangGraph Send 同时派发 security、style、logic 三路 run_review。security 关注注入、权限、敏感信息和危险调用；style 关注可读性、命名和维护性；logic 关注边界条件、错误处理和行为回归。每一路使用不同 Prompt 和工具集合，并写入 security_issues、style_issues、logic_issues，最后由 reflect_pr 统一检查。',
'这样划分的原则不是按“角色听起来是否高级”，而是看目标是否独立、能否并行、输出能否结构化合并。如果多个 Agent 需要频繁互相解释上下文，通信成本可能大于收益，这时一个 Agent 配合清晰 Workflow 反而更稳定。'
]),
('4. LangGraph 和 LangChain 有什么区别？为什么项目选择 LangGraph？', [
'我的回答是：LangChain 更像一套通用的 LLM 应用组件，提供 Prompt、模型调用、Retriever、Tool 和 Agent 等抽象；传统 Chain 更适合顺序明确的线性流程。LangGraph 则更强调有状态图编排，把节点、边、条件路由、循环、并行和 Checkpoint 作为一等能力。两者不是互斥关系，LangGraph 可以使用 LangChain 生态中的模型和工具组件。',
'RepoMind 选择 LangGraph，是因为流程不是一条固定直线。Issue 分类后有四条分支；security 需要先经过 human_gate；bug 要做 regression 检测和 fix_type 分类；feature 要先检测能力是否已经存在；上下文收集要并行；Reflection 失败还要根据 issue_type 回到不同分析节点。PR 也有三路并行 Review 和低置信度回退。',
'如果只用普通 Chain，我需要在链外自己维护大量 if/else、循环次数、并行结果和恢复状态，最终还是会写出一个隐式状态机。LangGraph 的价值是把这些控制关系显式化，图结构可以直接看出任务如何流转，State 也可以被 Checkpointer 持久化。',
'它的代价是 State 设计和并行写入需要更谨慎。例如三路上下文不能同时覆盖 code_context，所以我让它们分别写入 _entry_ctx、_rag_ctx、_git_ctx，再由 merge_context 合并。也就是说，框架解决了编排能力，但状态字段和并发语义仍然需要开发者自己设计。'
]),
('5. 了解 ReAct 吗？讲一下 ReAct 范式，以及项目里是怎么实现的', [
'我的回答是：ReAct 的核心是让模型在推理和行动之间循环。模型不是一次性直接生成最终结论，而是先根据当前上下文判断下一步需要什么证据，发起工具调用；系统执行工具并把结果返回；模型再根据新的观察决定继续调用工具还是输出答案。可以把它概括为 Reason、Act、Observe 的循环。',
'RepoMind 在 analyze_bug、analyze_feature、answer_question 以及 PR Review 节点中使用了这种模式。以 bug 分析为例，模型可能先调用 grep_code 搜索 Issue 中的函数名，再用 find_definition 跳到定义，用 read_file 阅读附近实现，用 find_callers 向上追踪调用关系；如果判断是回归问题，还可以调用 git_log_file、git_blame 或 search_commits 检查时间线；如果根因在依赖库，则使用 fetch_dependency 检索依赖源码。',
'工程实现上，tools/base.py 把每个工具定义成 name、description、JSON Schema parameters 和 execute 函数。build_tools 同时生成模型可见的 Tool Schema 和本地 executor 映射。LLM Client 保存 assistant 返回的 tool_calls，再把执行结果以 tool role 和对应 tool_call_id 加回消息链，保证下一轮模型能够看到完整行动历史。',
'ReAct 的风险是重复搜索和无限循环，所以项目不会完全放任模型。不同节点设置 6 到 12 轮的最大工具轮数；工具集合按任务裁剪；外层还有 Reflection 和 iteration 上限。我的理解是，ReAct 负责提高探索上限，但预算、权限和停止条件必须由确定性代码控制。'
]),
('6. 你觉得 Planning 是 Agent 框架做的，还是大模型做的？', [
'我的回答是：Planning 不应该简单归到某一方，比较合理的是分层处理。框架适合做确定性的宏观规划，大模型适合做信息不完整情况下的动态微观规划。',
'RepoMind 的宏观计划由 LangGraph 决定。例如输入 Issue 后一定先 fetch_issue，再 build_project_map、describe_images、parse_stack_trace 和 route_issue；分类为 security 就先进入 human_gate；完成上下文收集后，bug 先 classify_fix，feature 先 detect_existing。这些步骤依赖稳定、需要恢复，应该写在框架中，而不是每次都让模型重新规划。',
'进入 analyze_bug 以后，下一步读哪个文件、搜索什么符号、是否查看调用者、是否追踪 Git 历史，就很难提前写死。因为后一步取决于前一个工具返回的证据，所以这里由 ReAct 模型动态规划更合适。',
'我的总结是：框架负责“必须经过哪些阶段、最多走多远、失败后回哪里”，模型负责“在当前阶段具体查什么、如何根据观察调整下一步”。如果全部交给框架，系统会僵化；如果全部交给模型，路径、成本和恢复都会失控。'
]),
('7. Agent 的任务规划在 RepoMind 里是怎么做的？', [
'我的回答是：RepoMind 没有让模型一开始生成一份很长的 To-Do List，然后机械执行到底，而是采用“预规划加增量规划”的方式。因为代码问题在读到真正实现之前，往往无法提前知道完整调用链，过早生成详细计划很容易建立在错误假设上。',
'第一层是问题级规划。route_issue 判断 bug、feature、question、security，决定后续需要什么证据。第二层是分析前规划：bug 通过 detect_regression 判断是否属于版本回归，再由 classify_fix 判断更可能是 code_fix、dependency_upgrade 还是 config_fix；feature 通过 detect_existing 先确认仓库中是否已有类似能力。规划结果会改变后续 Prompt 的关注点。',
'第三层是执行期增量规划。进入 ReAct 后，模型每一轮只根据已有证据决定下一步工具。例如先定位入口，再沿定义和调用者扩展；如果发现版本差异，再转向 Git；如果发现问题来自第三方包，再查依赖源码。这样计划会随着观察结果更新，而不是坚持最初假设。',
'外层代码还负责计划预算：工具轮数有限，Reflection 最多触发两次重分析。这个设计的取舍是，项目没有实现复杂的显式计划树或 Tree-of-Thought，而是用业务节点完成粗粒度规划，用 ReAct 完成细粒度探索，更符合当前代码分析场景。'
]),
('8. Agent 请求从进入服务到返回结果，完整链路是什么？', [
'我的回答是：以 Issue 为例，后端收到 URL 后先创建 job_id，并把它同时作为 LangGraph thread_id。任务以异步 Job 形式执行，前端可以根据 job_id 查询状态。Graph 从 fetch_issue 开始，通过 GitHub API 获取标题、正文、labels、评论和关联 PR。',
'随后 build_project_map 构建仓库结构，describe_images 处理 Issue 图片，parse_stack_trace 提取异常堆栈，route_issue 把任务分成 bug、feature、question 或 security。security 会先进入 human_gate；bug 会额外判断是否回归。memory_retrieve 召回历史相似案例，prepare_context 确保仓库已经 clone 和建立索引。',
'上下文阶段使用 Send 并行执行三路任务：entry 路径识别入口并用 grep 精确定位；RAG 路径完成多查询召回、RRF、CrossEncoder、依赖和调用图补充；Git 路径为回归问题补充提交历史。三路结果在 merge_context 汇总后，根据 Issue 类型进入 analyze_bug、analyze_feature 或 answer_question。',
'分析节点通过 ReAct 工具循环继续取证，随后 reflect 独立检查结论。置信度低于 0.7 且 iteration 小于 2 时回到对应分析节点，否则 generate_report 生成结构化报告，memory_save 保存案例，Graph 结束。后端把 final_report、confidence 和 error 等字段写入 Job 结果。需要说明的是，Graph 状态有 SQLite Checkpoint，但当前 Job 表仍在进程内存中，这是生产化需要继续完善的地方。'
]),
('9. 多 Agent 系统为什么经常比单 Agent 效果更差？', [
'我的回答是：多 Agent 并不天然比单 Agent 更强。它增加了通信、状态同步和决策冲突。如果任务本身没有清晰边界，多个 Agent 会重复检索同一份资料、互相传递压缩后的二手信息，甚至对同一个结论反复争论，最后延迟和 Token 成本上升，信息却在传递中丢失。',
'另外，每个 Agent 都有自己的 Prompt 和上下文，局部判断可能正确，但全局目标不一致。自然语言通信也不是可靠协议：上一个 Agent 输出的一段总结，可能漏掉下一个 Agent 真正需要的代码证据。出错后还更难定位，到底是某个 Agent 推理错误、调度错误，还是通信过程发生了信息变形。',
'RepoMind 因此没有把 Issue 分析拆成很多会聊天的角色。Issue 先通过 Workflow 路由到一个专职分析节点，节点内部直接使用工具访问一手代码证据。PR 只有 security、style、logic 三个目标天然独立，所以才并行拆分，而且它们不互相聊天，而是输出结构化 issues，最后统一 Reflection。',
'我的判断标准是：任务是否可以清晰分解、分支之间是否低依赖、是否真的能并行、输出是否可以结构化合并。如果答案是否定的，优先使用单 Agent 加工具；如果存在天然独立的专业视角，再考虑多 Agent。多 Agent 是组织复杂度的手段，不是提升效果的默认答案。'
]),
('10. 如何防止 Agent 工作流进入死循环？', [
'我的回答是：不能只在 Prompt 里告诉模型“不要重复”，循环控制必须落在确定性代码中。我会从步数、状态变化、预算、错误类型和图路由五个层面控制。',
'RepoMind 已经实现了两层硬停止条件。第一层是 ReAct 工具循环：不同分析节点设置最大轮数，达到上限后必须结束工具探索并输出当前结论。第二层是 Reflection：reflect 每次把 iteration 加一，只有 confidence 低于 0.7 且 iteration 小于 2 才允许回到分析节点，因此最多重分析两次。PR 的 reflect_pr 也使用同样阈值，retry 回 prepare_reviews，而不是重新执行前面的抓取和 Memory。',
'如果继续做生产化，我还会增加停滞检测。为每次工具调用计算 tool_name 加规范化参数的指纹；如果连续多次调用同一工具且参数相同，直接阻止。还可以比较关键 State 的变化量，如果多轮之后 affected_files、root_cause 或证据集合没有新增，就判定 Agent 已经停滞。',
'错误也要分类处理：网络超时可以指数退避，参数错误允许有限修复，权限错误直接停止，检索无结果只允许有限次数 Query Rewrite。最后再加总 Token、总工具成本和墙钟时间预算。这样合法的探索仍然可以继续，但不会因为模型反复“再试一次”而无限运行。'
]),
]

def set_font(run, size=11, bold=None, color=None):
    name='PingFang SC'; run.font.name=name; run.font.size=Pt(size)
    rf=run._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('ascii','hAnsi','eastAsia','cs'): rf.set(qn('w:'+a),name)
    if bold is not None: run.bold=bold
    if color is not None: run.font.color.rgb=color

doc=Document(); sec=doc.sections[0]
sec.page_width,sec.page_height=Inches(8.5),Inches(11)
sec.top_margin=sec.bottom_margin=Inches(.82); sec.left_margin=sec.right_margin=Inches(.88)
normal=doc.styles['Normal']; normal.font.name='PingFang SC'; normal.font.size=Pt(11); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for nm,sz,bef,aft,col in [('Heading 1',16,16,8,BLUE),('Heading 2',13,12,6,DARK)]:
    st=doc.styles[nm]; st.font.name='PingFang SC'; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=col; st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC'); st.paragraph_format.space_before=Pt(bef); st.paragraph_format.space_after=Pt(aft); st.paragraph_format.keep_with_next=True
h=sec.header.paragraphs[0]; set_font(h.add_run('RepoMind · Agent 开发面试问答'),9,color=GRAY)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_font(f.add_run('第一批 · 10 道项目题'),9,color=GRAY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70); p.paragraph_format.space_after=Pt(10); set_font(p.add_run('RepoMind Agent 开发面试问答'),24,True,DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(35); set_font(p.add_run('第一批：Agent 项目架构与设计（10 题）'),14,color=BLUE)
doc.add_paragraph('本批回答逐题结合 RepoMind 当前代码整理，使用适合面试口述的完整回答。已实现能力、设计取舍和当前边界会明确区分。')
doc.add_page_break()
for q,paras in QA:
    doc.add_heading(q,level=1)
    for i,t in enumerate(paras):
        p=doc.add_paragraph(t)
        p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第一批10题'
doc.core_properties.author='RepoMind'
doc.save(OUT)
print(OUT)
