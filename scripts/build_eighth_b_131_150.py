from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第八批B_131-150.docx')

QA = [
('131. Transformer 为什么比 RNN 更适合大模型？它的核心结构是什么？', [
'我的回答是：Transformer 的核心优势是把序列建模从逐步递归改成基于 Attention 的全局信息交互。RNN 在第 t 个位置必须依赖前 t−1 个位置的隐藏状态，训练时难以在序列维度并行，而且长距离信息要经过很多次状态传递，容易衰减。Transformer 可以让一个位置直接关注任意历史位置，训练阶段能一次处理整段 Token，因此更容易扩大数据量、参数量和计算规模。',
'一个标准 Transformer Block 主要由多头自注意力、前馈网络、残差连接和归一化组成。Attention 负责在 Token 之间交换信息，FFN 对每个位置做非线性特征变换；多头机制让不同子空间学习语法、指代、位置或任务关系。现代大模型通常采用 Decoder-only 架构，并使用因果 Mask，保证生成第 t 个 Token 时只能读取前面的 Token。残差和归一化则改善深层网络的梯度传播与数值稳定性。',
'RepoMind 并没有从头训练 Transformer，而是通过统一 LLM Client 调用具备工具调用和结构化输出能力的模型。理解 Transformer 对项目仍有价值：Issue、代码片段、工具结果都会占用上下文窗口；长上下文会提高 Attention 计算和推理成本。因此 RepoMind 先用 RAG 缩小候选范围，再由 Agent 主动读取必要文件，而不是把整个仓库一次性塞给模型。'
]),
('132. Self-Attention 的计算过程是什么？为什么要除以 √d_k？', [
'我的回答是：Self-Attention 先把输入矩阵 X 分别乘以三个参数矩阵得到 Q、K、V。Q 表示当前位置希望查询什么，K 表示每个位置可以用什么特征被匹配，V 是真正用于聚合的信息。计算公式是 Attention(Q,K,V)=softmax(QKᵀ/√d_k)V：先算每个 Query 与所有 Key 的相似度，经 Mask 和 softmax 变成权重，再对 Value 做加权求和。',
'除以 √d_k 是为了控制点积的方差。若 Q、K 各维近似独立且方差为 1，d_k 维点积的方差会随 d_k 增大；数值过大时 softmax 很快饱和成接近 one-hot，非最大位置的梯度会非常小，训练不稳定。缩放后 logits 的量级更稳定。这不是为了改变排序，而是为了让 softmax 处在更有梯度的数值区间。实现中还要先减去最大值，避免指数运算溢出。',
'代码分析场景中，Attention 能让模型把 Issue 中的报错、函数名与后文代码建立联系，但它并不等于可靠检索。上下文过长时，相关信息可能被噪声稀释，且模型无法看到窗口外内容。RepoMind 因此通过向量召回、RRF 和 CrossEncoder 先选证据，再让 Self-Attention 在较小且相关的上下文内完成推理。'
]),
('133. Multi-Head Attention 为什么需要多个头？头越多越好吗？', [
'我的回答是：多头 Attention 会把隐藏维度投影到多个子空间，各个头独立计算注意力，再拼接并做输出投影。这样模型可以在同一层并行表达不同关系，例如一个头关注局部语法，一个头关注远距离依赖，另一个头关注变量定义与使用。相比单个大头，多头提供了更丰富的关系模式和归纳偏置。',
'头数并不是越多越好。在总隐藏维度固定时，头越多意味着单头维度越小，单个头的表达能力可能不足；很多头还会学习重复或无效模式，并增加 Kernel 调度、KV Cache 元数据和通信成本。选择要结合 hidden_size、head_dim、硬件和模型规模。推理模型常使用 MQA 或 GQA，让多个 Query 头共享较少的 Key/Value 头，在保持大部分质量的同时减少 KV Cache 和内存带宽。',
'RepoMind 不控制底层模型的头数，但模型选型时会间接受这些结构影响。项目中的代码分析请求往往上下文较长，吞吐更容易受 KV Cache 和显存约束；若自部署模型，我会优先评估采用 GQA 的模型，并以真实 Issue 数据比较正确率、首 Token 延迟、每秒 Token 和并发容量，而不会只按参数量选模型。'
]),
('134. Transformer 中的位置编码有什么作用？RoPE 与绝对位置编码有什么区别？', [
'我的回答是：纯 Self-Attention 对输入顺序本身不敏感，如果交换 Token，同时交换对应向量，计算结构不会知道哪个词在前。因此必须注入位置信息。绝对位置编码把位置向量直接加到 Token Embedding 上，模型学习“第几个位置”；经典实现可以是固定正弦编码，也可以是可学习向量，但超出训练长度后的泛化通常有限。',
'RoPE 不把位置向量简单相加，而是在计算 Attention 前，按位置对 Q、K 的二维子空间做旋转。两个位置的点积会自然包含相对距离，因此更适合表示相对位置关系，也便于配合 KV Cache。扩展上下文时可使用位置插值或频率缩放，但把窗口机械拉长并不保证模型真的会利用远端信息，可能出现短上下文能力下降或长距离检索准确率不足。',
'RepoMind 使用现成模型，没有修改位置编码。项目层面的应对是不要把“支持 128K”理解成可以无成本输入整个仓库：索引按照代码语义切片，检索只注入 top 证据，工具按需读取完整文件。这样既减少位置外推风险，也避免重要根因被大量无关代码淹没。'
]),
('135. Transformer 的 LayerNorm、残差连接和 Pre-Norm 分别解决什么问题？', [
'我的回答是：残差连接让每个子层学习对输入的增量，而不是从头重建表示，梯度也可以沿恒等路径传播，是训练深层网络的基础。LayerNorm 在单个样本的隐藏维度上归一化，不依赖 batch 大小，适合长度变化和小 batch 的语言模型。它能控制激活尺度，但并不能替代合理初始化、学习率和梯度裁剪。',
'Post-Norm 是先执行子层和残差相加，再做归一化；Pre-Norm 是先归一化，再进入 Attention/FFN，随后与残差相加。Pre-Norm 的恒等残差路径更直接，深层训练通常更稳定，因此许多 Decoder-only 模型采用 Pre-Norm 或 RMSNorm。代价是层数很深时表示尺度可能累积，所以还会配合最终归一化、残差缩放等设计。',
'这些属于 RepoMind 所调用基础模型的内部机制，项目没有训练网络层。RepoMind 自己更关注应用层稳定性：限制 Agent 轮数、约束结构化输出、保存 Checkpoint、区分可重试错误。面试中我会把两层可靠性分开，避免用模型内部归一化来解释 Agent 工作流层面的状态一致性。'
]),
('136. Decoder-only 模型如何训练和生成？Causal Mask 起什么作用？', [
'我的回答是：Decoder-only 模型通常做自回归语言建模。训练样本是一段 Token，输入右移后预测下一个 Token，所有位置的损失可以并行计算；例如输入“定位 根因”，模型分别预测后续 Token。Causal Mask 会屏蔽当前位置右侧的 Key，防止训练时偷看到答案。训练阶段虽然位置并行，但每个位置仍只使用历史上下文。',
'生成阶段则是串行的：模型根据 Prompt 预测一个 Token，把它追加到序列，再预测下一个。为了避免每轮重算全部历史 Token 的 K、V，服务会使用 KV Cache。停止条件除了 EOS，还包括最大长度、结构化输出完成、工具调用等；temperature、top-p 会影响采样随机性，而代码分析和 JSON 输出通常使用较低温度以提高一致性。',
'RepoMind 的 ReAct 循环建立在这种生成机制之上：模型一轮可能产生 Tool Call，系统执行工具后把结果作为新消息送回模型，再继续生成。外层 LangGraph 决定节点路由和重试，模型自身只负责当前轮决策。因此即使模型是自回归的，也不能把无限循环的控制权全部交给它。'
]),
('137. KV Cache 是什么？它为什么能加速推理，又带来什么代价？', [
'我的回答是：自回归生成时，历史 Token 在每一层对应的 K、V 不会因为新 Token 到来而改变。KV Cache 把这些张量保存下来，新一步只计算新 Token 的 Q、K、V，并让新 Q 与缓存中的所有 K 做 Attention。这样避免对历史前缀重复投影，显著降低每 Token 生成计算量。Prompt 首次处理叫 prefill，后续逐 Token 阶段叫 decode，两者的算力和带宽特征不同。',
'代价是显存占用随 batch、序列长度、层数、KV 头数和 head_dim 近似线性增长；长上下文和高并发时，KV Cache 往往比权重更快成为瓶颈。它还需要分页管理、请求调度和回收，碎片化会降低利用率。MQA/GQA、低精度 KV、PagedAttention 和前缀缓存都可降低或复用成本，但可能引入质量、命中率或实现复杂度的权衡。',
'RepoMind 当前调用外部模型 API，不自行管理 KV Cache。应用侧可以减少无效 Cache 压力：RAG 控制上下文长度，工具返回进行截断和结构化，只把真正需要的历史消息传入下一轮。若未来自部署，需要基于 RepoMind 的多轮工具调用特征测试并发，而不是仅用短 Prompt 的公开吞吐数据。'
]),
('138. FlashAttention 的原理是什么？它是否改变了 Attention 的结果？', [
'我的回答是：普通 Attention 会显式生成长度为 n 的 QKᵀ矩阵，写入显存，再读取做 softmax 和与 V 相乘；长序列时大量时间耗在高带宽显存读写，而不只是乘法。FlashAttention 使用分块计算，把 Q、K、V 的小块放进片上 SRAM，并通过 online softmax 维护每行最大值与归一化因子，避免物化完整 n×n 注意力矩阵。',
'它属于 IO-aware 的精确 Attention 算法，在相同精度和 Mask 下数学结果与标准 Attention 等价，只会有浮点计算顺序导致的微小误差，不是稀疏近似。它显著降低中间显存并提高训练和 prefill 效率，但 Attention 的理论计算量仍大致是 O(n²)；decode 阶段往往更受 KV Cache 读取和内存带宽限制。',
'RepoMind 没有自行集成 FlashAttention，因为模型推理由供应商负责。如果转为自托管，我会先确认模型、CUDA、dtype 和 Mask 是否兼容，再用真实的代码上下文长度测 prefill 延迟与峰值显存。应用层仍要保留检索和上下文压缩，因为 Kernel 优化不能解决无关上下文导致的推理质量下降。'
]),
('139. 什么是 LoRA？为什么它能用少量参数完成微调？', [
'我的回答是：LoRA 冻结预训练权重 W，不直接更新完整的 d×k 矩阵，而是在前向中增加低秩增量 ΔW=BA，其中 A、B 的秩 r 远小于 d、k。训练时只优化 A、B，参数量从 d×k 降到 r(d+k)。它基于一个经验假设：针对具体下游任务所需的权重变化通常位于较低维的子空间。',
'LoRA 常应用于 Attention 的 q、k、v、o 投影，也可以扩展到 FFN。关键超参数包括 rank r、缩放 alpha、dropout、target modules、学习率和数据质量。rank 过小可能欠拟合，过大则增加显存并不一定提升泛化。部署时可把 Adapter 合并到基础权重以减少额外计算，也可动态挂载多个 Adapter，但多租户调度和版本管理会更复杂。',
'RepoMind 当前没有进行 LoRA 训练，主要通过 Prompt、RAG、工具和工作流提升效果。原因是代码根因分析高度依赖仓库实时证据，单靠微调无法记住不断变化的代码。只有在积累足够多、经过审核的 Issue—证据—结论数据，并确认错误来自模型行为而非检索后，我才会考虑用 LoRA 学习稳定的输出格式、工具选择或审查偏好。'
]),
('140. QLoRA 与 LoRA 有什么区别？量化训练为什么仍能工作？', [
'我的回答是：QLoRA 的核心是在 LoRA 基础上把冻结的基础模型权重量化到 4 bit，以减少加载权重所需显存；可训练的 LoRA 参数和关键计算通常仍使用 BF16/FP16。经典方案使用适合近似正态权重分布的 NF4、double quantization，以及 paged optimizer。反向传播会经过反量化后的计算路径更新 Adapter，但不会更新 4 bit 基座权重。',
'它能工作是因为量化误差主要作用于冻结基座，而任务适配由较高精度的低秩参数完成；预训练模型本身也有一定冗余。但 QLoRA 不是无损方案，效果取决于模型规模、量化格式、目标任务和数据。训练显存下降不等于训练速度必然更快，反量化和 Kernel 支持可能造成额外开销；推理部署也要确认训练格式与推理引擎兼容。',
'RepoMind 没有使用 QLoRA。若资源有限且确实需要训练工具调用模型，我会先建立可复现基线，再比较 Prompt/RAG、LoRA 和 QLoRA 的任务成功率、JSON 合法率、工具选择正确率以及显存成本。没有评测集就直接量化微调，很难判断提升来自训练还是样本泄漏。'
]),
('141. SFT 是什么？高质量 SFT 数据应该如何构造？', [
'我的回答是：SFT 即监督微调，用“输入—理想输出”样本对预训练模型做交叉熵训练，使模型学会指令遵循、格式、领域表达或工具调用行为。它更像把模型分布拉向示范答案，不直接优化人类偏好，也不保证事实正确。数据质量通常比盲目扩大数量更重要，错误或风格冲突的示范会被模型稳定复现。',
'构造 Agent 数据时不能只保留最终答案，还应包含任务上下文、合法 Tool Schema、正确工具及参数、工具观察和停止条件。样本应覆盖正常路径、无结果、工具超时、参数错误、证据不足和拒绝高风险动作，并划分仓库级而非随机行级训练/验证集，防止相似代码泄漏。对于思维过程，应优先监督可审计的计划、证据引用和动作，不依赖不可验证的长篇内心推理。',
'RepoMind 目前没有 SFT，已有的 Reflection、结构化结果和人工审核可用于沉淀候选数据，但不能直接把模型自产结果回灌。需要先去重、脱敏、校验代码证据与 commit，并由人工确认结论。如果将来微调，我会先聚焦工具选择和结构化报告这类可客观评测的目标，而不是声称训练一个模型就能替代 RAG。'
]),
('142. 指令微调时如何设计 loss mask？为什么不能对所有 Token 都计算损失？', [
'我的回答是：一条对话训练样本通常包含 system、user、assistant 和 tool 消息。语言模型前向仍读取完整序列，但训练时可用 loss mask 只对希望模型学习生成的 Token 计算交叉熵，例如 assistant 的答案或 Tool Call；用户问题、系统规则和工具返回用于条件输入，标签设为 ignore_index。否则模型会浪费容量学习复述用户文本，还可能把工具返回误当作需要生成的内容。',
'多轮 Agent 数据需要更细致：assistant 的工具调用通常应计入 loss，tool observation 不计入；最终答案计入。是否训练自然语言 reasoning 要根据数据可信度决定。还要处理模板边界、特殊 Token、padding mask 和序列截断，避免把答案开头截掉或在不同 Chat Template 下标签错位。可通过随机抽样解码带 mask 的 Token 做数据单测。',
'RepoMind 未进行这类训练，但它的 ReAct Trace 能提供消息角色清晰的原始材料。如果未来制作 SFT 数据，我会从 LLM Client 的结构化消息生成样本，并显式记录哪些 Token 是 Agent 动作、哪些是环境观察。不能简单把完整运行日志拼接后全部算 loss，否则模型可能学习伪造工具结果。'
]),
('143. DPO 的原理是什么？它和 SFT、PPO 有什么区别？', [
'我的回答是：DPO 使用同一 Prompt 下的 chosen 与 rejected 回答对，直接提高模型相对参考模型生成 chosen 的对数概率差，并降低 rejected 的相对优势。它从带 KL 约束的偏好优化目标推导出一个分类式损失，不需要显式训练 Reward Model，也不需要像 PPO 那样在线采样并维护 Critic，因此训练流程更简单、稳定。',
'SFT 学习“示范答案是什么”，只有单个目标答案；DPO 学习“两个可接受程度不同的答案中更偏好哪个”。PPO 则通常用 Reward Model 给在线生成打分，通过策略梯度更新，可优化更灵活的序列级奖励，但工程复杂、容易出现奖励投机。DPO 的局限是依赖高质量偏好对和参考策略支持，离线数据覆盖不足时也难纠正模型没有探索过的行为。',
'RepoMind 没有使用 DPO。若积累“同一 Issue 的两份分析及人工偏好”，可以用 DPO优化证据引用、简洁度或拒绝编造，但偏好标签必须基于 commit 级证据，而不能只偏好文风更自信的答案。工具执行是否正确更适合规则评测或执行反馈，不能全部压缩成主观偏好。'
]),
('144. PPO 用于 RLHF 时的训练链路是什么？主要难点有哪些？', [
'我的回答是：典型 RLHF 先对基础模型做 SFT，再用人工偏好训练 Reward Model。PPO 阶段由当前策略对 Prompt 采样回答，Reward Model 给序列奖励，同时用参考模型计算 KL 惩罚，Critic 估计 value；随后通过 clipped surrogate objective 更新策略，并训练 value。KL 约束用于避免策略为了奖励偏离语言模型分布过远。',
'难点首先是成本：策略、参考模型、Reward Model 和 Critic 会同时占用大量显存，在线 rollout 也很慢。其次是稳定性，需要调 KL、clip range、advantage normalization、学习率和奖励尺度；奖励模型存在偏差时，策略会利用漏洞产生 reward hacking。长度偏好、格式奖励和训练数据泄漏都可能让奖励上升但真实质量下降，因此必须用独立人工评测与可执行指标复核。',
'RepoMind 没有 PPO/RLHF 训练链路，项目更适合先用确定性约束控制工具边界。若未来优化 Agent 策略，我会把“找到正确文件、调用参数合法、报告证据可验证”拆成可观测指标，并在离线沙箱中训练，绝不会让探索策略直接对真实 GitHub 仓库执行写操作。'
]),
('145. GRPO 与 PPO 有什么区别？为什么近期推理模型常讨论 GRPO？', [
'我的回答是：GRPO 通常对同一个 Prompt 采样一组回答，用组内奖励的相对位置构造 advantage，例如对奖励做组内均值和标准差归一化，而不额外训练一个 value/critic 模型。随后仍用重要性比率裁剪和 KL 约束更新策略。相比 PPO，它减少了 Critic 的训练和显存开销，也让相对可验证奖励更容易用于数学、代码等多解任务。',
'它并不是天然优于 PPO。组内多个 rollout 会增加生成成本；如果所有样本奖励相同或奖励噪声很大，组内 advantage 信号很弱。奖励设计仍可能被投机，组大小、采样温度、KL 系数和长度归一化都会影响结果。不同论文或实现对“GRPO”的细节也可能不同，面试时应先说明采用的具体目标，而不是把名称当成统一算法。',
'RepoMind 没有使用 GRPO。理论上可以针对 Issue 生成多条工具轨迹，用“是否定位正确文件、证据是否来自目标 commit、工具成本”构造组内奖励；但仓库分析的真实答案常不唯一，自动判分也不充分。因此在训练前要先建立沙箱、轨迹回放和人工复核，现阶段项目采用 Reflection 与有限重试更可控。'
]),
('146. 如何防止强化学习中的 Reward Hacking？', [
'我的回答是：Reward Hacking 是策略找到了提高代理奖励、却没有完成真实目标的捷径。例如报告写得更长而 Judge 给分更高，或大量调用工具碰巧覆盖参考关键词。根因不是模型“作弊”这么简单，而是奖励只观测了真实目标的一部分，且策略会主动寻找其中漏洞。',
'防护需要多层奖励与约束：把结果正确性、证据可验证性、工具成本和安全策略分开；能执行验证的指标优先用单元测试、文件位置和 Schema 校验，不只用 LLM Judge。保留隐藏测试集和对抗集，定期人工抽检高分轨迹；对长度、重复调用和越权行为设置硬限制。训练中监控奖励分项、KL、答案长度及分布漂移，而不是只看总奖励。',
'在 RepoMind 中，即使没有 RL，Reflection 也可能出现类似问题：模型可能偏爱措辞完整但证据薄弱的报告。因此反思节点要求引用文件和代码证据，重试次数由代码限制，评测还应比较根因与受影响文件。若未来做 RL，这些现有可观测信号可以成为奖励候选，但必须防止把“调用更多工具”误当成“分析更深入”。'
]),
('147. MoE 模型是什么？路由、负载均衡和推理有什么挑战？', [
'我的回答是：MoE 把 Transformer 某些 FFN 层替换为多个 Expert。Router 根据每个 Token 的隐藏状态选择 top-k Expert，只有被选中的 Expert 参与计算，因此模型总参数可以很大，而每个 Token 的激活计算量相对可控。Attention 通常仍是共享的；MoE 的优势主要是以稀疏激活扩大模型容量，而不是让每个专家显式对应某个固定业务领域。',
'训练时如果 Router 总选择少数 Expert，会造成热门 Expert 溢出、其他 Expert 学不到东西，所以常加入负载均衡辅助损失、容量因子或路由扰动。分布式训练还需要 all-to-all 把 Token 发送到对应 Expert，通信可能成为瓶颈。推理时虽然激活参数少，但完整权重仍需存储或跨节点分片；小 batch 下路由和通信开销也可能抵消理论 FLOPs 优势。',
'RepoMind 调用的模型是否采用 MoE 由模型供应方决定，项目没有训练或部署 MoE。选型时我会关心实际任务准确率、价格、延迟和工具调用稳定性，而不会把“总参数更多”直接等同于效果更好。自部署时还要结合硬件互联和并发测算，单卡放不下全部 Expert 时，部署复杂度会显著增加。'
]),
('148. 大模型训练中的数据并行、张量并行和流水线并行有什么区别？', [
'我的回答是：数据并行把完整模型复制到多个设备，每个设备处理不同 mini-batch，再同步梯度；适合模型能放进单卡但希望扩大吞吐。ZeRO/FSDP 会进一步把优化器状态、梯度和参数分片，降低每卡内存。张量并行则把单层矩阵乘法沿行或列切到多卡，解决一层权重放不下的问题，但每层都可能发生通信。',
'流水线并行按层把模型切成多个 stage，不同 micro-batch 像流水线一样通过各 stage，适合模型纵向分割，但存在 pipeline bubble 和调度复杂度。实际大规模训练通常组合 3D parallelism，并加上 sequence/context parallel。选择取决于模型大小、序列长度、batch、节点内 NVLink 与节点间网络；并行度越高不代表线性加速，通信和负载不均会吞噬收益。',
'RepoMind 没有训练基础模型，因此没有实现这些并行策略。项目中的“并行”是 LangGraph 任务并行：Issue 的多路上下文收集以及 PR 的 security、logic、style 审查。二者不能混淆；前者解决 GPU 上的模型容量和训练吞吐，后者解决应用工作流中独立 I/O 与模型调用的端到端延迟。'
]),
('149. 大模型训练为什么常用 BF16？FP16、BF16 和 FP32 如何取舍？', [
'我的回答是：FP32 有 8 位指数和 23 位尾数，范围和精度都高，但显存与计算成本大。FP16 只有 5 位指数、10 位尾数，精度尚可但动态范围小，训练中容易上溢或下溢，通常需要 loss scaling。BF16 保留与 FP32 相同的 8 位指数，只减少尾数到 7 位，因此动态范围大、训练稳定，现代 GPU/TPU 上常用于大模型训练。',
'BF16 不是所有计算都只保留 BF16：矩阵输入可以是 BF16，累加、优化器状态、部分归一化和 master weights 可能保留 FP32。推理还可以使用 FP16、INT8、INT4，但低比特会在显存、吞吐和精度之间权衡。选择应看硬件支持、Kernel、任务敏感度和端到端成本，而不是只比较理论位宽。',
'RepoMind 使用 API 模型，不控制供应商内部 dtype。若自部署用于代码分析，我会用同一评测集比较不同精度下的根因准确率、工具调用 JSON 合法率和长上下文稳定性，同时记录显存、P95 延迟和吞吐。代码中的标识符和数字对量化误差可能更敏感，因此不能只用通用困惑度判断。'
]),
('150. 如果要为 RepoMind 选择自部署模型，你会如何评估训练与推理方案？', [
'我的回答是：我会先明确目标，不会从“选哪个最大模型”开始。RepoMind 至少有意图分类、检索查询生成、ReAct 工具决策、结构化根因报告和 Reflection，不同节点对能力、延迟和成本要求不同。先用当前 API 模型建立仓库隔离的基线，指标包括根因正确率、相关文件命中、工具选择与参数正确率、结构化输出合法率、任务成功率、Token 成本和 P95 延迟。',
'推理侧会用真实长度分布做压测，比较模型在 BF16/量化下的 prefill、decode、KV Cache 占用和并发容量，并验证 vLLM 或其他引擎对 Tool Call、JSON Schema、长上下文和前缀缓存的支持。架构上可以让小模型处理分类和 Query Rewrite，大模型只处理复杂分析；路由失败时再升级模型。部署必须有超时、限流、灰度、回滚和隐私隔离。',
'训练侧遵循从轻到重：先改 Prompt、RAG 和工具 Schema；再根据 Badcase 判断是否需要 LoRA/SFT；只有积累可靠偏好或可验证奖励后才考虑 DPO/GRPO。RepoMind 当前没有自训练和自部署，这是演进方案而非现状。每次训练都必须固定索引版本和评测集，做离线回归与线上灰度，确保收益来自模型而不是测试泄漏或更长上下文。'
]),
]

B = RGBColor(46, 116, 181); D = RGBColor(31, 77, 120); G = RGBColor(95, 103, 112)
def sf(run, size=11, bold=None, color=None):
    name = 'Arial Unicode MS'; run.font.name = name; run.font.size = Pt(size)
    rf = run._element.get_or_add_rPr().get_or_add_rFonts()
    for a in ('ascii','hAnsi','eastAsia','cs'): rf.set(qn('w:'+a), name)
    if bold is not None: run.bold = bold
    if color is not None: run.font.color.rgb = color
def add_page_field(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sf(p.add_run('第八批B · 第 131—150 题  |  '), 9, color=G)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    p._p.append(fld)

doc = Document(); s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
s.top_margin = s.bottom_margin = Inches(.82); s.left_margin = s.right_margin = Inches(.88)
st = doc.styles['Normal']; st.font.name='Arial Unicode MS'; st.font.size=Pt(11)
st._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial Unicode MS')
st.paragraph_format.space_after=Pt(6); st.paragraph_format.line_spacing=1.25
for nm,z,bf,af,c in [('Heading 1',16,16,8,B),('Heading 2',13,12,6,D)]:
    st=doc.styles[nm]; st.font.name='Arial Unicode MS'; st.font.size=Pt(z); st.font.bold=True; st.font.color.rgb=c
    st._element.rPr.rFonts.set(qn('w:eastAsia'),'Arial Unicode MS')
    st.paragraph_format.space_before=Pt(bf); st.paragraph_format.space_after=Pt(af); st.paragraph_format.keep_with_next=True
sf(s.header.paragraphs[0].add_run('RepoMind · Agent 开发面试问答'),9,color=G)
add_page_field(s.footer.paragraphs[0])
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70); p.paragraph_format.space_after=Pt(10)
sf(p.add_run('RepoMind Agent 开发面试问答'),24,True,D)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(35)
sf(p.add_run('第八批 B：大模型原理、训练与推理（20 题）'),14,color=B)
doc.add_paragraph('本批为第 131—150 题，覆盖 Transformer、Attention、LoRA、SFT、偏好优化、MoE 与训练推理工程。涉及 RepoMind 未采用的训练或自部署技术时，回答会明确说明项目边界。')
doc.add_page_break()
for q, ps in QA:
    doc.add_heading(q, level=1)
    for t in ps:
        p=doc.add_paragraph(t); p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第八批B 131—150题'
doc.core_properties.author='RepoMind'; doc.core_properties.subject='大模型原理、训练与推理'
doc.save(OUT)
print(OUT)
print('questions',len(QA))
print('chars',sum(len(q)+sum(map(len,ps)) for q,ps in QA))
