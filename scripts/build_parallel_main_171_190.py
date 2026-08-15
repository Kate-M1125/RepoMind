from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches,Pt,RGBColor
OUT=Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第八批D_171-190.docx')
QA=[
('171. 什么是 Agentic RAG？它与一次检索后直接生成有什么区别？',[
'我的回答是：固定 RAG 通常执行一次 Query、召回、精排和生成，路径在请求前就确定；Agentic RAG 允许模型根据当前证据判断是否继续搜索、改写 Query、读取文件、追踪调用关系或查询外部依赖。区别不在是否使用向量库，而在检索过程是否形成可控的决策循环。',
'RepoMind 是两者的组合。retrieve.py 先执行确定性的多查询召回、RRF 和 CrossEncoder，为模型提供高质量初始候选；进入 analyze 节点后，ReAct Agent 可以继续调用 grep_code、read_file、find_definition、find_callers、Git 和依赖工具，直到证据足够或达到轮数上限。',
'这种组合避免两个极端：只做一次 RAG 可能漏掉跨文件根因，从空白开始完全 Agentic 搜索又会消耗太多轮次。固定检索负责召回效率，Agent 负责沿证据动态扩展，外层 Workflow 控制预算和停止。'
]),
('172. 什么情况下应该坚持固定 RAG，而不是升级成 Agentic RAG？',[
'我的回答是：当问题类型稳定、答案集中在单篇文档、一次召回命中率已经很高，并且业务强调整体延迟、成本和可预测性时，固定 RAG 更合适。例如 FAQ、规则查询和标准 API 文档问答，没有必要让模型自主循环。',
'Agentic RAG 适合需要多跳证据、跨文件调用链、工具验证和查询路径无法预先确定的任务。RepoMind 的 Issue 根因分析属于这种场景，因为模型可能先看到入口，再发现真正问题位于被调用函数、历史提交或依赖库。',
'是否升级应由 Badcase 数据决定。如果绝大多数失败来自单次检索漏证据，且补检能稳定改善，才值得增加 Agent Loop；如果错误来自文档质量或模型理解，增加循环只会放大成本。'
]),
('173. 如何防止 Agentic RAG 的多轮补检把延迟和成本拉得过高？',[
'我的回答是：需要同时限制轮数、候选规模、工具集合和总预算。每轮补检必须有新的检索意图，不能只是换一种表述重复搜；连续结果高度重合时应提前停止。',
'RepoMind 先用 RAG 缩小范围，再允许 ReAct 最多 6到12轮。不同节点只暴露必要工具，CrossEncoder 只处理 RRF 后的少量候选，最终上下文保留 top-3。Reflection 低分最多重试两次。',
'生产化还可记录 Query 指纹和结果 Jaccard 相似度，若多轮无新增文件或证据就终止；按任务设置 Token、工具费用和墙钟预算。预算用尽时返回当前结论、缺失证据和低置信度，而不是假装完整。'
]),
('174. RepoMind 的评测数据从哪里来？如何构建高质量测试集？',[
'我的回答是：代码分析评测需要可验证的真实案例。可以从已经关闭并有修复 PR 的 GitHub Issue 构建样本，以合并 diff、维护者评论和版本记录作为参考；PR Review 则使用已知问题、人工标注或历史 Review 评论。',
'RepoMind 的 eval/archive_v1 已保存多个开源仓库的 bug 结果，dataset、runner、judge 和 report 构成离线评测链路，tests/regression 还有 smoke cases 与 baselines。高质量样本要覆盖不同仓库、问题类型、代码规模和根因类别，并避免同一模式大量重复。',
'金标准不能简单复制 PR 全部改动，因为一个 PR 可能包含顺手重构。需要人工确认核心根因、关键文件、严重度和可接受修复，再将开发集、回归集与独立留出集分开。'
]),
('175. 如果把 RepoMind 流程封装成 Skill，Skill 路由怎么测试？',[
'我的回答是：Skill 路由测试要同时覆盖该触发和不该触发。正例包括输入 GitHub Issue/PR URL、要求根因分析或 Code Review；负例包括普通代码解释、非 GitHub 链接和仅要求修改本地文件。还要测试多个 Skill 都可能匹配时的冲突。',
'数据集应保存 user_query、expected_skill、required_inputs 和是否需要澄清。指标包括路由准确率、拒识率、误触发率和澄清正确率；仅看总体准确率会掩盖高风险误触发。',
'RepoMind 当前通过 main.py URL 规则选择 IssueGraph 或 PRGraph，还没有通用 Skill Router。若封装 Skill，可以保留 URL 正则作为确定性高精度路由，再由模型处理自然语言别名。'
]),
('176. 什么是 Harness Engineering？在 RepoMind 中体现在哪里？',[
'我的回答是：Harness Engineering 指围绕模型建立一套执行与约束环境，让模型能力可以稳定落地。它包括上下文构建、工具协议、状态机、权限、重试、结构化输出、Trace、评测和恢复，而不只是 Prompt。',
'RepoMind 的 LangGraph、TypedDict State、ToolRegistry、LLM Client、RAG、Reflection、Checkpoint 和 eval 共同构成 Harness。模型负责语义判断，但工具执行、条件路由、最大轮数和报告结构由系统控制。',
'项目目前的 Harness 仍有边界：Job 元数据未外置，权限治理和工具版本不足，Trace 主要是节点日志。说明这些不足并给出改进，比把 Harness 简化成“写了几个 Prompt”更符合工程实际。'
]),
('177. Agent 上下文压缩怎么做？如何避免压缩后丢失之前的流程？',[
'我的回答是：不能直接摘要整个 messages 后丢弃原文。应把信息分成不可压缩的结构化状态、可摘要的对话过程和可外置的证据。目标、用户约束、当前计划、已完成步骤和工具副作用必须保存在 State。',
'RepoMind 已将 issue_type、fix_type、affected_files、confidence 和 iteration 等关键内容显式存储；长代码保留文件路径和行号，必要时可重新 read_file。ReAct 历史可以压缩为“调用过哪些工具、得到哪些确认结论和未解决问题”，但 tool_call 与关键结果引用需要保留。',
'压缩后用一致性检查验证目标、待办和已完成步骤数量，并保存摘要版本和原 Trace 引用。这样即使 Prompt 中不再包含全部历史，恢复和审计仍可访问原始证据。'
]),
('178. To-Do List 为什么能让 Agent 更聚焦？RepoMind 是否实现了？',[
'我的回答是：To-Do List 把隐含计划外置为 pending、in_progress、done，使模型每轮知道当前目标和剩余任务，减少重复搜索和遗漏。它尤其适合长任务和多个独立子目标。',
'RepoMind 没有实现通用可编辑 To-Do List，但图节点本身承担了宏观计划：分类、上下文、分析、Reflection、报告都有明确阶段；ReAct 则在节点内增量决策。iteration 和 State 保存执行进度。',
'若加入 To-Do，应把列表放在结构化 State，工具完成后由代码校验再更新，而不是允许模型随意把未完成项标为 done。对于短的代码探索，过细待办也会增加 Token，应按任务复杂度启用。'
]),
('179. build_project_map 是怎么构建的？它对 Agent 有什么价值？',[
'我的回答是：project_map 提供仓库宏观导航，让模型在检索前理解主要目录、代码、测试、配置和文档分布。没有地图时，模型容易只根据文件名盲搜，也不知道召回文件在项目中的角色。',
'RepoMind 在 fetch 后执行 build_project_map，并把结果写入 State。后续入口识别、分析 Prompt 和 PR Review 都能使用这份结构；它与 RAG 的区别是，地图提供全局骨架，RAG 提供局部相关代码。',
'项目地图要控制深度和噪声，排除 .git、构建产物和大型依赖目录；大仓库可先展示顶层和关键模块，再按需 list_dir 展开，不能把完整文件树全部塞入上下文。'
]),
('180. detect_regression 是怎么判断回归问题的？为什么要单独规划？',[
'我的回答是：回归问题的证据需求与普通 bug 不同，需要关注“之前正常、某版本后失败”的时间线、版本号、相关提交和 blame。detect_regression 根据 Issue 文本中的版本与行为变化线索判断 is_regression。',
'如果是回归，Git 上下文分支会检索近期提交、文件历史或 blame，分析 Agent 也获得 Git 工具。这样模型不会只从静态代码猜根因，而会寻找引入行为变化的具体 diff。',
'单独规划能节省工具轮次：非回归问题不必无差别查询大量提交。判断可能不完全准确，因此 Git 证据应作为补充；当堆栈和静态代码明显指向近期改动时，Agent仍可以按需调用 Git。'
]),
('181. classify_fix 为什么在 analyze_bug 之前？',[
'我的回答是：classify_fix 先把可能修复分为 code_fix、dependency_upgrade 和 config_fix，目的是给分析阶段一个高层关注方向，而不是提前决定最终答案。',
'dependency_upgrade 时 Prompt 会关注 pyproject.toml、lockfile 和依赖源码；config_fix 关注配置路径和默认值；code_fix 则继续追踪本仓库实现。这样 ReAct 不需要浪费前几轮反复判断问题属于哪种修复。',
'分类结果必须允许被证据推翻。它是规划先验，不是硬规则；最终 root_cause 和 fix_suggestion 仍由真实代码和工具结果决定。'
]),
('182. Feature 分支为什么先执行 detect_existing？',[
'我的回答是：用户请求的功能可能已经存在，只是 API 名称、版本或使用方式不明确。如果直接规划新代码，Agent可能重复实现、破坏兼容性或给出错误文件。',
'detect_existing 会从 Feature 描述中提取可验证的搜索词和路径线索，在仓库中检查现有 API。结果写入 feature_exists 和 existing_api，analyze_feature 再决定回答使用方法、补文档，还是设计真正的新功能。',
'搜索词不能只取对象名。例如“给 MemoryUsage 补测试”应搜索 test_memusage 和 tests 路径，而不是只搜索 MemoryUsage。这个节点的价值是把“是否已有能力”变成可验证证据。'
]),
('183. Question 分支与 Bug 分支有什么不同？',[
'我的回答是：Question 的目标不是寻找待修复缺陷，而是解释当前框架为什么这样行为、正确用法是什么以及有哪些限制。它更重视文档、现有 API 和设计语义，不能强行给出代码修改。',
'RepoMind 的 question 直接进入 answer_question，不执行 regression 和 classify_fix。它仍使用项目地图、Memory、RAG 和代码工具，因为高质量解释必须基于真实实现，而不是模型常识。',
'如果分析中发现确实存在实现错误，可以在答案中说明证据并降低对原分类的信任，但当前 Graph 不支持动态改 issue_type，这是可改进边界。'
]),
('184. Security 分支如何处理？为什么不能完全自动化？',[
'我的回答是：Security Issue 可能包含未公开漏洞、利用方式和敏感影响，错误传播或自动执行都可能造成风险。因此 RepoMind 在 route 后先进入 human_gate，确认后才继续 Memory 和分析。',
'批准后的分析复用 bug 路径：收集 entry/RAG/Git 证据、classify_fix、analyze_bug 和 Reflection。报告仍应控制敏感细节，不把可直接利用的信息写入公共 Memory。',
'当前 human_gate 较基础。生产化应验证审批者身份、记录决定、限制 Memory 作用域，并对高危工具单独授权。HITL 是风险控制的一层，不替代权限和审计。'
]),
('185. memory_retrieve 放在上下文准备之前有什么考虑？',[
'我的回答是：先召回历史案例可以为后续检索提供仓库特有术语、类似根因和已验证文件，帮助 Query 与分析聚焦；但它不能覆盖当前代码证据。',
'RepoMind 在分类后执行 memory_retrieve，再 prepare_context 和三路证据收集。Memory 作为单独字段进入 Prompt，与 code_context 分开，模型可以区分历史经验和当前证据。',
'如果历史案例与当前 commit 冲突，应优先当前代码。更成熟的实现还应把 Memory 的版本和置信度用于过滤，避免先验让检索产生确认偏差。'
]),
('186. Reflection 重试为什么回分析节点，而不是从头执行？',[
'我的回答是：Reflection 质疑的通常是结论证据不足，而不是 GitHub 数据和仓库索引失效。从头重跑 fetch、project_map 和 Memory只会增加成本，且输入没有变化。',
'Issue 根据 issue_type 回 analyze_bug、analyze_feature 或 answer_question；PR 回 prepare_reviews 重新并行 Review。已有 code_context 和 reflection_note 可以用于第二轮补证据。',
'如果 critic 指出的是检索遗漏，理想实现应允许回 prepare_context 或指定工具，而当前 Issue Graph 主要回分析节点。这是可以进一步细化的错误路由。'
]),
('187. RepoMind 如何保证最终报告结构稳定？',[
'我的回答是：分析阶段先生成结构化字段，报告节点再消费 root_cause、affected_files、severity、fix_suggestion、code_changes、confidence 等字段，而不是直接把整个 ReAct 自由文本原样返回。',
'LLM 输出经过 JSON 解析和 Pydantic/字段校验，confidence 做范围限制。PR 的 Verdict 由严重度规则确定，critical/high 才 REQUEST CHANGES，避免模型随意改变合并策略。',
'结构稳定还需要语义校验，例如 affected_files 必须存在、行号合法、严重度与描述一致。当前项目已有基本 Guardrail，生产化可增加引用验证和 Schema 版本。'
]),
('188. RepoAgent Server 的作用是什么？',[
'我的回答是：RepoAgent Server 把单仓库分析能力封装成服务，使跨仓库 Orchestrator 可以通过统一 Job 接口提交 Issue 并获取报告。每个服务只负责自己仓库范围，降低跨仓库上下文混杂。',
'agents/repo_agent_server.py 与主 backend 类似，为任务生成 job_id/thread_id，执行 IssueGraph 并返回 confidence 和报告。Orchestrator 可以并行调用多个 Repo Agent 后合并。',
'当前服务仍偏原型，任务状态、认证、超时和服务发现需要生产化。跨 Agent 通信应传结构化结果和证据引用，而不是完整自然语言历史。'
]),
('189. RepoMind 为什么没有直接做流式 Token 输出？应该怎么扩展？',[
'我的回答是：RepoMind 当前是异步 Job 模式，主要返回节点完成后的结构化报告，而不是将模型每个 Token 直接推给前端。代码分析中间包含工具调用和多分支，原始流式文本不等于可靠进度。',
'更有价值的流式设计是发送结构化事件：job_started、node_started、tool_called、evidence_found、reflection_retry、report_chunk 和 completed。前端可用 SSE 展示进度，最终结果仍以持久化 State 为准。',
'不能把模型思维链原样流出，也不能在工具尚未验证时把临时结论显示成事实。事件需要序号和重连游标，服务重启后可以从持久化日志继续。'
]),
('190. 如果把 RepoMind 推向生产，优先改造哪些部分？',[
'我的回答是：第一优先级是可靠性与安全：把 jobs 外置到数据库，加入 MQ/Worker、任务租约、取消与恢复；为私有仓库和工具加入权限、密钥管理、审计和敏感数据过滤。',
'第二是检索与语言覆盖：索引绑定 commit，增加增量更新与版本切换，接入 BM25/混合检索，为多语言使用 Tree-sitter或 LSP，并完善父子上下文和依赖版本。',
'第三是质量闭环：统一 Trace，扩大人工金标准和线上采样，组合规则、Judge 与人工评测，治理 Memory 版本和污染。最后再优化模型路由、批处理与成本，不能在任务状态和权限仍不可靠时只追求吞吐。'
]),
]
B=RGBColor(46,116,181);D=RGBColor(31,77,120);G=RGBColor(95,103,112)
def sf(r,z=11,b=None,c=None):
 n='PingFang SC';r.font.name=n;r.font.size=Pt(z);rf=r._element.get_or_add_rPr().get_or_add_rFonts()
 for a in('ascii','hAnsi','eastAsia','cs'):rf.set(qn('w:'+a),n)
 if b is not None:r.bold=b
 if c is not None:r.font.color.rgb=c
doc=Document();s=doc.sections[0];s.page_width,s.page_height=Inches(8.5),Inches(11);s.top_margin=s.bottom_margin=Inches(.82);s.left_margin=s.right_margin=Inches(.88)
st=doc.styles['Normal'];st.font.name='PingFang SC';st.font.size=Pt(11);st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_after=Pt(6);st.paragraph_format.line_spacing=1.25
for nm,z,bf,af,c in [('Heading 1',16,16,8,B),('Heading 2',13,12,6,D)]:
 st=doc.styles[nm];st.font.name='PingFang SC';st.font.size=Pt(z);st.font.bold=True;st.font.color.rgb=c;st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_before=Pt(bf);st.paragraph_format.space_after=Pt(af);st.paragraph_format.keep_with_next=True
sf(s.header.paragraphs[0].add_run('RepoMind · Agent 开发面试问答'),9,c=G);f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;sf(f.add_run('并行批次 D · 第 171—190 题'),9,c=G)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(70);p.paragraph_format.space_after=Pt(10);sf(p.add_run('RepoMind Agent 开发面试问答'),24,True,D)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(35);sf(p.add_run('第八批 D：RepoMind 进阶项目题（20 题）'),14,c=B)
doc.add_paragraph('本批为第 171—190 题，逐题依据 RepoMind 当前代码、设计文档和真实边界整理。');doc.add_page_break()
for q,ps in QA:
 doc.add_heading(q,level=1)
 for t in ps:p=doc.add_paragraph(t);p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第八批D';doc.core_properties.author='RepoMind';doc.save(OUT);print(OUT)
