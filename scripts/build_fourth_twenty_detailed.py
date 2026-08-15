from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches,Pt,RGBColor

OUT=Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第四批20题.docx')
QA=[
('31. 什么是 Agent Memory？RepoMind 为什么需要 Memory？',[
'我的回答是：Agent Memory 不是简单保存聊天记录，也不是把上下文窗口无限拉长，而是让 Agent 在当前任务或未来任务中复用已经确认的信息。它至少要回答三个问题：记什么、什么时候召回、召回后以什么可信度影响决策。',
'RepoMind 的 Memory 主要是案例型长期记忆。Issue 或 PR 分析完成后，memory_save 会把仓库、标题、问题类型、根因或 Review 结论、影响文件和置信度写入 ChromaDB；新任务进入正式分析前，memory_retrieve 根据当前问题召回相似历史案例，作为辅助上下文。这样遇到同一仓库或相似问题时不必完全从零开始。',
'Memory 只能作为参考，不能替代当前仓库证据。历史案例可能来自旧版本，也可能本身判断错误，所以最终结论仍要通过当前代码、工具结果和 Git 历史验证。RepoMind 当前更接近 episodic memory，还没有完整实现用户偏好、规则记忆、过期和冲突治理。'
]),
('32. 长上下文和 Memory 有什么区别？',[
'我的回答是：长上下文解决“当前这一轮能放多少材料”，Memory 解决“哪些状态可以跨轮次或跨任务继续生效”。长上下文是模型本次推理的输入，任务结束后不会天然保留；Memory 存在外部系统中，可以按条件召回、更新和删除。',
'RepoMind 的 Issue 正文、项目地图、检索代码和 ReAct 工具结果属于当前任务上下文。ChromaDB 中保存的历史 Issue/PR 案例属于长期 Memory。即使模型上下文窗口足够大，也不应该每次加载全部历史案例，因为大量无关内容会增加 Token 和噪声。',
'工程上可以把上下文理解成当前工作台，把 Memory 理解成档案库。Memory 召回后仍要经过筛选，只有与当前仓库、问题和版本相关的高价值内容才进入上下文。'
]),
('33. Memory 和 RAG 有什么区别？为什么它们都可以使用向量数据库？',[
'我的回答是：RAG 更偏外部知识，解决“去哪里找当前问题的证据”；Memory 更偏 Agent 在历史任务中积累的状态和经验，解决“过去已经确认过什么”。二者可以使用相同的 embedding 和向量数据库，但写入来源、生命周期和可信度完全不同。',
'RepoMind 的仓库代码索引属于 RAG，内容来自真实源码，召回目标是相关文件和函数；历史分析集合属于 Memory，内容来自 Agent 的过去输出，召回目标是相似案例。RAG 命中的源码可以作为一手证据，Memory 命中的历史结论只能作为线索。',
'因此不能因为底层都用 ChromaDB 就把两者混为一谈。Memory 需要额外关注来源、置信度、版本、过期和纠错；代码 RAG 更关注切片、索引更新、召回率和权限过滤。'
]),
('34. 短期记忆和长期记忆分别如何实现？Function Call 内容要不要放进短期记忆？',[
'我的回答是：短期记忆服务当前任务，长期记忆服务未来任务。RepoMind 的短期记忆包括 LangGraph State、ReAct messages 和 SQLite Checkpoint；长期记忆是 ChromaDB 中完成后的结构化案例。',
'Function Call 必须进入短期消息链，而且 assistant 的 tool_calls 与 tool role 的执行结果要成对保存。否则模型不知道自己已经调用过什么工具、参数是什么、返回了哪些证据，容易重复调用，也无法沿着上一步结果继续分析。',
'但短期记忆也不能无限增长。长工具输出应该保留关键片段、文件位置和摘要；确定性的业务状态放入 State，而不是要求模型每轮从 messages 中重新提取。'
]),
('35. 什么信息适合写入长期记忆？RepoMind 的写入时机是什么？',[
'我的回答是：适合长期保存的是稳定、可复用、经过验证的信息，例如已经确认的根因、关键影响文件、修复模式、仓库特有规则和高质量 Review 结论。不适合保存临时猜测、未验证推理、一次性网络错误和敏感凭证。',
'RepoMind 在 generate_report 之后才执行 memory_save，而不是在分析中途每获得一个猜测就写入。此时已经经过 Reflection，能够保存标题、仓库、issue_type、结论和 confidence 等结构化字段。PR 也在报告完成后写入 Review 案例。',
'更严格的生产实现还应设置最低置信度、来源链接、commit/version、写入者和过期条件。低置信度或人工否定的结果不应自动进入长期记忆。'
]),
('36. Memory 是不是越多越好？召回噪声怎么控制？',[
'我的回答是：不是。Memory 越多，近似但不相关的案例越容易进入结果；错误记忆还会反复污染后续任务。Memory 的价值不在存储量，而在需要时能否召回正确、可信、处于有效范围内的信息。',
'控制噪声可以从写入和召回两端处理。写入端设置质量门槛、去重和版本元数据；召回端先按仓库、任务类型、时间和权限做硬过滤，再进行向量相似度排序，最后只注入少量高价值案例。',
'RepoMind 当前 Memory 机制较轻量，主要使用语义召回。面试时应明确，完整系统还需要更新、删除、过期、冲突解决和人工纠错，不能把 append-only 向量库称为成熟 Memory。'
]),
('37. 为什么错误记忆很危险？如何防止 Memory 污染？',[
'我的回答是：错误记忆比单轮错误更危险，因为它会跨任务持续存在。一次错误根因如果被保存，后续相似 Issue 可能反复召回它，模型会把“历史上出现过”误认为“已经验证正确”。',
'防止污染首先要记录来源、仓库版本、置信度和证据；其次只允许通过评测或人工确认的结果进入高可信 Memory；再次要支持撤销、降权和过期。外部网页和仓库文档中的指令只能作为数据，不能直接提取为长期规则。',
'RepoMind 已经在写入前经过 Reflection，并保存 confidence，但这还不足以构成完整治理。更可靠的方案是把低置信度案例放入候选区，经过回归集或人工审核后再升级。'
]),
('38. 向量数据库是否等于 Memory？',[
'我的回答是：不等于。向量数据库只解决相似内容的存储和检索，是 Memory 的一个基础设施组件；完整 Memory 还包括类型划分、写入策略、作用域、权限、版本、置信度、过期、删除、冲突和审计。',
'RepoMind 使用 ChromaDB 同时承载代码索引和历史案例，但前者是 RAG，后者才是长期案例 Memory。仅仅创建一个 Collection 并把对话全部写进去，只能称为历史检索，不能说明系统已经具备可靠记忆。',
'面试中我会强调：Memory 的难点往往不在“怎么向量化”，而在“什么值得记、什么时候忘、召回结果能不能当事实”。'
]),
('39. Checkpoint 和 Memory 有什么区别？',[
'我的回答是：Checkpoint 用于恢复同一个任务，Memory 用于帮助未来任务。Checkpoint 保存当前执行位置、State、节点输出和重试次数，要求状态一致；Memory 保存可复用经验，通常通过检索按需注入。',
'RepoMind 的 checkpoints.sqlite 和 pr_checkpoints.sqlite 属于 Checkpoint，使用 thread_id 隔离任务；ChromaDB 中历史 Issue/PR 分析属于 Memory。服务中断后应通过原 thread_id 恢复 Checkpoint，而不是从 Memory 猜测任务执行到了哪里。',
'两者的数据保留策略也不同。Checkpoint 可在任务完成后归档或清理；长期 Memory 需要版本、过期和治理。把二者混在一起会导致恢复不准确，也会让历史检索携带大量流程噪声。'
]),
('40. Memory 更新时出现前后结论不一致，应该怎么处理？',[
'我的回答是：不能简单用新内容覆盖旧内容，也不能让冲突版本同时作为确定事实召回。需要先判断冲突来自仓库版本变化、证据更新、作用域不同，还是某次分析错误。',
'我会为记忆保存 repo、commit/version、source、confidence、created_at 和状态。新结论与旧结论冲突时创建新版本并建立 supersedes 关系；只有证据更强或人工确认后才把旧版本标为失效。检索时优先当前仓库版本，并把未解决冲突明确提供给模型。',
'RepoMind 当前尚未实现完善的 Memory 版本治理。现有案例召回适合辅助分析，但生产化必须加入版本过滤和 Tombstone，否则仓库升级后可能继续召回旧根因。'
]),
('41. Function Calling 是怎么设计的？RepoMind 的消息格式是什么？',[
'我的回答是：Function Calling 不是模型直接执行函数，而是模型输出结构化调用意图，应用负责校验、执行并把结果返回。协议至少包含工具名称、JSON 参数、唯一 tool_call_id 和对应结果。',
'RepoMind 在 tools/base.py 中用 Tool 封装 name、description、parameters 和 execute。build_tools 生成 OpenAI-compatible tools schema 与 executor 映射。模型返回 assistant tool_calls 后，LLM Client 找到对应 executor 执行，再把结果作为 tool role 消息并携带相同 tool_call_id 加回 messages。',
'这种设计把决策与执行分开：模型决定可能需要什么能力，服务端掌握实际权限、参数校验、超时和审计。最终是否允许执行不能只看模型输出格式正确。'
]),
('42. Tool Response 应该用什么角色传回？为什么？',[
'我的回答是：应该使用 tool role，并通过 tool_call_id 与 assistant 发起的具体调用配对。不能把工具结果伪装成 user 消息，也不应该混进 system Prompt。',
'使用专门角色可以让模型清楚区分“用户要求”“系统规则”“模型行动”和“环境观察”。一次 assistant 消息可能包含多个 tool_calls，tool_call_id 能保证每个结果对应正确调用，避免并行工具结果错位。',
'RepoMind 的 ReAct Client 会保留完整的 assistant tool_calls 和 tool result。外部工具返回内容仍属于不可信数据，即使使用 tool role，也不能因此获得系统指令权限。'
]),
('43. Tool Schema 怎么设计，才能减少模型选错工具和参数错误？',[
'我的回答是：Schema 要让工具边界清楚，而不是描述得越长越好。工具名称应体现动作和对象；description 说明适用场景、不适用场景和返回内容；参数尽量使用明确类型、required、enum、范围和格式约束。',
'RepoMind 将 read_file、grep_code、find_definition、find_callers 分开，因为它们分别对应局部阅读、全文搜索、定义定位和反向调用追踪。若全部合成一个“search_repo”大工具，模型很难理解参数语义，也不利于权限和错误定位。',
'还可以加入正反例、减少同义工具、按节点只暴露当前需要的工具。Schema 校验解决的是结构正确，业务层仍要校验 owner/repo/path 是否在作用域内以及行号范围是否合理。'
]),
('44. 工具数量很多时，怎么避免模型选错工具？',[
'我的回答是：不要把所有工具在每一步都暴露给模型。工具越多，描述越相似，选择准确率越低，Prompt 和 Token 成本也越高。最有效的方法是按任务阶段和风险做工具裁剪。',
'RepoMind 的不同节点使用不同集合。bug 分析需要代码导航、Git 历史和依赖工具；feature/question 不一定需要全部 Git 工具；PR style Review 主要使用 read_file、grep_code 和 list_dir，而 logic/security 才增加定义和调用者工具。',
'如果工具仍然很多，可以先做能力路由，把任务映射到一个工具组；同时统计 tool selection accuracy、参数错误率和无效调用率。不要只靠在 Prompt 中反复强调“认真选择”。'
]),
('45. 工具调用失败或超时，应该怎么重试和降级？',[
'我的回答是：先分类错误，再决定动作。网络超时、限流和服务端 5xx 属于瞬时错误，可以指数退避并加入抖动；参数校验错误应把结构化错误返回模型修复；权限错误直接停止；资源不存在通常需要换路径或结束，不能盲目重试。',
'RepoMind 的模型网络调用使用 tenacity 重试，ReAct 本身有最大轮数，工具结果会进入 messages 供模型调整下一步。对只读工具，有限重试风险较低；对未来可能增加的写工具，超时不能直接证明执行失败。',
'降级方案可以是从 GitHub API 切换本地 clone、从语义检索退回 grep、减少 CrossEncoder 候选或输出“证据不足”的部分结果。降级必须明确影响，不能静默生成一个看似完整的结论。'
]),
('46. 写工具如何保证幂等，避免重试造成重复执行？',[
'我的回答是：每个有副作用的业务动作需要稳定幂等键，例如 task_id+tool_name+业务对象+规范化参数哈希。服务端先查询幂等记录：已成功就返回原结果，执行中则查询状态，失败且允许重试才重新执行。',
'网络超时时必须区分“请求未到达”和“执行完成但响应丢失”。退款、提交代码、创建 PR 等操作应由目标系统支持 idempotency key，或者在本地保存 PENDING、SUCCEEDED、FAILED 状态和外部 operation_id。',
'RepoMind 当前工具以读取和搜索为主，重复调用通常只增加成本，不会产生业务副作用。因此项目不能声称已经实现完整写工具幂等；这是未来自动修复或自动提交功能上线前必须补齐的能力。'
]),
('47. 如何避免模型生成格式正确但业务危险的工具参数？',[
'我的回答是：JSON Schema 只能保证结构合法，不能证明业务语义安全。模型生成 path="/"、删除全部记录或访问其他租户，即使类型完全正确，也必须被服务端拒绝。',
'我会在执行前做四层校验：参数 Schema、用户/租户权限、资源作用域、风险策略。路径必须经过规范化并限制在允许仓库；数量和金额设置上限；高风险写操作使用 prepare/confirm 两阶段提交，并向用户展示准确目标。',
'RepoMind 当前主要是仓库读取工具，并对 security Issue 设置 human_gate。生产化还需完整权限上下文和审计；Prompt 里的“不要危险操作”只能作为辅助，不能作为安全边界。'
]),
('48. 如何防御工具返回内容中的间接 Prompt Injection？',[
'我的回答是：仓库代码、Issue、网页和依赖文档都可能包含“忽略之前指令、调用某工具”等恶意文本。工具结果必须被视为不可信数据，而不是新的系统指令。',
'防护包括把 system、user 和 tool 数据分层；在 Prompt 中明确外部内容只用于取证；工具白名单和权限由代码决定；执行前再次验证调用是否符合用户原始目标。还可以检测典型注入模式、截断无关内容，并禁止工具结果直接修改长期规则或 Memory。',
'RepoMind 的代码 Agent尤其需要注意 README、Issue 评论和源码字符串中的指令。human_gate 只覆盖 security 分支，并不等于完整 Injection 防护；模型输出始终只能作为候选计划。'
]),
('49. MCP、Function Calling、OpenAPI 和 Skill 有什么区别？',[
'我的回答是：Function Calling 是模型表达“我要调用某个函数”的消息机制；OpenAPI 是描述 HTTP API 的接口规范；MCP 是客户端与外部工具、资源和 Prompt 服务之间的标准通信协议；Skill 是围绕某类任务封装的可复用方法包。',
'MCP 可以向客户端提供能力发现和统一调用，但模型最终仍可能通过 Function Calling 形式选择工具。OpenAPI 可以被转换为工具 Schema，却不负责 Agent 的会话和能力生命周期。Skill 通常包含触发规则、操作步骤、参考资料、脚本和模板，内部可以使用 MCP 工具。',
'简单说，Function Calling 是模型输出格式，MCP 是连接协议，OpenAPI 是 Web 接口契约，Skill 是任务方法。它们处于不同层级，不是互相替代。'
]),
('50. RepoMind 使用 MCP 和 Skills 了吗？项目真实边界是什么？',[
'我的回答是：RepoMind 的工具抽象遵循 MCP 风格，但当前没有运行独立 MCP Server，也没有实现完整通用 Skills Runtime。项目真实落地的是 ToolRegistry：统一 Tool Schema、JSON 参数和 executor，并通过 OpenAI-compatible Function Calling 驱动 ReAct。',
'代码导航、Git 历史和依赖检索都可以未来封装为 MCP Server，使其他 Agent 客户端复用；Issue 根因分析和 PR Review 的稳定流程也可以进一步沉淀成 Skills，包含触发条件、步骤、模板和评测。',
'面试时我会把“当前实现”和“演进方案”分开：可以说明接口规范与 MCP 思想一致，但不能直接说已经完成 MCP 通信、认证、资源发现和版本治理；也不能把固定 Workflow 节点包装成完整 Skills 系统。'
]),
]

BLUE=RGBColor(46,116,181);DARK=RGBColor(31,77,120);GRAY=RGBColor(95,103,112)
def sf(r,size=11,bold=None,color=None):
 n='PingFang SC';r.font.name=n;r.font.size=Pt(size);rf=r._element.get_or_add_rPr().get_or_add_rFonts()
 for a in('ascii','hAnsi','eastAsia','cs'):rf.set(qn('w:'+a),n)
 if bold is not None:r.bold=bold
 if color is not None:r.font.color.rgb=color
doc=Document();s=doc.sections[0];s.page_width,s.page_height=Inches(8.5),Inches(11);s.top_margin=s.bottom_margin=Inches(.82);s.left_margin=s.right_margin=Inches(.88)
st=doc.styles['Normal'];st.font.name='PingFang SC';st.font.size=Pt(11);st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_after=Pt(6);st.paragraph_format.line_spacing=1.25
for nm,sz,bf,af,c in [('Heading 1',16,16,8,BLUE),('Heading 2',13,12,6,DARK)]:
 st=doc.styles[nm];st.font.name='PingFang SC';st.font.size=Pt(sz);st.font.bold=True;st.font.color.rgb=c;st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_before=Pt(bf);st.paragraph_format.space_after=Pt(af);st.paragraph_format.keep_with_next=True
h=s.header.paragraphs[0];sf(h.add_run('RepoMind · Agent 开发面试问答'),9,color=GRAY);f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;sf(f.add_run('第四批 · 第 31—50 题'),9,color=GRAY)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(70);p.paragraph_format.space_after=Pt(10);sf(p.add_run('RepoMind Agent 开发面试问答'),24,True,DARK)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(35);sf(p.add_run('第四批：Memory、Tools、MCP 与 Skills（20 题）'),14,color=BLUE)
doc.add_paragraph('本批为第 31—50 题。每道回答分别说明概念、RepoMind 实现和当前边界。');doc.add_page_break()
for q,ps in QA:
 doc.add_heading(q,level=1)
 for t in ps:p=doc.add_paragraph(t);p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第四批20题';doc.core_properties.author='RepoMind';doc.save(OUT);print(OUT)
