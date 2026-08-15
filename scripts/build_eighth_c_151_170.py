from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第八批C_151-170.docx')

QA = [
('151. 在有序数组中查找目标值第一次出现的位置。', [
'我的回答是：这道题的关键不是“找到 target 就返回”，而是把问题转化成寻找第一个大于等于 target 的位置，也就是 lower_bound。普通二分命中后立即返回，得到的可能是重复元素中任意一个位置；lower_bound 在 nums[mid] >= target 时仍然收缩右边界，才能继续向左寻找。',
'''实现时使用左闭右开区间 [left, right)，初始 right = len(nums)。循环结束时 left 是第一个 >= target 的下标，再检查它是否越界且值是否等于 target：

def first_position(nums, target):
    left, right = 0, len(nums)
    while left < right:
        mid = left + (right - left) // 2
        if nums[mid] >= target:
            right = mid
        else:
            left = mid + 1
    return left if left < len(nums) and nums[left] == target else -1''',
'时间复杂度是 O(log n)，额外空间 O(1)。边界包括空数组、目标小于所有元素、目标大于所有元素、数组全部等于目标。若问最后一次出现，可以先找第一个大于 target 的位置再减一。'
]),
('152. 股票买卖：只能交易一次与可以交易多次分别怎么做？', [
'我的回答是：只能交易一次时，需要维护扫描到当前位置之前的最低价格，并尝试把今天作为卖出日，更新最大利润。因为买入必须发生在卖出之前，不能简单用全局最大值减全局最小值。',
'''def max_profit_once(prices):
    min_price = float('inf')
    ans = 0
    for price in prices:
        min_price = min(min_price, price)
        ans = max(ans, price - min_price)
    return ans

def max_profit_many(prices):
    return sum(max(0, prices[i] - prices[i-1])
               for i in range(1, len(prices)))''',
'允许多次交易且不能同时持有多只股票时，可以把每一段正收益都累加。它等价于在每个上升段谷底买入、峰顶卖出，但逐日累加更简洁。两种算法都是 O(n) 时间；一次交易 O(1) 空间，多次交易同样 O(1)。空数组、单日价格和单调下降序列都返回 0。若存在手续费、冷冻期或最多 k 次交易，就不能直接套贪心，需要状态机动态规划。'
]),
('153. 求无重复字符的最长子串。', [
'我的回答是：使用滑动窗口。右指针逐个读字符，用 last 保存该字符最近一次出现的位置；如果它落在当前窗口内，就把左边界直接跳到上次位置的下一位。左边界只能向右移动，所以必须写成 max(left, last[ch] + 1)，否则会被窗口外的旧记录拉回。',
'''def length_of_longest_substring(s):
    last = {}
    left = ans = 0
    for right, ch in enumerate(s):
        if ch in last:
            left = max(left, last[ch] + 1)
        last[ch] = right
        ans = max(ans, right - left + 1)
    return ans''',
'每个字符只被右指针访问一次，时间复杂度 O(n)，哈希表空间 O(min(n, 字符集大小))。例如 abba 扫到第二个 a 时，a 的旧位置已经在当前窗口外，left 不能从 2 退回 1。若题目要求返回子串，还需在刷新最大长度时记录 best_left。Python 按 Unicode 字符迭代；若语言按字节处理 UTF-8，需要明确“字符”的定义。'
]),
('154. 求两个字符串的最长公共子串。', [
'我的回答是：最长公共子串要求连续，与最长公共子序列不同。定义 dp[i][j] 为 s1 以 i-1 结尾、s2 以 j-1 结尾的公共子串长度：字符相等时 dp[i][j] = dp[i-1][j-1] + 1，否则必须清零，因为连续性已经断开。',
'''def longest_common_substring(a, b):
    dp = [0] * (len(b) + 1)
    best = end = 0
    for i in range(1, len(a) + 1):
        for j in range(len(b), 0, -1):
            if a[i-1] == b[j-1]:
                dp[j] = dp[j-1] + 1
                if dp[j] > best:
                    best, end = dp[j], i
            else:
                dp[j] = 0
    return a[end-best:end]''',
'时间复杂度 O(mn)，滚动数组空间 O(n)。一维数组必须让 j 从右向左遍历，否则 dp[j-1] 已被当前行覆盖。若只要长度，不需要 end；若要返回所有并列结果，需要记录多个结束位置。超长字符串可讨论后缀自动机或后缀数组，但面试中应先把 DP 的定义和连续性讲清楚。'
]),
('155. 不使用内置大整数，如何实现大数加法和大数减法？', [
'我的回答是：把数字当字符串，从最低位向最高位逐位计算。加法维护 carry；减法维护 borrow。减法前先比较绝对值大小，决定结果符号，并保证用较大的绝对值减较小的绝对值。这样不会依赖语言整数上限。',
'''def add_str(a, b):
    i, j, carry, out = len(a)-1, len(b)-1, 0, []
    while i >= 0 or j >= 0 or carry:
        x = ord(a[i])-48 if i >= 0 else 0
        y = ord(b[j])-48 if j >= 0 else 0
        carry, digit = divmod(x + y + carry, 10)
        out.append(str(digit)); i -= 1; j -= 1
    return ''.join(reversed(out))

def sub_abs(a, b):  # 前提：a 的绝对值 >= b
    i, j, borrow, out = len(a)-1, len(b)-1, 0, []
    while i >= 0:
        x = ord(a[i])-48 - borrow
        y = ord(b[j])-48 if j >= 0 else 0
        if x < y: x, borrow = x + 10, 1
        else: borrow = 0
        out.append(str(x-y)); i -= 1; j -= 1
    while len(out) > 1 and out[-1] == '0': out.pop()
    return ''.join(reversed(out))''',
'加减法时间复杂度 O(max(m,n))，结果空间同阶。工程实现还要先规范化前导零、校验非法字符，并统一处理 -0。若要求完整带符号运算，可先解析符号：同号做绝对值加法，异号转成绝对值减法。'
]),
('156. 合并两个有序数组。', [
'我的回答是：如果 nums1 尾部预留了 n 个位置，应从后向前合并。正向写会覆盖 nums1 里尚未比较的有效元素；从尾部写时，每次放入两个数组当前最大值，目标位置天然是空闲的。',
'''def merge(nums1, m, nums2, n):
    i, j, k = m - 1, n - 1, m + n - 1
    while i >= 0 and j >= 0:
        if nums1[i] > nums2[j]:
            nums1[k] = nums1[i]; i -= 1
        else:
            nums1[k] = nums2[j]; j -= 1
        k -= 1
    while j >= 0:
        nums1[k] = nums2[j]
        j -= 1; k -= 1''',
'时间复杂度 O(m+n)，额外空间 O(1)。只需要补 nums2 的剩余元素：nums1 若还有剩余，本来就在正确位置。边界要覆盖 m=0、n=0、两数组值完全交错以及大量重复值。如果题目不允许修改输入，则用新数组和双指针，空间为 O(m+n)。'
]),
('157. 三数之和为 0，如何去重？', [
'我的回答是：先排序，再枚举第一个数 i，对剩余区间使用左右双指针。排序同时解决了移动方向和去重问题：和小于 0 时左指针右移，和大于 0 时右指针左移；命中后跳过左右两侧所有相同值。',
'''def three_sum(nums):
    nums.sort(); ans = []
    for i in range(len(nums)-2):
        if i > 0 and nums[i] == nums[i-1]: continue
        if nums[i] > 0: break
        left, right = i+1, len(nums)-1
        while left < right:
            total = nums[i] + nums[left] + nums[right]
            if total < 0: left += 1
            elif total > 0: right -= 1
            else:
                ans.append([nums[i], nums[left], nums[right]])
                lv, rv = nums[left], nums[right]
                while left < right and nums[left] == lv: left += 1
                while left < right and nums[right] == rv: right -= 1
    return ans''',
'排序 O(n log n)，双指针主体 O(n²)，额外空间取决于排序实现，不计输出通常为 O(log n) 或 O(1)。不能用结果集合事后去重来替代指针去重，因为会产生大量重复计算。还可用最小/最大可能和做剪枝，但应先保证基础版本正确。'
]),
('158. 求二维网格中岛屿的最大面积。', [
'我的回答是：把每个值为 1 的格子视为图节点，四联通边表示相邻。遍历网格，遇到未访问陆地就用 DFS/BFS 搜索整个连通分量并计数，更新最大面积。访问标记可以单独维护，也可以在允许修改输入时把访问过的 1 改成 0。',
'''def max_area_of_island(grid):
    if not grid or not grid[0]: return 0
    rows, cols, best = len(grid), len(grid[0]), 0
    for r in range(rows):
        for c in range(cols):
            if grid[r][c] != 1: continue
            grid[r][c] = 0; stack = [(r, c)]; area = 0
            while stack:
                x, y = stack.pop(); area += 1
                for nx, ny in ((x-1,y),(x+1,y),(x,y-1),(x,y+1)):
                    if 0 <= nx < rows and 0 <= ny < cols and grid[nx][ny] == 1:
                        grid[nx][ny] = 0; stack.append((nx, ny))
            best = max(best, area)
    return best''',
'时间复杂度 O(rows×cols)，每个格子最多入栈一次；最坏空间 O(rows×cols)。我更倾向显式栈，避免大岛屿导致递归栈溢出。边界包括空网格、全水、全陆；如果不允许修改 grid，就使用 visited 集合或布尔矩阵。若对角线也算相邻，只需扩展方向集合。'
]),
('159. 编辑距离怎么实现？', [
'我的回答是：定义 dp[i][j] 为把 word1 的前 i 个字符变成 word2 的前 j 个字符所需最少操作数。最后一个字符相同就继承 dp[i-1][j-1]；不同则在删除 dp[i-1][j]、插入 dp[i][j-1]、替换 dp[i-1][j-1] 三种方案中取最小再加 1。',
'''def min_distance(a, b):
    prev = list(range(len(b)+1))
    for i, ca in enumerate(a, 1):
        cur = [i] + [0] * len(b)
        for j, cb in enumerate(b, 1):
            if ca == cb: cur[j] = prev[j-1]
            else: cur[j] = 1 + min(prev[j], cur[j-1], prev[j-1])
        prev = cur
    return prev[-1]''',
'初始化 dp[i][0]=i、dp[0][j]=j，表示与空串转换只能连续删除或插入。时间复杂度 O(mn)，滚动数组空间 O(n)，可以让 b 为较短字符串降到 O(min(m,n))。若要恢复具体操作序列，需要保留完整 DP 表并反向追踪；若三种操作成本不同，就把对应转移加上各自权重。'
]),
('160. 爬楼梯问题如何解？', [
'我的回答是：如果每次能走 1 或 2 级，到第 n 级的最后一步只能来自 n-1 或 n-2，因此 f(n)=f(n-1)+f(n-2)。这是斐波那契型动态规划，不需要保存整张数组，只保留前两个状态。',
'''def climb_stairs(n):
    if n <= 0: return 0
    prev2, prev1 = 1, 1  # f(0), f(1)
    for _ in range(2, n+1):
        prev2, prev1 = prev1, prev1 + prev2
    return prev1''',
'时间复杂度 O(n)，空间 O(1)。这里把 f(0)=1 理解为“什么都不走”这一种组合，使递推统一；对外若 n<=0 是否返回 0 要按题意约定。若允许步长集合 steps，可以用 dp[i]=sum(dp[i-step])；若 n 极大，可用矩阵快速幂做到 O(log n)。面试时也应指出朴素递归会产生指数级重复计算。'
]),
('161. 合并重叠区间。', [
'我的回答是：先按起点排序，随后线性扫描。当前区间起点小于等于已合并结果最后一个区间的终点，说明存在重叠或接触，就更新终点为两者最大值；否则开启一个新区间。排序把可能相交的区间放到一起，使局部判断成立。',
'''def merge_intervals(intervals):
    if not intervals: return []
    intervals.sort(key=lambda x: x[0])
    merged = [intervals[0][:]]
    for start, end in intervals[1:]:
        if start <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], end)
        else:
            merged.append([start, end])
    return merged''',
'时间复杂度 O(n log n)，主要来自排序；扫描 O(n)，结果空间 O(n)。需要先确认区间是闭区间还是半开区间：[1,2] 和 [2,3] 在闭区间中应合并，在半开区间中通常不合并。还要考虑输入是否允许原地排序、起止点是否合法，以及输出是否要保持不可变结构。'
]),
('162. 两个链表表示两个非负整数，求它们之和。', [
'我的回答是：若链表按低位在前存储，就同时遍历两个链表并维护进位 carry，每轮把两个当前节点值和 carry 相加，当前结果位是 sum % 10，新进位是 sum // 10。循环条件必须包含 carry，否则 5+5 最后的进位会丢失。',
'''def add_two_numbers(l1, l2):
    dummy = ListNode(0); tail = dummy; carry = 0
    while l1 or l2 or carry:
        x = l1.val if l1 else 0
        y = l2.val if l2 else 0
        carry, digit = divmod(x + y + carry, 10)
        tail.next = ListNode(digit); tail = tail.next
        l1 = l1.next if l1 else None
        l2 = l2.next if l2 else None
    return dummy.next''',
'时间复杂度 O(max(m,n))，新链表空间同阶。dummy 节点用于统一首节点和后续节点的连接逻辑。输入长度不同按缺位补 0；题目通常保证不存在多余前导零。若数字是高位在前，可以用两个栈逆序处理，或在允许修改输入时先反转链表。'
]),
('163. 有一个 10GB 日志文件，但只有 4GB 内存，如何统计或排序？', [
'我的回答是：先澄清任务是统计、去重、TopK 还是全量排序，因为方案不同。共同原则是流式读取和外部存储，绝不能把整个文件一次性加载进内存。简单计数或聚合可逐行解析并更新有界状态；高基数 Key 需要按 hash 分桶到多个小文件，再逐桶处理。',
'如果要求全量排序，我会使用外部归并排序：按可用内存切成多个块，每块读入、排序后写成有序 run；第二阶段用最小堆做 k 路归并，顺序写出结果。4GB 不能全部给业务数据，还要预留运行时、缓冲区和堆内存，因此块大小可能控制在数百 MB，并限制同时打开的文件数，必要时分层归并。',
'复杂度方面，CPU 仍约为 O(n log n)，但瓶颈主要是磁盘 I/O。要使用大块顺序读写、压缩权衡 CPU 与磁盘、临时文件校验和断点续跑。若只求 TopK，用大小为 k 的小顶堆可把内存降到 O(k)；若求 UV，可以按 hash 精确去重，或在允许误差时用 HyperLogLog。面试中最重要的是根据目标选择算法，而不是机械回答“分块”。'
]),
('164. 有一千条数据需要求和，如何设计并行方案？', [
'我的回答是：一千个普通数字求和非常轻，单线程循环通常比创建线程、序列化和调度更快，所以第一判断是不应为了并行而并行。只有每条“数据”的获取或计算本身昂贵，或者题目明确要求分布式处理，才值得拆分。',
'若每条数据需要远程读取，我会按批次并发拉取，设置并发上限、超时和重试，再在本地归约；若是 CPU 密集型计算，可把 1000 条按核数切成若干连续分片，每个 Worker 返回部分和，Coordinator 再求和。分布式情况下给每个分片稳定 task_id，结果幂等写入，失败只重算分片，最终归约前检查所有分片齐全。',
'数值类型也要明确：整数可能溢出，应使用足够宽的类型；浮点数并行归约改变加法顺序，结果可能有微小差异，可用 pairwise summation 或 Kahan 补偿。设计评价指标是总耗时、资源和正确性，而不是线程数量。'
]),
('165. 设计一个短链系统。', [
'我的回答是：核心接口是创建短链和重定向。创建时校验原 URL、生成全局唯一 ID，再用 Base62 编码为短码，保存 short_code→long_url、创建者、过期时间和状态。跳转时读取映射并返回 301 或 302；若需要实时变更和点击统计，通常选 302，永久且强缓存的链接才适合 301。',
'读流量远大于写流量，链路是 CDN/网关→本地缓存或 Redis→持久化数据库。数据库按 short_code 主键查询，可按 ID 范围或一致性哈希分片。ID 可来自号段服务或 Snowflake；如果用随机码，则需足够位数并依赖唯一索引处理碰撞。热门短链要缓存和请求合并，不存在或已失效的短码也可做短期负缓存。',
'点击统计不应阻塞跳转，请求只异步投递事件，由流处理聚合 PV、UV、来源等指标。安全上要做恶意网址检测、域名黑名单、限流、权限和防枚举；删除采用状态失效并让缓存快速失效。容量估算要从链接新增 QPS、跳转峰值 QPS、保留年限和每条记录大小推导，而不是直接报固定机器数量。'
]),
('166. 如何设计一个支持百万 QPS 的缓存系统？', [
'我的回答是：先拆清百万 QPS 是总量还是单 Key 峰值、读写比例、对象大小、允许的延迟和一致性。典型架构是客户端本地缓存 + 分布式缓存集群 + 数据源，多级缓存分别吸收进程内热点和全局流量。请求通过一致性哈希或 Redis Cluster 槽位分片，节点部署副本并跨故障域。',
'热点 Key 是主要风险：单 Key 不会因普通分片而摊开。我会使用本地缓存、热点复制、请求合并和逻辑过期，避免大量请求同时回源。缓存穿透用参数校验、布隆过滤器和短期负缓存；击穿用 singleflight/互斥重建；雪崩用 TTL 抖动、分批预热、限流、熔断和返回陈旧值。',
'一致性策略按业务选择 Cache Aside 或写穿。更新数据库后删除缓存时要处理失败，可用事务消息/CDC 做最终一致；强一致数据不应依赖普通缓存直接决策。容量要按有效 QPS、平均对象大小、命中率、副本数和网络带宽估算，并监控命中率、P99、热 Key、驱逐率、连接数、复制延迟和回源 QPS。百万 QPS 的答案是分层与消峰，不是单纯“多加 Redis”。'
]),
('167. 设计一个自动售卖机。', [
'我的回答是：我会先把它建模为状态机，而不是堆 if-else。核心状态包括 Idle、Selecting、Paying、Dispensing、Refunding、OutOfService；事件包括选择商品、投入金额、支付成功、取消、出货成功/失败和超时。每次转换有明确前置条件和副作用，非法事件在当前状态下被拒绝。',
'领域对象至少有商品、货道库存、订单、支付和找零模块。用户选择后先检查价格与可售库存；支付成功后生成不可重复消费的 payment_id，出货使用幂等 order_id 驱动电机。只有传感器确认出货成功才完成订单；失败则重试有限次数并退款/转人工，同时记录故障。现金场景还要检查找零能力，电子支付则要处理支付成功回调晚于用户取消的竞态。',
'并发上，一台机器同一时刻通常串行处理一个会话，但后台补货和云端配置可能并发，需要版本号或锁保护库存。断网时可允许现金交易或缓存白名单商品，电子支付根据风险策略降级。审计日志、库存对账、支付对账和设备心跳保证可运营性。代码设计可用 State Pattern 隔离各状态行为，用硬件接口抽象真实电机与测试模拟器。'
]),
('168. 线上接口延迟突然飙升，你如何排查？', [
'我的回答是：先止损再定位。先确认影响范围、开始时间、P50/P95/P99、错误率和吞吐，判断是全部请求、某个机房、某个版本还是特定依赖；必要时回滚最近发布、关闭高成本功能、限流或扩容。只看平均延迟会掩盖长尾。',
'排查按请求链路自外向内：网关排队与连接、应用线程池/事件循环、CPU/内存/GC、数据库慢查询与连接池、缓存命中率、下游 RPC、网络丢包和磁盘。结合 Trace 找耗时跨度最大的 Span，再用指标和日志验证。若吞吐不变但延迟升高，常见是资源饱和或下游变慢；吞吐同时激增则要检查流量与热 Key；错误率升高还要看超时重试是否形成放大。',
'以 RepoMind 为例，我会拆分 GitHub API、clone、embedding、RAG、CrossEncoder、LLM 各阶段耗时，检查是否某仓库过大、模型限流、工具重试或并行分支等待最慢节点。恢复后需要留下时间线和根因证据，并补充分阶段超时、并发上限、熔断降级、容量告警及回归压测，不能把“扩容后好了”当成根因。'
]),
('169. 如何用一个公平硬币实现指定概率？', [
'我的回答是：如果目标概率是 a/b，可以生成 [0,b) 上的均匀整数，落在前 a 个值时返回成功。用公平硬币生成 k=ceil(log2 b) 位二进制数 x；因为 2^k 通常不是 b 的倍数，所以 x>=b 时必须拒绝并重试，不能直接取模，否则前面的余数会概率更高。',
'''def bernoulli(a, b, flip):
    if not 0 <= a <= b: raise ValueError
    k = (b - 1).bit_length()
    while True:
        x = 0
        for _ in range(k):
            x = (x << 1) | flip()  # flip() 等概率返回 0/1
        if x < b:
            return x < a''',
'该方法结果严格等于 a/b，单轮接受率 b/2^k，大于 1/2，所以期望硬币次数小于 2k。若概率是有限二进制小数，可直接按位比较；任意实数概率只能在给定精度下近似，或使用其可计算二进制展开。边界 a=0 和 a=b 可以直接返回，减少无意义抛掷。'
]),
('170. 如果要做一个 TUI 视频剪辑 MVP，你会怎么设计？', [
'我的回答是：MVP 目标不是在终端复制完整 Premiere，而是完成一条可靠的核心链路：导入视频→查看元数据/低成本预览→设置入点和出点→拼接片段→导出。首版只支持本地文件、单视频轨和少量转场，明确不做复杂时间线、多层特效和云协作。',
'界面可分为媒体列表、时间线/片段列表、预览信息和命令/状态区。TUI 框架负责键盘交互与状态渲染，ffprobe 读取时长、编码和关键帧，ffmpeg 执行裁剪、拼接和转码。预览不强求终端实时播放，可生成低分辨率代理、ASCII 帧或调用外部播放器跳转到当前时间点。项目文件保存源文件路径、片段 in/out、顺序和导出配置，采用 JSON 便于恢复与测试。',
'执行层把每次编辑表示为不可变操作，便于 undo/redo；导出作为后台 Job，显示进度、允许取消，并保留 ffmpeg 日志。能 stream copy 时快速无损裁剪，但非关键帧精确裁剪通常需要转码，界面要明确提示速度与精度取舍。MVP 验收指标包括：能处理常见 MP4、时间边界正确、失败信息可理解、项目可恢复、导出文件音画同步。若继续演进，再增加波形、字幕、代理缓存和插件系统。'
]),
]

B = RGBColor(46, 116, 181)
D = RGBColor(31, 77, 120)
G = RGBColor(95, 103, 112)
CODE_FILL = 'F3F5F7'

def set_font(run, size=11, bold=None, color=None, name='Arial Unicode MS'):
    run.font.name = name
    run.font.size = Pt(size)
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
        fonts.set(qn('w:' + attr), name)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color

def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    ppr.append(shd)

def add_content_paragraph(doc, text):
    if '\n' not in text:
        p = doc.add_paragraph(text)
        p.paragraph_format.keep_together = True
        return
    intro, code = text.split('\n', 1)
    if intro.strip():
        p = doc.add_paragraph(intro)
        p.paragraph_format.keep_together = True
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(.18)
    p.paragraph_format.right_indent = Inches(.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.keep_together = True
    shade_paragraph(p, CODE_FILL)
    r = p.add_run(code)
    set_font(r, size=9.2, name='Menlo')

doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin = section.bottom_margin = Inches(.82)
section.left_margin = section.right_margin = Inches(.88)
section.header_distance = section.footer_distance = Inches(.42)

normal = doc.styles['Normal']
normal.font.name = 'Arial Unicode MS'; normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, before, after, color in [('Heading 1', 16, 16, 8, B), ('Heading 2', 13, 12, 6, D)]:
    style = doc.styles[name]
    style.font.name = 'Arial Unicode MS'; style.font.size = Pt(size); style.font.bold = True
    style.font.color.rgb = color; style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

set_font(section.header.paragraphs[0].add_run('RepoMind · Agent 开发面试问答'), 9, color=G)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run('第八批 C · 第 151—170 题'), 9, color=G)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(70); p.paragraph_format.space_after = Pt(10)
set_font(p.add_run('RepoMind Agent 开发面试问答'), 24, True, D)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(35)
set_font(p.add_run('第八批 C：算法与场景设计（20 题）'), 14, color=B)
doc.add_paragraph('本批为第 151—170 题。算法题覆盖思路、关键实现、复杂度与边界；场景题强调分层设计、可靠性和工程取舍。')
doc.add_page_break()

for question, paragraphs in QA:
    doc.add_heading(question, level=1)
    for text in paragraphs:
        add_content_paragraph(doc, text)

doc.core_properties.title = 'RepoMind Agent 开发面试问答：第八批C 151—170题'
doc.core_properties.author = 'RepoMind'
doc.core_properties.subject = '算法与场景设计面试问答'
doc.save(OUT)
print(OUT)
