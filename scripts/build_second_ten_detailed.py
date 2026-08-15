from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT=Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第二批10题.docx')
BLUE=RGBColor(46,116,181); DARK=RGBColor(31,77,120); GRAY=RGBColor(95,103,112)

QA=[
('11. Agent 的状态应该如何建模？为什么不能只保存 messages？',[
'我的回答是：Agent 的状态应该围绕“业务事实、执行进度和控制信息”显式建模，messages 只适合保存模型对话和工具交互记录，不能承担全部状态。因为 messages 本质上是按时间排列的自然语言历史，如果路由节点想知道当前 Issue 类型、重试次数或已经找到的文件，就必须重新让模型从对话里提取，不仅浪费 Token，也会引入不确定性。',
'RepoMind 使用 TypedDict 分别定义 IssueState 和 PRState。IssueState 中除了 issue_url、title、body、comments 等输入数据，还显式保存 issue_type、is_regression、fix_type、root_cause、affected_files、severity、confidence、iteration 和 error。上下文并行阶段还使用 _entry_ctx、_rag_ctx、_git_ctx 三个临时字段，最后由 merge_context 合并成 code_context。',
'这种设计有三个好处。第一，条件路由可以直接读取 issue_type 和 confidence，不依赖模型再次判断；第二，Checkpoint 可以恢复到明确的业务状态；第三，评测和排错能够知道错误发生在哪个字段和节点，而不是只看到一长串对话。messages 仍然需要保存，因为 ReAct 必须知道模型调用过什么工具、工具返回了什么，但它只是状态的一部分。',
'状态建模也不能无节制地把所有内容都塞进去。长代码和原始响应应保存引用、摘要或外部存储地址；真正影响路由和恢复的字段才进入核心 State。否则 Checkpoint 会越来越大，并行合并和版本升级也会更困难。'
]),
('12. Workflow 和 Agent 的区别是什么？RepoMind 为什么把两者结合使用？',[
'我的回答是：Workflow 的控制权主要在代码，开发者提前定义节点、顺序和条件；Agent 的控制权更多在模型，模型根据目标和环境反馈动态决定下一步。Workflow 追求确定性、可恢复和可观测，Agent 擅长处理无法提前枚举完整路径的问题。',
'RepoMind 中，获取 GitHub 数据、构建项目地图、分类 Issue、并行收集上下文、Reflection 和生成报告，都属于 Workflow。它们的先后关系明确，失败后也需要知道从哪里恢复。代码分析阶段则属于 Agent：拿到候选上下文后，模型可能需要读文件、搜索符号、查调用者、看 Git 历史或检索依赖源码，下一步取决于刚刚获得的证据，无法完全写死。',
'如果全部做成 Workflow，我必须提前枚举大量代码搜索路径，面对不同仓库会非常僵硬；如果全部交给 Agent，模型可能跳过必要步骤、重复调用工具或无限循环，Checkpoint 也难以表达稳定状态。因此项目使用外层 LangGraph Workflow 兜住业务流程，内层 ReAct Loop 提供动态探索能力。',
'这不是简单地“固定流程加一个模型节点”，而是明确控制权边界：框架决定必须经过哪些阶段、工具预算是多少、什么时候停止和重试；模型只在允许的阶段和工具集合内自主决策。'
]),
('13. Workflow 的 Loop 和 Agent 的 Loop 各自怎么设计？业务 Workflow 怎么往 Loop 上设计？',[
'我的回答是：Workflow Loop 和 Agent Loop 最大的区别是循环条件由谁决定。Workflow Loop 通常由确定性条件控制，例如重试次数、分页游标、状态是否完成；Agent Loop 则由模型根据观察结果判断是否继续调用工具。两者都必须有代码级停止条件，不能把退出完全交给模型。',
'RepoMind 有两类 Loop。第一类是 ReAct 工具循环：模型生成 tool_calls，系统执行后把 tool result 加回 messages，再让模型继续判断；不同节点最多执行 6 到 12 轮。第二类是 Reflection 回路：critic 给出 confidence，低于 0.7 且 iteration 小于 2 时，Issue 根据 issue_type 回到对应分析节点，PR 回到 prepare_reviews 重新并行审查。',
'业务 Workflow 如果要嵌入 Agent Loop，我会把 Loop 封装成一个职责明确的节点。节点输入是结构化任务和上下文，内部允许模型进行有限次 ReAct，输出必须经过 Schema 校验，再交还外层 Workflow。外层只关心节点是否成功、输出是否合格和是否需要重试，不需要理解模型每一轮的思考。',
'这样设计的好处是，外层图仍然可观测、可持久化，内层 Agent 又能处理不确定问题。需要注意的是，工具调用过程也要记录 trace，否则只在 Loop 结束时保存最终答案，中间失败仍然无法定位。'
]),
('14. Workflow 中的原子性在 RepoMind 里怎么体现？',[
'我的回答是：原子性首先体现在节点职责单一，一个节点只完成一种可以清楚描述、独立测试的状态转换；其次体现在节点输出要么形成完整的新状态，要么明确失败，不能悄悄留下半成品。',
'RepoMind 把获取 Issue、构建项目地图、解析堆栈、问题分类、回归判断、Memory 召回、上下文准备、分析、Reflection 和报告生成拆成独立节点。例如 route_issue 只负责输出 issue_type，不同时做代码检索；classify_fix 只判断 code_fix、dependency_upgrade 或 config_fix，不直接生成最终修复方案；generate_report 只消费已经确认的分析字段并组织报告。',
'并行上下文收集更能体现这个原则。entry、rag、git 三个分支不会同时写 code_context，而是分别写 _entry_ctx、_rag_ctx、_git_ctx；每个分支完成自己的证据收集后，再由 merge_context 一次性组合。这样避免两个并行节点互相覆盖，也让任一路失败时更容易识别。',
'不过 RepoMind 当前大部分工具是读取型操作，所以补偿问题相对简单。如果未来加入自动提交代码、创建 PR 等写操作，单纯的节点拆分还不够，还需要幂等键、执行前后状态、查询确认和 Saga 补偿。面试时我会明确，这是生产化扩展，而不是当前已经实现的能力。'
]),
('15. Workflow 的每个节点是怎么控制的？Eval 的结果如何影响后续？',[
'我的回答是：节点控制不能只依赖模型输出一句“完成了”。我会把控制分成输入约束、执行过程、输出校验和路由决策四层。确定性节点使用代码校验；模型节点要求结构化输出并通过 JSON/Pydantic 校验；超时和网络错误按错误类型重试；质量是否达标则交给 Reflection 或评测节点。',
'RepoMind 的 LLM Client 会解析模型 JSON，处理代码块包裹和格式异常，并用 Pydantic 对关键输出做字段类型和枚举约束。图中的路由不会读取一段自然语言再猜结果，而是读取 issue_type、confidence、iteration 等明确字段。例如 route_issue 决定进入 bug、feature、question 还是 security；Reflection 把 confidence 夹在 0 到 1，并增加 iteration。',
'Eval 对后续的直接影响体现在 Reflection 条件边。confidence 低于 0.7 且重试次数未达上限时，系统不会生成报告，而是回到相应分析节点；通过后才进入 generate_report。PR 的 retry 回 prepare_reviews，跳过已经完成的 fetch、项目地图和 Memory，避免为了重评质量把昂贵的前置步骤全部重跑。',
'当前 Reflection 主要依赖独立 LLM critic，优点是能判断证据是否充分，缺点是评分可能漂移。更稳妥的生产方案是把规则校验、Schema、引用存在性、工具执行状态和 LLM Judge 组合起来，而不是把所有 Eval 都交给另一个模型。'
]),
('16. 如何设计一个能够恢复中断的长任务 Agent？',[
'我的回答是：长任务不能只把执行进度保存在进程内存或对话上下文中，而要把它建模成可持久化的状态机。恢复时需要知道任务身份、当前节点、已经完成的节点、关键输出、重试次数以及外部副作用是否已经发生。',
'RepoMind 使用 LangGraph SqliteSaver 作为 Checkpointer。Issue 和 PR 分别写入 checkpoints.sqlite 和 pr_checkpoints.sqlite。后端为每次任务生成 job_id，并把它作为 configurable.thread_id 传给 Graph；同一个 thread_id 对应一条可恢复的执行历史。State 中的 issue_type、code_context、analysis、confidence 和 iteration 等字段会随图执行保存。',
'恢复设计的关键是不要把 Checkpoint 和 Memory 混淆。Checkpoint 服务当前任务的中断续跑，保存的是执行状态；ChromaDB Memory 服务未来相似任务的经验召回，保存的是完成后的案例。二者生命周期和一致性要求不同。',
'还要诚实说明项目边界：当前 LangGraph 图状态可以持久化，但 backend/app.py 中的 jobs 字典仍然在单进程内存里，服务重启后 Web 层任务列表不会自动恢复。生产化需要把 Job 元数据放到数据库或队列中，由 Worker 根据 thread_id 恢复 Graph，并为写工具增加幂等和状态查询。'
]),
('17. 服务重启后如何恢复任务？怎样避免已经完成的步骤重新执行？',[
'我的回答是：服务重启后不能重新拿原始输入从 START 执行，而要使用原来的任务标识读取最近一次稳定 Checkpoint。编排器恢复当前 State 和下一待执行节点，只调度未完成部分。',
'在 RepoMind 中，这个任务标识就是 thread_id，后端目前使用 job_id 作为 thread_id。SqliteSaver 会按 thread_id 保存图执行状态。如果任务在 analyze_bug 前已经完成 fetch_issue、build_project_map 和上下文收集，恢复时应从持久化位置继续，而不是重新拉取和索引仓库。PR 的 Reflection retry 也专门回 prepare_reviews，跳过 Memory 和其他前置节点。',
'不过“节点已经执行”不一定等于“副作用没有问题”。读取 GitHub、构建索引等操作通常可以安全重试；创建 PR、写数据库、退款等写操作则要为每次业务动作生成幂等键，并在执行前记录 PENDING，执行后记录 SUCCEEDED。网络超时时先按幂等键查询状态，不能直接重放。',
'因此完整恢复依赖三部分：Graph Checkpoint 确定流程位置，持久化 Job 表负责重新发现未完成任务，工具执行日志和幂等键保证副作用安全。RepoMind 当前主要完成第一部分，后两部分是上线前需要补齐的工程能力。'
]),
('18. Agent 服务同时收到两次对同一会话的请求，如何避免状态覆盖？',[
'我的回答是：不能采用简单的“读取整个 State—修改—整体覆盖写回”。两个请求可能都基于相同旧版本执行，后提交者会覆盖先提交者的结果，甚至让 Graph 出现不一致的节点状态。',
'第一种方案是按 conversation_id 或 thread_id 串行化。同一会话的请求进入同一个队列分区，或者获取带租约的分布式锁；不同会话仍然可以并行。第二种方案是乐观锁，在状态表中保存 version，更新时使用 where version=old_version；失败说明状态已被别人修改，需要重新加载或提示冲突。',
'RepoMind 已经使用 thread_id 隔离不同任务，但当前 backend 的 jobs 是内存字典，并没有实现多实例下的会话锁或版本控制。因此它可以避免不同 job 混用同一个 Checkpoint，却不能声称已经解决分布式并发写覆盖。',
'如果后续支持用户在同一任务上追加问题，我会把用户事件作为追加日志保存，而不是直接覆盖 State；调度器按序消费事件并生成新版本 Checkpoint。对于不可合并的并发修改，宁可明确返回冲突，也不要静默采用最后写入者覆盖。'
]),
('19. 首次生成和多轮补充的链路路由，应该怎么区分和实现？',[
'我的回答是：首次生成和多轮补充的目标不同。首次生成需要完成任务初始化、数据抓取、索引和完整分析；多轮补充通常应该复用已有仓库、证据和结论，只针对用户的新问题增量执行，不能每次都从头跑完整链路。',
'路由时我会显式保存 request_type、parent_run_id、用户新问题和上一轮结构化结果。首次请求从 START 进入 fetch；补充请求先判断新问题是要求解释原结论、补充证据、修改分析范围，还是完全切换任务。如果只是解释，可以直接读取原 State 和报告；如果要求补证据，可以回到 prepare_context 或 ReAct；如果输入了新的 Issue/PR，则创建新 thread。',
'RepoMind 当前 CLI 和 Job API 主要按一次 URL 对应一次完整任务设计，还没有完整的多轮补充路由。已有可复用基础是 thread_id、Checkpoint、Memory 和显式 State；但需要新增事件入口、父子运行关系及不同补充意图的条件边。',
'面试时我不会把这个能力说成已经上线。我会说明当前实现支持任务状态持久化和相似案例召回，多轮补充是合理演进方向；实现重点是复用确定性成果，同时避免旧结论被用户新问题无意覆盖。'
]),
('20. 你如何看待 Agent 的自进化？RepoMind 能不能做自进化？',[
'我的回答是：我对“自进化”会比较谨慎。它不应该等同于 Agent 自动把所有历史对话写进 Memory，也不应该让模型自行修改 Prompt 和工具后直接上线。更可靠的定义是：系统从执行日志和 Badcase 中提取候选经验，经过评测和审核后，更新可召回案例、Prompt、工具或固定流程。',
'RepoMind 已经具备自进化闭环的一部分基础。任务完成后会把结构化分析结果写入 ChromaDB Memory，新任务开始时召回相似案例；eval 目录包含数据集、Runner、LLM Judge 和报告，可以比较修改前后的根因、文件、严重度和修复建议；Reflection 也会记录低置信度反馈。',
'但当前 Memory 写入更接近案例积累，并不等于系统已经验证这些经验一定正确。错误案例如果长期保存，反而会污染后续任务。因此更稳妥的流程是：从低分 Trace 和人工确认案例中生成候选经验，去重并绑定来源、仓库、版本和置信度；在隔离评测集上验证；只有稳定提升后才升级为 Prompt 规则、检索策略或 Skill。',
'我认为现阶段最现实的是“人工监督的持续改进”，而不是完全自主进化。模型可以提出修改建议，但是否写入长期 Memory、是否改变工具权限、是否发布新策略，必须经过确定性评测和人工审批。'
]),
]

def set_font(run,size=11,bold=None,color=None):
 name='PingFang SC'; run.font.name=name; run.font.size=Pt(size); rf=run._element.get_or_add_rPr().get_or_add_rFonts()
 for a in ('ascii','hAnsi','eastAsia','cs'): rf.set(qn('w:'+a),name)
 if bold is not None: run.bold=bold
 if color is not None: run.font.color.rgb=color

doc=Document(); sec=doc.sections[0]; sec.page_width,sec.page_height=Inches(8.5),Inches(11); sec.top_margin=sec.bottom_margin=Inches(.82); sec.left_margin=sec.right_margin=Inches(.88)
normal=doc.styles['Normal']; normal.font.name='PingFang SC'; normal.font.size=Pt(11); normal._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC'); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.25
for nm,sz,bef,aft,col in [('Heading 1',16,16,8,BLUE),('Heading 2',13,12,6,DARK)]:
 st=doc.styles[nm]; st.font.name='PingFang SC'; st.font.size=Pt(sz); st.font.bold=True; st.font.color.rgb=col; st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC'); st.paragraph_format.space_before=Pt(bef); st.paragraph_format.space_after=Pt(aft); st.paragraph_format.keep_with_next=True
h=sec.header.paragraphs[0]; set_font(h.add_run('RepoMind · Agent 开发面试问答'),9,color=GRAY)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; set_font(f.add_run('第二批 · 第 11—20 题'),9,color=GRAY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70); p.paragraph_format.space_after=Pt(10); set_font(p.add_run('RepoMind Agent 开发面试问答'),24,True,DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(35); set_font(p.add_run('第二批：状态、Workflow 与可靠性（10 题）'),14,color=BLUE)
doc.add_paragraph('本批为第 11—20 题，逐题结合 RepoMind 当前代码整理。项目已实现能力与生产化改进会明确区分。'); doc.add_page_break()
for q,paras in QA:
 doc.add_heading(q,level=1)
 for t in paras:
  p=doc.add_paragraph(t); p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第二批10题'; doc.core_properties.author='RepoMind'; doc.save(OUT); print(OUT)
