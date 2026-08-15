from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches,Pt,RGBColor
OUT=Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第五批20题.docx')
QA=[
('51. 如何设计 AI Agent 的完整评测体系？',[
'我的回答是：Agent 评测不能只看最终答案，因为一次任务还包含意图识别、检索、规划、工具选择、参数生成、错误恢复和成本。合理做法是分层评测，并把最终失败归因到具体环节。',
'RepoMind 可以分三层。检索层看关键文件或函数 Recall@K、MRR、NDCG；分析层看 issue_type、root_cause、affected_files、severity 和 fix_suggestion；链路层看工具选择、参数正确率、循环次数、恢复成功率、延迟和 Token。PR 还要分别评 security、logic、style 的问题检出率和误报率。',
'项目 eval 目录已经提供 Dataset、Runner、LLM Judge 和 Report，tests/regression 保存 smoke cases 与 baseline。完整体系还应加入人工金标准、独立留出集、线上采样和版本化 Trace，不能只依赖一个 Judge 分数。'
]),
('52. 没有用户反馈时，怎么对 Agent 做有效抽检？',[
'我的回答是：没有显式反馈不代表没有评测信号。可以按风险、置信度和行为异常分层抽样，而不是纯随机查看少量“看起来正常”的结果。',
'我会重点抽检低 confidence、Reflection 重试、工具轮数异常、耗时过长、检索无结果、结论与历史案例冲突的任务；再随机抽取一部分高分任务，防止 Judge 与模型共同偏差导致虚高。每条样本由人工检查证据是否真实、根因是否可复现、建议是否可执行。',
'RepoMind 已记录 confidence、iteration 和结构化结果，适合作为抽样条件。后续应把节点 Trace、工具调用和候选检索结果一起保存，否则人工只能看到最终报告，无法判断结论是如何产生的。'
]),
('53. 出现 Badcase 时，怎么快速定位是检索、规划、工具还是模型的问题？',[
'我的回答是：要沿数据链路逐层排查，不能看到最终答案错了就直接改 Prompt。第一步确认 GitHub 数据和仓库版本是否正确；第二步看关键文件是否进入初始召回和精排结果；第三步检查 Agent 是否选择正确工具、参数是否有效；第四步才判断模型是否在证据充分时推理错误。',
'如果关键代码没有进入 top-k，问题在切片、Query、embedding 或 rerank；正确代码已进入 code_context，但模型没有阅读，可能是上下文组织或规划问题；模型调用了正确工具但工具返回错误，则查 API、路径和版本；证据正确且完整仍得出错误结论，才是分析 Prompt 或模型能力问题。',
'RepoMind 的节点边界比较清晰，适合按 fetch、retrieve、analyze、reflect、report 做归因。修复后要新增回归 Case，并检查相邻类别，避免只修当前问题。'
]),
('54. LLM-as-a-Judge 有哪些系统性偏差？怎么降低？',[
'我的回答是：常见偏差包括位置偏差、偏爱更长或更流畅的答案、自我偏好、对格式敏感，以及缺少真实代码环境时无法验证事实。Judge 还可能把“听起来合理”误判为“证据正确”。',
'降低偏差可以随机交换候选顺序、隐藏模型身份、使用明确 Rubric 和分项评分、要求 Judge 引用证据、对关键样本使用多 Judge 或人工复核。确定性指标应优先用代码计算，例如文件是否命中、JSON 是否合法、工具是否成功。',
'RepoMind 的 Judge 适合评根因、文件、严重度和修复可执行性，但不能单独证明系统正确。高风险或分数边界样本需要人工核验，Judge Prompt 和模型版本也要固定并纳入版本记录。'
]),
('55. 如何证明一次 Agent 修改是正向优化，而不是在评测集上过拟合？',[
'我的回答是：必须把开发集、回归集和独立留出集分开。Badcase 用于开发，固定回归集保证旧能力不退化，留出集在方案确定后才评测；如果只在反复调试的同一批样本上提升，不能说明泛化。',
'比较时不仅看平均总分，还要按 bug/feature/question/security、仓库和难度分层，并同时观察最差分位、失败率、延迟、Token 和工具次数。统计显著性和逐 Case 差异比单一平均数更有价值。',
'RepoMind 可以利用 eval 数据集和 regression baseline 做前后对比。每次修改应记录 Prompt、模型、检索参数和代码版本，新增 Badcase 回归，但不能把所有新增样本继续当作独立测试集。'
]),
('56. Prompt 调优遇到“修好一类、坏了另一类”怎么解决？',[
'我的回答是：这通常说明一个大 Prompt 同时承担了相互冲突的任务，或者修改只针对局部案例。先用回归集确认退化属于哪种类型，再考虑路由、上下文和结构拆分，而不是继续堆补丁规则。',
'RepoMind 已经按 issue_type 拆分 analyze_bug、analyze_feature 和 answer_question，PR 也按 security、style、logic 分路 Prompt。某个 bug 规则不应直接污染 question 分支；公共要求只保留证据、格式和安全约束。',
'每次 Prompt 修改要版本化并跑全量分层回归。若规则越来越多，应把确定性要求移到 Schema 或代码校验，把专业策略下沉到独立节点，减少 Prompt 内部冲突。'
]),
('57. Prompt 调到极限后，优先换模型、调参数还是做 SFT？',[
'我的回答是：先定位瓶颈。若缺少证据，换模型和 SFT 都不会解决，应该修检索或工具；若输出格式不稳，优先结构化输出和校验；若复杂推理持续失败，再评估更强模型。',
'参数调整适合控制随机性、长度和成本，但不会创造缺失能力。SFT 适合有大量稳定高质量样本、目标行为明确且 Prompt 已难以表达的场景，成本和维护要求最高，不应作为第一反应。',
'RepoMind 当前通过 API 使用基座模型，没有声称完成 SFT。合理顺序是修链路与数据、优化 Prompt/Schema、做模型路由，最后才用真实 Badcase 数据评估 SFT。'
]),
('58. RepoMind 的 Reflection 模块是怎么设计的？',[
'我的回答是：Reflection 是独立 critic 节点，不与 analyze 节点共享即时推理过程。它读取 Issue、代码上下文和结构化分析结果，检查根因是否有证据、影响文件是否合理、是否存在其他可能原因，并输出 confidence 和 note。',
'reflect.py 会把 confidence 限制在 0 到 1，并把 iteration 加一。graph.py 只有在 confidence<0.7 且 iteration<2 时才按 issue_type 回到对应分析节点，否则生成报告。PR 的 reflect_pr 也采用同样阈值，retry 回 prepare_reviews。',
'这种 Actor-Critic 分离能减少自我确认偏差，但仍可能出现 critic 漂移。生产化应加入引用存在性、字段完整性和规则检查，并把 reflection_note 明确注入下一轮分析。'
]),
('59. 有没有探索不同基座模型对 Reflection 的影响？',[
'我的回答是：RepoMind 当前没有形成严格的多模型 Reflection 对照实验，所以不能虚构结论。理论上，Actor 与 Critic 使用相同模型成本低、风格一致，但可能共享同类偏差；使用不同模型能增加判断多样性，却会提高成本和输出差异。',
'实验应固定分析结果和 Rubric，对多个 critic 比较与人工金标准的一致率、错误召回率、误触发重试率、延迟和费用。还要检查不同 issue_type 上是否稳定，不能只看总体平均。',
'如果更换 critic，我会先在 eval 数据集离线回放，不直接修改线上阈值。只有在低分识别和重试收益上稳定提升，才考虑替换。'
]),
('60. 为什么不直接用一个微调后的小模型做 Reflection？',[
'我的回答是：可以考虑，但前提是有足够高质量的批评样本和明确 Rubric。小模型成本和延迟更低，适合格式、字段和常见错误检查；但复杂代码因果、跨文件证据和替代根因判断可能超出能力。',
'更合理的是分层 critic：规则和小模型先检查 JSON、文件引用、严重度冲突等确定性问题；复杂根因再交给强模型。这样不会让昂贵模型处理所有简单校验。',
'RepoMind 当前使用通用 LLM critic，没有训练专用小模型。若要微调，应从人工确认的分析/批评对构建数据，并用独立仓库测试泛化，避免小模型只记住已有 Case。'
]),
('61. 为什么模型调用会失败？如何判断失败原因？',[
'我的回答是：失败要分网络基础设施、请求协议、模型输出和业务执行四层。网络层有超时、DNS、连接池、限流和 5xx；请求层有模型名、权限、Token 超限和 Tool Schema 错误；输出层有截断、非法 JSON 和字段缺失；业务层则是工具或后续节点失败。',
'每次调用应记录 trace_id、模型和 Prompt 版本、输入输出 Token、耗时、HTTP 状态、finish_reason、重试次数和错误类型。网络和限流可指数退避；参数与权限错误不应重复重试；Token 超限要裁剪或摘要。',
'RepoMind 的 LLM Client 使用 tenacity 和结构解析，但生产化仍应把“模型请求成功”和“业务任务成功”分开监控。模型成功返回 tool_call，并不代表工具执行成功。'
]),
('62. 模型输出合法 JSON，但业务仍解析失败，怎么排查？',[
'我的回答是：合法 JSON 只说明语法正确，不代表符合业务 Schema。需要依次检查字段类型、必填项、enum、数组长度、字段间约束、Schema 版本和流式截断。',
'RepoMind 对关键输出使用 Pydantic/结构校验，并对 confidence 做数值夹取。还应保留原始响应、解析错误路径、Prompt 与 Schema 版本，向模型返回结构化错误，而不是一句“格式不对”。',
'如果同一错误持续出现，优先简化 Schema、使用原生结构化输出或把确定性字段交给代码，不要无限增加“请严格输出 JSON”的 Prompt 和重试。'
]),
('63. Trace 应该保存哪些信息？如何兼顾可观测性和隐私？',[
'我的回答是：Trace 要能重建“输入经过哪些节点和工具变成最终答案”。至少保存 task/thread id、节点、状态转移、模型/Prompt 版本、Token、耗时、tool_call、参数摘要、结果摘要、错误、重试、confidence 和最终评测。',
'RepoMind 当前节点 wrapper 会打印节点及返回字段预览，State 和 Checkpoint 保存主要流程数据；要做生产诊断，还需要统一结构化 Trace 和跨 LLM、工具、检索的 span。',
'隐私方面，密钥绝不记录，用户数据和源码按权限脱敏或只保存哈希/引用；工具参数也应分字段处理。Trace 需要访问控制、保留期限和审计，不能为了排错永久保存全部敏感上下文。'
]),
('64. RepoMind 为什么使用异步 Job API？当前有什么问题？',[
'我的回答是：Issue 分析包含仓库拉取、索引、模型和多轮工具调用，可能持续较长时间，不适合让 HTTP 请求一直阻塞。Job API 接收 URL 后立即返回 job_id，后台执行 Graph，客户端轮询任务状态。',
'backend/app.py 使用 job_id 作为 thread_id，运行结束后保存 final_report、confidence 和 error。这样把提交与结果读取分离，也为超时、取消和进度展示提供接口基础。',
'当前问题是 jobs 使用进程内字典，多实例之间不共享，服务重启后 Job 元数据丢失，也缺少真正队列、租约和取消信号。生产化应使用数据库+消息队列+Worker，并依靠 LangGraph Checkpoint 恢复。'
]),
('65. 消息队列在 Agent 系统中有什么作用？为什么不直接用数据库通信？',[
'我的回答是：消息队列负责异步解耦、削峰、重试和任务分发；数据库负责持久化业务状态和查询。让 Worker 不断轮询数据库既增加负载，也缺少成熟的投递确认、消费组和死信机制。',
'RepoMind 当前没有 MQ，后台进程直接执行 Job。生产化可以在 API 写入 Job 后发送 task_id，Worker 消费并用 thread_id 执行 Graph；状态仍写数据库，消息只携带标识，避免大上下文在队列中复制。',
'需要接受至少一次投递并让消费者幂等。消息确认应在状态安全落库后完成，失败进入重试或死信队列，同一会话可按 thread_id 分区保证顺序。'
]),
('66. Agent 系统中的缓存一致性怎么保证？',[
'我的回答是：先区分缓存对象。GitHub 元数据、仓库索引、检索结果和最终答案的更新频率与风险不同，不能使用同一 TTL。缓存 Key 必须包含 repo、commit/version、模型或索引版本。',
'常用 cache-aside 是先更新数据源，再删除缓存；索引更新则适合构建新快照、校验完成后原子切换别名，旧索引保留回滚。删除文档需要向向量和关键词索引传播 Tombstone。',
'RepoMind 以当前仓库索引和 Chroma Collection 为主，尚未实现完整分布式缓存层。若缓存 RAG 结果，必须绑定 commit，否则仓库更新后会继续返回旧代码。'
]),
('67. 令牌桶、漏桶和滑动窗口限流有什么区别？Agent 场景怎么选？',[
'我的回答是：令牌桶按速率补充令牌，允许一定突发；漏桶以固定速率流出，输出平滑但突发请求需要排队或丢弃；滑动窗口统计最近时间段请求数，语义直观，但高精度实现存储成本更高。',
'Agent API 入口适合令牌桶，允许正常短突发并限制长期速率；昂贵模型或 CrossEncoder 可再加并发信号量；任务级还要限制 Token、工具轮数和总成本。',
'限流维度应包含用户、租户、模型和工具，而不是只有全局 QPS。高风险工具还需要更严格配额，不能因为入口请求少就允许一个任务无限调用。'
]),
('68. Agent 服务在高并发下怎么设计？',[
'我的回答是：先做容量估算并拆分 API、调度、Worker、模型网关、索引和存储。API 无状态水平扩展，长任务进入队列，Worker 按资源类型分池；模型和 CrossEncoder 受并发与批处理控制。',
'高频确定性请求优先缓存或规则直达，复杂任务才进入 Agent；设置超时、熔断、隔离、降级和优先级。仓库索引按 repo+commit 复用，避免每个请求重复 clone 和 embedding。',
'RepoMind 当前是演示型单机 Job API。可以说明已有并行 Send 和 Checkpoint 基础，但不能把它描述为百万 QPS 架构；生产化重点是队列、外置状态、索引服务和成本治理。'
]),
('69. 如何设计 Agent 的超时、熔断和降级策略？',[
'我的回答是：超时应分模型调用、工具调用、节点和任务总时限；熔断按依赖统计连续失败与错误率；降级则保证返回范围缩小但语义诚实的结果。',
'RepoMind 可以在 GitHub API 不可用时使用已有本地 clone，在 CrossEncoder 失败时退回 RRF 排名，在依赖检索失败时继续本仓库证据，在工具预算耗尽时输出低置信度和缺失说明。',
'降级不能伪装成完整成功。报告应标记哪些证据缺失，安全分析不能因超时自动放行。恢复后采用半开探测，避免大量任务同时冲击刚恢复的依赖。'
]),
('70. Agent 系统应该监控哪些工程指标？',[
'我的回答是：指标分业务成功、模型、检索、工具、工作流和资源成本。业务层看任务成功率、人工接受率和严重错误；模型层看延迟、Token、错误和结构化输出率；检索看 Recall@K、无结果率和 rerank 延迟。',
'工具层看选择准确率、参数错误率、成功/超时/重试；工作流看节点耗时、循环次数、Reflection 重试、Checkpoint 恢复和死信；资源层看 CPU/GPU、队列长度、并发和单任务成本。',
'RepoMind 已有 confidence、iteration、eval 分数和节点日志，但还需统一 metrics、logs、traces，并以 thread_id 串联。告警应对应可行动问题，例如某工具超时率突增，而不是只有总请求失败率。'
]),
]

B=RGBColor(46,116,181);D=RGBColor(31,77,120);G=RGBColor(95,103,112)
def sf(r,z=11,b=None,c=None):
 n='PingFang SC';r.font.name=n;r.font.size=Pt(z);rf=r._element.get_or_add_rPr().get_or_add_rFonts()
 for a in('ascii','hAnsi','eastAsia','cs'):rf.set(qn('w:'+a),n)
 if b is not None:r.bold=b
 if c is not None:r.font.color.rgb=c
doc=Document();s=doc.sections[0];s.page_width,s.page_height=Inches(8.5),Inches(11);s.top_margin=s.bottom_margin=Inches(.82);s.left_margin=s.right_margin=Inches(.88)
st=doc.styles['Normal'];st.font.name='PingFang SC';st.font.size=Pt(11);st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_after=Pt(6);st.paragraph_format.line_spacing=1.25
for nm,z,bf,af,c in [('Heading 1',16,16,8,B),('Heading 2',13,12,6,D)]:
 st=doc.styles[nm];st.font.name='PingFang SC';st.font.size=Pt(z);st.font.bold=True;st.font.color.rgb=c;st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_before=Pt(bf);st.paragraph_format.space_after=Pt(af);st.paragraph_format.keep_with_next=True
sf(s.header.paragraphs[0].add_run('RepoMind · Agent 开发面试问答'),9,c=G);f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;sf(f.add_run('第五批 · 第 51—70 题'),9,c=G)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(70);p.paragraph_format.space_after=Pt(10);sf(p.add_run('RepoMind Agent 开发面试问答'),24,True,D)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(35);sf(p.add_run('第五批：评测、可观测性与可靠性（20 题）'),14,c=B)
doc.add_paragraph('本批为第 51—70 题，覆盖 Agent 评测、Reflection、故障诊断、异步任务和高并发工程。');doc.add_page_break()
for q,ps in QA:
 doc.add_heading(q,level=1)
 for t in ps:p=doc.add_paragraph(t);p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第五批20题';doc.core_properties.author='RepoMind';doc.save(OUT);print(OUT)
