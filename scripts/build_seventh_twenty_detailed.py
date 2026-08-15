from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches,Pt,RGBColor
OUT=Path('/Users/msy/Desktop/github/RepoMind/RepoMind_Agent开发面试问答_第七批20题.docx')
QA=[
('91. RepoMind 的提示词模板是怎么构建的？上下文工程有哪些实践？',[
'我的回答是：Prompt 不应该把所有规则、代码和历史材料堆成一段长文本，而要围绕当前节点的单一职责组织。RepoMind 按 Issue 类型和 PR Review 类型拆分 Prompt，让模型明确目标、可用证据、工具集合、输出 Schema 和停止要求。',
'bug Prompt 关注根因、调用链、回归和修复类型；feature 关注功能是否已存在、改动位置和兼容性；question 强调基于现有实现解释；PR 的 security、logic、style 使用不同检查维度。公共约束只保留证据优先、不得编造文件、结构化输出等内容。',
'上下文按来源分块：Issue 信息、project_map、memory_context、entry/RAG/Git 证据和 reflection_note。工具结果在 ReAct 消息链中按调用顺序返回。这样模型能区分任务、历史案例和当前代码证据，也便于发现哪一块导致 Badcase。'
]),
('92. 做过 Query Rewrite 吗？需要用户补充信息时怎么实现？',[
'我的回答是：RepoMind 的多查询构造可以视为受控 Query Rewrite：从 Issue 生成 HyDE、标题 Query、类型+正文 Query 和堆栈 Query。它不是把原问题随意改写成一个新问题，而是保留多种表达并通过 RRF 融合。',
'如果改写需要用户补充信息，我会让模型输出结构化 missing_fields、reason 和 question，再由规则判断是否真的阻塞。仓库 URL、目标分支、版本和写操作范围缺失时必须澄清；可以从 GitHub API 或项目配置确定的信息不应反复询问用户。',
'RepoMind 当前以标准 GitHub URL 为入口，还没有完整交互式澄清节点。已有 human_gate 和 Checkpoint 可以支撑中断/恢复，但要扩展多轮澄清还需保存 pending_question 与用户回答，并从相同 thread_id 继续。'
]),
('93. 什么是并行化意图识别？为什么要做？RepoMind 使用了吗？',[
'我的回答是：并行化意图识别是让多个独立分类器或视角同时判断任务属性，例如主意图、风险、是否需要检索和是否需要工具，再由规则合并。它适合属性相互独立且单次识别成为延迟瓶颈的场景。',
'RepoMind 的 route_issue 当前主要一次输出 bug、feature、question、security，并没有把意图识别拆成多个并行模型调用。项目的真实并行发生在分类之后：Issue 上下文 entry/RAG/Git 三路并行，PR security/style/logic 三路并行。',
'若扩展并行识别，我会把“业务类型”和“风险等级”分开，因为 security 可能同时具有 bug 特征。合并必须由确定性规则处理冲突，并评估额外模型调用是否真的降低总延迟，而不是为了并行增加成本。'
]),
('94. Agent 配置十多个工具时，如何系统性降低选错工具的概率？',[
'我的回答是：核心不是继续写更长的工具描述，而是减少每一步的候选集合。先根据节点和任务类型做工具路由，只暴露当前可能需要的能力；其次消除功能重叠，统一名称、参数和错误返回。',
'RepoMind 的 analyze_bug 暴露代码导航、Git 历史和依赖工具；feature/question 的集合更小；PR style 主要使用 read_file、grep_code、list_dir，logic/security 才增加定义与调用者。模型不需要在每一轮从全部工具中选择。',
'还应离线评测 tool selection accuracy，收集误选对并改进 description 的适用/禁用条件。模型连续选择无效工具时由代码检测停滞，而不是允许它靠随机试错找到正确工具。'
]),
('95. Tool 执行结果不确定时应该怎么处理？',[
'我的回答是：工具返回“超时”或连接断开，并不说明动作一定没有发生。读取工具通常可以重试；写工具必须进入 UNKNOWN 状态，优先调用状态查询或使用幂等键确认，不能直接重复执行。',
'RepoMind 当前工具以读取源码、搜索和 Git 查询为主，不确定结果主要影响证据完整性。系统应把 tool_status、error_type 和 partial_result 明确返回模型；如果关键证据缺失，降低 confidence，并在报告中说明。',
'未来自动创建 PR 或修改文件时，需要 prepare/execute/query 三阶段接口。恢复任务时先检查外部 operation_id，再决定返回旧结果、继续等待、补偿还是人工介入。'
]),
('96. 只读、低风险写和高风险写工具如何分级？',[
'我的回答是：只读工具只获取信息，例如 read_file、grep_code、git_log；低风险写工具影响范围有限且容易回滚，例如在临时分支写文件；高风险写工具包括删除、合并、发布、转账和修改权限。',
'分级后应设置不同策略：只读可以自动执行但仍受作用域限制；低风险写要求幂等、diff 预览和自动回滚；高风险写必须二次确认、权限校验、审计和两阶段提交。',
'RepoMind 当前核心工具基本属于只读，security Issue 还经过 human_gate。若增加自动修复，应该默认只在隔离工作树生成 Patch，先让用户审阅，再允许创建 PR，不能直接写目标分支。'
]),
('97. 工具版本升级后，运行中的 Agent 任务怎么办？',[
'我的回答是：Checkpoint 只保存状态还不够，运行中任务还依赖当时的工具协议。若 Schema、语义或返回格式变化，恢复后用新工具继续可能破坏消息链和业务结果。',
'应为每个工具记录 tool_name、semantic_version、Schema hash 和执行环境版本，并把 workflow_version 写入任务。兼容升级可以由新版本读取旧参数；不兼容升级则保留旧 Worker，或通过显式迁移转换待执行 State。',
'RepoMind 当前 ToolRegistry 没有完整版本治理，SQLite Checkpoint 也主要面向同版本恢复。生产化需要版本化 Graph 和 Tool Schema，不能假设所有历史任务都能被最新代码直接继续。'
]),
('98. MCP Server 如何做认证、授权和审计？',[
'我的回答是：认证解决调用者是谁，授权解决能访问什么，审计记录实际做了什么。MCP Server 不应因为连接来自 Agent 就默认信任，客户端必须携带用户或服务身份，并把身份传递到每次工具调用。',
'授权可以结合 RBAC 与资源级策略：用户是否能调用该工具、能访问哪个仓库、分支和字段。高风险能力需要短期委托凭证和明确 scope，不能把长期管理员密钥交给模型。',
'审计至少保存主体、工具版本、参数摘要、目标资源、决策结果、耗时和返回状态，并对敏感字段脱敏。RepoMind 当前未运行独立 MCP Server，这属于若把 ToolRegistry 服务化后必须补齐的能力。'
]),
('99. 如何防范恶意 MCP Server 和数据外泄？',[
'我的回答是：MCP Server 本身是外部信任边界，可能返回恶意指令、夸大能力或收集传入数据。客户端不能仅凭工具描述就把用户上下文、源码和凭证全部发送过去。',
'防护包括 Server allowlist、签名/来源验证、最小权限、网络隔离、出站域名限制、参数数据分类和返回内容按不可信数据处理。调用前根据工具声明和数据敏感度决定是否允许，并记录数据流向。',
'对于 RepoMind，私有仓库源码不能发送给未知依赖服务。即使未来接入 MCP，也应优先本地代码导航服务，第三方 Server 只接收完成任务所需的最小片段。'
]),
('100. A2A 解决什么问题？它与 MCP 在多 Agent 系统中如何分工？',[
'我的回答是：MCP 主要标准化 Agent/客户端如何发现和调用工具、资源与 Prompt；A2A 更关注 Agent 与 Agent 之间如何发现能力、委派任务、交换状态和返回结果。一个连接能力，一个连接协作者。',
'例如 RepoMind Orchestrator 如果服务化，可以通过 A2A 把不同仓库任务委派给 Repo Agent；每个 Repo Agent 内部再通过 MCP 使用代码搜索、Git 和依赖工具。Agent 间传递应是结构化任务和结果，不是无限自然语言聊天。',
'项目当前的跨仓库协作使用 LangGraph Send 和 HTTP RepoAgent 服务，并未实现标准 A2A。面试时可以说明协议分工和演进方向，但不能把现有内部调用直接称为完整 A2A。'
]),
('101. 知识库更新后，缓存、向量索引和答案如何保持一致？',[
'我的回答是：知识更新涉及原始文档、Chunk、embedding、关键词索引和缓存，不能只覆盖其中一层。每份文档需要版本，Chunk 继承版本，缓存 Key 也包含索引快照或 commit。',
'更新时先构建新索引快照并做数量、抽样召回和权限校验，完成后原子切换 alias；旧索引保留一段时间用于回滚。删除需要传播 Tombstone，查询层在异步清理完成前过滤失效版本。',
'RepoMind 的仓库索引应绑定 owner/repo/commit，而不是只按仓库名复用。当前版本治理仍较轻，生产化必须避免用最新索引回答针对旧 Issue 版本的问题。'
]),
('102. 向量数据库数据不断增长，查询为什么不一定线性变慢？',[
'我的回答是：向量数据库通常使用 HNSW、IVF 等近似最近邻索引，不需要每次与所有向量暴力比较。查询复杂度更多受索引结构、efSearch、nprobe、维度、过滤条件和内存布局影响。',
'数据增长仍会带来构建时间、内存、缓存命中和索引维护压力；高选择性 metadata filter 如果不能下推，也可能先搜索大量无效候选。分片过多还会增加结果合并成本。',
'RepoMind 按仓库管理 Collection，本身就是缩小搜索空间的一种方式。继续扩展时应按租户/仓库隔离、监控 P95 延迟与 Recall，而不是只看平均查询时间。'
]),
('103. BM25 和向量分数不在同一量纲，混合检索怎么融合？',[
'我的回答是：不能直接把 BM25 分数和余弦相似度相加，因为分布随 Query、索引和实现变化。无标注数据时最稳妥的是 RRF，只使用各路排名，通过 1/(k+rank) 累加。',
'如果有标注数据，可以训练 Learning to Rank，把 BM25、向量相似度、文件类型、仓库版本、路径匹配和新鲜度作为特征。硬约束如 repo、commit 和权限应先过滤，不能只作为弱排序特征。',
'RepoMind 已使用 RRF 融合多路向量 Query，但尚未加入独立 BM25。后续接入 BM25 时可以复用 RRF，再让 CrossEncoder 统一精排。'
]),
('104. 文档更新后新旧向量同时被召回，怎么办？',[
'我的回答是：根因通常是更新采用 append 而没有失效旧 Chunk，或者异步索引切换期间查询同时访问两版。仅依赖内容相似度去重不可靠，因为旧版和新版可能只有一行差异。',
'应为 document_id、chunk_id 和 version 建稳定关系。新版本完成后原子切换 active_version，查询强制过滤；旧 Chunk 标记 tombstone 后异步删除。结果层按逻辑 chunk_id 只保留当前版本。',
'代码场景还要绑定 commit。RepoMind 如果分析历史 Issue，可能有意使用旧 commit；因此不是永远只保留最新，而是根据任务版本选择一致快照。'
]),
('105. 多租户 RAG 如何做文档权限过滤？',[
'我的回答是：权限过滤必须在检索层执行，不能先召回所有租户内容再让模型忽略。每个 Chunk 保存 tenant_id、resource_id、ACL 或可推导的权限标签，查询使用当前用户身份生成强制 Filter。',
'向量库不支持复杂 ACL 时，可以先得到允许资源集合再检索，或按租户/安全域分 Collection。缓存 Key 必须包含权限上下文，避免 A 用户的结果被 B 用户复用。',
'RepoMind 当前主要分析用户明确指定的公开仓库，没有完整多租户授权。若支持私有仓库，需要把 GitHub token 对应权限传入 loader、检索和工具层，并禁止 Memory 跨租户召回。'
]),
('106. 代码 Agent 的文件、命令和网络沙箱怎么设计？',[
'我的回答是：沙箱应默认拒绝，只开放任务所需目录、命令和域名。文件系统使用独立临时工作树，限制路径规范化后的根目录；命令使用 allowlist、资源限额和无特权用户；网络只允许必要 API。',
'执行环境设置 CPU、内存、进程数、磁盘、时间和输出上限，敏感挂载与宿主凭证不可见。写操作生成 Patch 或提交到临时分支，任务结束后可销毁环境；所有命令和文件变更进入审计。',
'RepoMind 当前工具以 GitHub API和本地读取为主，没有开放任意 shell Agent，所以风险较低。但未来自动修复不能直接复用开发机权限，必须单独隔离。'
]),
('107. 敏感数据进入模型前如何识别和脱敏？',[
'我的回答是：先进行数据分类，识别密钥、Token、邮箱、手机号、身份证、数据库连接串和业务敏感字段。检测可结合正则、熵、密钥模式、字段名和专用分类器。',
'脱敏策略取决于任务：完全无关的字段删除；需要保持关联的值使用稳定占位符；必须参与计算的数据尽量在本地工具处理，只把结果交给模型。日志和 Trace 同样需要脱敏。',
'RepoMind 分析仓库时尤其要防止 .env、证书和测试密钥进入索引。loader 应默认排除敏感文件，并在 Chunk 写入和工具返回两个边界重复扫描。'
]),
('108. Text2SQL 如何做 AST 校验、只读限制、成本控制和脱敏？',[
'我的回答是：模型生成 SQL 只能作为候选。服务端先用 SQL Parser 生成 AST，拒绝 INSERT、UPDATE、DELETE、DDL、多语句、危险函数和未授权表；再按用户身份注入行列权限。',
'执行前使用 EXPLAIN 估算扫描行数和成本，强制 LIMIT、超时和只读事务，连接使用只读账号。结果返回模型前对敏感列脱敏并限制行数，防止通过宽查询外泄。',
'RepoMind 不包含 Text2SQL，这属于通用工具安全题。可以借鉴同一原则：模型负责建议，确定性网关负责授权、静态校验和执行。'
]),
('109. 浏览器 Agent 的高风险操作如何确认？',[
'我的回答是：浏览、搜索和读取属于低风险；发送消息、购买、发布、删除、授权和提交表单属于高风险。确认应发生在最终动作前，并展示目标、关键参数和不可逆影响。',
'采用 prepare/confirm/execute：Agent 先准备操作和预览，用户确认后生成一次性授权 Token，执行层验证 Token 与参数摘要一致，防止确认后模型偷偷更改目标。',
'RepoMind 没有浏览器 Agent，但 security human_gate 展示了同样思想。若未来自动创建 GitHub PR，也应先展示 diff、仓库和分支，再由用户确认执行。'
]),
('110. Agent 系统如何管理密钥和用户授权委托？',[
'我的回答是：模型上下文中不能出现长期密钥。凭证由服务端 Secret Manager 保存，工具执行时按当前用户和 scope 获取短期 Token；模型只能看到凭证引用或能力是否可用。',
'用户授权采用最小权限和短有效期，区分读取仓库、写分支、创建 PR 等 scope，并支持撤销。工具日志记录凭证 ID 而不是密钥值，错误信息也要防止回显。',
'RepoMind 当前通过配置使用 GitHub/API 凭证。生产化应从环境变量进一步升级到密钥管理服务，并避免私有仓库 Token 被写入 Checkpoint、Memory 或 Trace。'
]),
]
B=RGBColor(46,116,181);D=RGBColor(31,77,120);G=RGBColor(95,103,112)
def sf(r,z=11,b=None,c=None):
 n='PingFang SC';r.font.name=n;r.font.size=Pt(z);rf=r._element.get_or_add_rPr().get_or_add_rFonts()
 for a in('ascii','hAnsi','eastAsia','cs'):rf.set(qn('w:'+a),n)
 if b is not None:r.bold=b
 if c is not None:r.font.color.rgb=c
doc=Document();s=doc.sections[0];s.page_width,s.page_height=Inches(8.5),Inches(11);s.top_margin=s.bottom_margin=Inches(.82);s.left_margin=s.right_margin=Inches(.88)
st=doc.styles['Normal'];st.font.name='PingFang SC';st.font.size=Pt(11);st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_after=Pt(6);st.paragraph_format.line_spacing=1.25
for nm,z,bf,af,c in [('Heading 1',16,16,8,B),('Heading 2',13,12,6,D)]:
 st=doc.styles[nm];st.font.name='PingFang SC';st.font.size=Pt(z);st.font.bold=True;st.font.color.rgb=c;st._element.rPr.rFonts.set(qn('w:eastAsia'),'PingFang SC');st.paragraph_format.space_before=Pt(bf);st.paragraph_format.space_after=Pt(af);st.paragraph_format.keep_with_next=True
sf(s.header.paragraphs[0].add_run('RepoMind · Agent 开发面试问答'),9,c=G);f=s.footer.paragraphs[0];f.alignment=WD_ALIGN_PARAGRAPH.RIGHT;sf(f.add_run('第七批 · 第 91—110 题'),9,c=G)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(70);p.paragraph_format.space_after=Pt(10);sf(p.add_run('RepoMind Agent 开发面试问答'),24,True,D)
p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(35);sf(p.add_run('第七批：工具治理、知识一致性与安全（20 题）'),14,c=B)
doc.add_paragraph('本批为第 91—110 题，覆盖 Prompt、工具运行时、MCP/A2A、知识库版本、多租户和安全边界。');doc.add_page_break()
for q,ps in QA:
 doc.add_heading(q,level=1)
 for t in ps:p=doc.add_paragraph(t);p.paragraph_format.keep_together=True
doc.core_properties.title='RepoMind Agent 开发面试问答：第七批20题';doc.core_properties.author='RepoMind';doc.save(OUT);print(OUT)
